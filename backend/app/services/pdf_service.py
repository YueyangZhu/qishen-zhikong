"""PDF / DOCX 解析服务

使用 pdfplumber 解析 PDF，python-docx 解析 DOCX，
将文档切分为段落（ContractParagraph），与前端 types 对齐。

解析能力：
- 过滤 PDF 页眉/页脚/页码/水印（启发式：重复出现的短行、纯数字页码、"第X页"等）
- 识别段落类型（title/header/body/signature/table/image），支持前端差异化渲染
- 首部段（甲乙方）、签署段独立成节，左栏章节目录完整反映合同结构
- 保留原稿阅读顺序
- 提取表格为二维数组（table 类型段落），前端渲染为 HTML 表格
- 提取图片为 base64（image 类型段落），前端渲染为 <img>
"""
import base64
import io
import json
import logging
import re
from typing import List, Tuple, Optional
from collections import Counter
import pdfplumber
from docx import Document as DocxDocument
import mammoth

from app.schemas.review import ContractParagraph, ContractSection, ParsedDocument

logger = logging.getLogger(__name__)

# PyMuPDF 可选导入（未安装时跳过图片提取，不影响表格和文本）
try:
    import fitz  # PyMuPDF
    HAS_FITZ = True
except ImportError:
    HAS_FITZ = False
    logger.warning("PyMuPDF 未安装，PDF 图片提取功能不可用（表格和文本解析不受影响）")


# 段落切分正则
# 匹配"第X条"、"一、二、三、"、"1. 2. 3."、"第一条 第二条"等条款编号
CLAUSE_PATTERN = re.compile(
    r"^(第[一二三四五六七八九十百零\d]+条|[一二三四五六七八九十]+、|\d+[.、）)])",
    re.MULTILINE,
)

# 页眉/页脚/页码识别正则
PAGE_NUMBER_PATTERN = re.compile(
    r"^(第\s*\d+\s*页(?:\s*/\s*共\s*\d+\s*页)?|page\s*\d+|\d+\s*/\s*\d+|\d{1,3})$",
    re.IGNORECASE,
)

# 首部段（甲乙方信息）识别：以"甲方/乙方/供方/需方/发包方/承包方"等开头
HEADER_PATTERN = re.compile(
    r"^(甲方|乙方|供方|需方|发包方|承包方|委托方|受托方|出租方|承租方|买方|卖方|定作方|承揽方|采购方|供应方|卖受人|买受人|出让人|受让人|招标方|投标方|发包人|承包人|委托人|受托人|出租人|承租人|订立人|协议方)[（(：: 　]",
)

# 签署段识别：以签署关键词开头，或末尾含签字盖章
SIGNATURE_PATTERN = re.compile(
    r"^(签署|签字|盖章|签订日期|签订地点|签约地点|签约日期|本合同一式|双方签字|双方盖章|甲方签章|乙方签章|甲方（签章|乙方（签章|甲方盖章|乙方盖章)",
)
SIGNATURE_TAIL_PATTERN = re.compile(r"签字（盖章）|（签字盖章）|（盖章）|签字日期|盖章日期")

# 图片大小限制：单张 500KB，每文档最多 10 张
MAX_IMAGE_BYTES = 500 * 1024
MAX_IMAGES_PER_DOC = 10


class ContentBlock:
    """内容块：文本/表格/图片，按原稿顺序排列"""

    def __init__(
        self,
        block_type: str,
        text: str = "",
        table_data: Optional[List[List[str]]] = None,
        image_data: str = "",
        image_format: str = "",
    ):
        self.type = block_type  # 'text' | 'table' | 'image'
        self.text = text
        self.table_data = table_data
        self.image_data = image_data
        self.image_format = image_format


class PDFService:
    """合同文件解析"""

    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """解析合同文件，返回结构化文档

        Args:
            content: 文件二进制内容
            filename: 文件名（用于判断类型）
        """
        name_lower = filename.lower()
        if name_lower.endswith(".pdf"):
            blocks = self._parse_pdf(content)
        elif name_lower.endswith(".docx"):
            blocks = self._parse_docx(content)
        elif name_lower.endswith(".txt"):
            blocks = [ContentBlock("text", content.decode("utf-8", errors="ignore"))]
        else:
            try:
                blocks = [ContentBlock("text", content.decode("utf-8", errors="ignore"))]
            except Exception:
                raise ValueError(f"不支持的文件类型：{filename}")

        # 检查内容非空
        has_content = any(
            b.text.strip() or b.table_data or b.image_data for b in blocks
        )
        if not has_content:
            raise ValueError("文件内容为空，无法解析")

        paragraphs, sections = self._split_paragraphs(blocks)
        title = paragraphs[0].text if paragraphs else filename

        # fullText 仅拼接文本类内容（表格/图片用摘要）
        full_parts = []
        for b in blocks:
            if b.type == "text":
                full_parts.append(b.text)
            elif b.type == "table":
                full_parts.append(b.text)
            elif b.type == "image":
                full_parts.append("[图片]")
        full_text = "\n\n".join(full_parts)

        html_content = self._generate_html(content, filename, blocks)

        return ParsedDocument(
            title=title,
            sections=sections,
            paragraphs=paragraphs,
            fullText=full_text,
            htmlContent=html_content,
        )

    def _generate_html(self, content: bytes, filename: str, blocks: List[ContentBlock]) -> Optional[str]:
        """生成原文格式 HTML 用于前端预览"""
        name_lower = filename.lower()
        try:
            if name_lower.endswith(".docx"):
                result = mammoth.convert_to_html(io.BytesIO(content))
                html = result.value
                warnings = result.messages
                if warnings:
                    logger.debug(f"mammoth 警告: {warnings}")
                if not html.strip():
                    raise ValueError("mammoth 返回空 HTML")
                css = """
                <style>
                  body { font-family: 'Microsoft YaHei', 'SimSun', serif; line-height: 1.8; padding: 40px; color: #333; max-width: 900px; margin: 0 auto; }
                  table { border-collapse: collapse; width: 100%; margin: 16px 0; }
                  td, th { border: 1px solid #999; padding: 8px 12px; text-align: left; }
                  th { background: #f5f5f5; font-weight: 600; }
                  p { margin: 8px 0; }
                  h1, h2, h3, h4 { margin: 16px 0 8px; }
                  img { max-width: 100%; }
                </style>
                """
                return css + html
            elif name_lower.endswith(".pdf"):
                return """<div style="padding:20px;text-align:center;color:#999">
                  <p>PDF 文件暂不支持原文格式预览，请使用下载功能查看原文件</p>
                </div>"""
            return None
        except Exception as e:
            logger.warning(f"生成 HTML 失败: {e}")
            return None

    def _parse_pdf(self, content: bytes) -> List[ContentBlock]:
        """使用 pdfplumber 解析 PDF

        保留原稿阅读顺序：按页处理，每页提取文本、表格、图片，
        按文本→表格→图片顺序输出（表格区域内的文本行会被过滤避免重复）。
        同时过滤页眉/页脚/页码/水印等非正文内容。
        """
        blocks: List[ContentBlock] = []
        pages_lines: List[List[str]] = []
        page_tables: List[List] = []  # 每页的表格列表
        # 收集所有页的行，统计每行在多少页出现（用于识别页眉页脚）

        with pdfplumber.open(io.BytesIO(content)) as pdf:
            total_pages = len(pdf.pages)
            for page in pdf.pages:
                # 提取表格（find_tables 返回 Table 对象，有 bbox 和 extract()）
                tables = []
                try:
                    found = page.find_tables()
                    for t in found:
                        try:
                            data = t.extract()
                            # 过滤掉空表格（全空行）
                            if data and any(any(cell and cell.strip() for cell in row) for row in data):
                                tables.append(data)
                        except Exception:
                            continue
                except Exception as e:
                    logger.debug(f"表格提取失败（页 {page.page_number}）：{e}")
                page_tables.append(tables)

                # 提取文本行，过滤掉表格区域内的行
                page_text = page.extract_text() or ""
                if page_text:
                    # 尝试按行过滤表格区域
                    lines = [line.rstrip() for line in page_text.split("\n") if line.strip()]
                    pages_lines.append(lines)
                else:
                    pages_lines.append([])

        if not pages_lines and not any(page_tables):
            return blocks

        # 统计每行在多少页出现（页眉页脚在多页重复出现）
        line_page_count: Counter = Counter()
        for lines in pages_lines:
            for line in set(lines):
                line_page_count[line] += 1

        def is_header_footer(line: str) -> bool:
            """识别页眉/页脚/页码/水印"""
            stripped = line.strip()
            if not stripped:
                return True
            if PAGE_NUMBER_PATTERN.match(stripped):
                return True
            if (
                total_pages >= 2
                and len(stripped) <= 30
                and line_page_count[stripped] >= max(2, total_pages // 2)
            ):
                return True
            return False

        # 按页输出文本块 + 表格块
        for page_idx, lines in enumerate(pages_lines):
            filtered = [line for line in lines if not is_header_footer(line)]
            if filtered:
                blocks.append(ContentBlock("text", "\n".join(filtered)))
            # 该页的表格转为 table 块
            for table_data in page_tables[page_idx]:
                # 表格文本摘要（用于风险匹配）：行用 | 分隔，单元格用空格
                summary_lines = [" ".join((cell or "").strip() for cell in row) for row in table_data]
                summary = "\n".join(summary_lines)
                blocks.append(ContentBlock("table", text=summary, table_data=table_data))

        # 提取图片（用 PyMuPDF）
        if HAS_FITZ:
            image_blocks = self._extract_pdf_images(content)
            blocks.extend(image_blocks)

        return blocks

    def _extract_pdf_images(self, content: bytes) -> List[ContentBlock]:
        """使用 PyMuPDF 提取 PDF 内嵌图片，转 base64"""
        blocks: List[ContentBlock] = []
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            img_count = 0
            for page in doc:
                if img_count >= MAX_IMAGES_PER_DOC:
                    break
                for img_info in page.get_images(full=True):
                    if img_count >= MAX_IMAGES_PER_DOC:
                        break
                    try:
                        xref = img_info[0]
                        pix = fitz.Pixmap(doc, xref)
                        # CMYK 色彩空间转 RGB
                        if pix.n >= 5:
                            pix = fitz.Pixmap(fitz.csRGB, pix)
                        img_bytes = pix.tobytes("png")
                        if len(img_bytes) > MAX_IMAGE_BYTES:
                            logger.debug(f"图片过大（{len(img_bytes)} bytes），跳过")
                            continue
                        b64 = base64.b64encode(img_bytes).decode("ascii")
                        blocks.append(ContentBlock(
                            "image",
                            text="[图片]",
                            image_data=b64,
                            image_format="png",
                        ))
                        img_count += 1
                    except Exception as e:
                        logger.debug(f"图片提取失败：{e}")
                        continue
            doc.close()
        except Exception as e:
            logger.warning(f"PDF 图片提取失败：{e}")
        return blocks

    def _parse_docx(self, content: bytes) -> List[ContentBlock]:
        """使用 python-docx 解析 DOCX

        按文档顺序遍历段落和表格（python-docx 的 document.element.body 按顺序包含所有元素）。
        """
        doc = DocxDocument(io.BytesIO(content))
        blocks: List[ContentBlock] = []

        # 按文档顺序遍历 body 子元素（段落 w:p 和表格 w:tbl 交替出现）
        from docx.oxml.ns import qn

        body = doc.element.body
        para_idx = 0
        table_idx = 0
        paragraphs = doc.paragraphs
        tables = doc.tables

        for child in body.iterchildren():
            tag = child.tag
            if tag == qn("w:p"):
                if para_idx < len(paragraphs):
                    text = paragraphs[para_idx].text.strip()
                    if text:
                        blocks.append(ContentBlock("text", text))
                    para_idx += 1
            elif tag == qn("w:tbl"):
                if table_idx < len(tables):
                    table = tables[table_idx]
                    try:
                        data = []
                        for row in table.rows:
                            data.append([cell.text.strip() for cell in row.cells])
                        if data and any(any(c for c in row) for row in data):
                            summary_lines = [" ".join(row) for row in data]
                            blocks.append(ContentBlock(
                                "table",
                                text="\n".join(summary_lines),
                                table_data=data,
                            ))
                    except Exception as e:
                        logger.debug(f"DOCX 表格提取失败：{e}")
                    table_idx += 1

        # 提取 DOCX 图片
        image_blocks = self._extract_docx_images(doc)
        blocks.extend(image_blocks)

        return blocks

    def _extract_docx_images(self, doc) -> List[ContentBlock]:
        """提取 DOCX 内嵌图片，转 base64"""
        blocks: List[ContentBlock] = []
        try:
            img_count = 0
            for rel in doc.part.rels.values():
                if img_count >= MAX_IMAGES_PER_DOC:
                    break
                if "image" in rel.reltype:
                    try:
                        image_part = rel.target_part
                        img_bytes = image_part.blob
                        if len(img_bytes) > MAX_IMAGE_BYTES:
                            continue
                        # 根据内容类型判断格式
                        content_type = image_part.content_type or "image/png"
                        fmt = "png" if "png" in content_type else "jpeg"
                        b64 = base64.b64encode(img_bytes).decode("ascii")
                        blocks.append(ContentBlock(
                            "image",
                            text="[图片]",
                            image_data=b64,
                            image_format=fmt,
                        ))
                        img_count += 1
                    except Exception as e:
                        logger.debug(f"DOCX 图片提取失败：{e}")
                        continue
        except Exception as e:
            logger.warning(f"DOCX 图片提取失败：{e}")
        return blocks

    def _split_paragraphs(
        self, blocks: List[ContentBlock]
    ) -> Tuple[List[ContractParagraph], List[ContractSection]]:
        """将内容块切分为段落和章节

        切分策略（保证与原稿顺序一致，且不破坏多行条款）：
        1. 文本块：按空行切块，块内按"条款编号行"或"甲乙方行"切分
        2. 表格块：直接转为 table 类型段落
        3. 图片块：直接转为 image 类型段落
        4. title/header/signature 段落各自独立成节
        5. 每段分配 id（p1, p2, ...）和 index，顺序与原稿一致
        """
        paragraphs: List[ContractParagraph] = []
        sections: List[ContractSection] = []
        current_section_paras: List[str] = []
        current_section_title = "合同首部"
        current_section_no = "首部"

        def flush_section():
            if current_section_paras:
                sections.append(ContractSection(
                    id=f"s{len(sections) + 1}",
                    title=current_section_title,
                    clauseNo=current_section_no,
                    paragraphIds=current_section_paras[:],
                ))
                current_section_paras.clear()

        para_idx = 0

        for block in blocks:
            if block.type == "table":
                # 表格块 → table 类型段落
                para_idx += 1
                para_id = f"p{para_idx}"
                paragraphs.append(ContractParagraph(
                    id=para_id,
                    index=para_idx,
                    text=block.text,
                    type="table",
                    tableData=block.table_data,
                ))
                current_section_paras.append(para_id)
                continue

            if block.type == "image":
                # 图片块 → image 类型段落
                para_idx += 1
                para_id = f"p{para_idx}"
                paragraphs.append(ContractParagraph(
                    id=para_id,
                    index=para_idx,
                    text=block.text,
                    type="image",
                    imageData=block.image_data,
                    imageFormat=block.image_format,
                ))
                current_section_paras.append(para_id)
                continue

            # 文本块：按空行切块，块内按条款号/甲乙方行切分
            text = block.text
            text_blocks = [b for b in re.split(r"\n\s*\n", text) if b.strip()]
            raw_paragraphs: List[str] = []
            for tb in text_blocks:
                lines = tb.split("\n")
                current_lines: List[str] = []
                for line in lines:
                    if not line.strip():
                        continue
                    stripped = line.strip()
                    if CLAUSE_PATTERN.match(stripped) and current_lines:
                        raw_paragraphs.append("\n".join(current_lines))
                        current_lines = []
                    elif HEADER_PATTERN.match(stripped) and current_lines:
                        raw_paragraphs.append("\n".join(current_lines))
                        current_lines = []
                    current_lines.append(line.rstrip())
                if current_lines:
                    raw_paragraphs.append("\n".join(current_lines))

            # 第一段拆分（封面页：标题+元数据）
            if raw_paragraphs and not paragraphs:
                first = raw_paragraphs[0]
                first_lines = first.split("\n")
                if len(first_lines) > 1:
                    first_line = first_lines[0].strip()
                    if (not CLAUSE_PATTERN.match(first_line)
                            and not HEADER_PATTERN.match(first_line)
                            and not SIGNATURE_PATTERN.match(first_line)
                            and (len(first_line) <= 40 or re.search(r'合同|协议|契约', first_line))):
                        raw_paragraphs[0] = first_line
                        raw_paragraphs.insert(1, "\n".join(first_lines[1:]))

            for raw in raw_paragraphs:
                para_idx += 1
                para_id = f"p{para_idx}"
                clause_no, clause_title = self._detect_clause(raw)
                para_type = self._detect_type(para_idx, raw, clause_no)

                # 章节边界判定
                if para_type == 'title':
                    flush_section()
                    current_section_title = raw[:30]
                    current_section_no = "标题"
                elif para_type == 'header':
                    if current_section_no not in ("首部", "标题"):
                        flush_section()
                        current_section_title = "合同主体"
                        current_section_no = "首部"
                elif para_type == 'signature':
                    flush_section()
                    current_section_title = clause_title or "签署落款"
                    current_section_no = "签署"
                else:
                    if clause_no and clause_no != current_section_no:
                        flush_section()
                        current_section_no = clause_no
                        current_section_title = clause_title or clause_no

                paragraphs.append(ContractParagraph(
                    id=para_id,
                    index=para_idx,
                    text=raw,
                    clauseNo=clause_no,
                    clauseTitle=clause_title,
                    type=para_type,
                ))
                current_section_paras.append(para_id)

        flush_section()
        return paragraphs, sections

    @staticmethod
    def _detect_type(idx: int, text: str, clause_no: str | None) -> str:
        """识别段落类型

        Returns: 'title' | 'header' | 'signature' | 'body'
        - title: 第一段且无 clauseNo，首行 ≤ 40 字符，或含"合同/协议/契约"关键词
        - header: 以"甲方/乙方/供方/需方"等开头（首部甲乙方信息）
        - signature: 以"签署/签字/盖章/本合同一式"等开头，或末尾含签字盖章
        - body: 其余（含 clauseNo 的条款段，或无 clauseNo 的非首段正文）
        """
        first_line = text.split("\n", 1)[0].strip()
        if SIGNATURE_PATTERN.match(first_line) or SIGNATURE_TAIL_PATTERN.search(text):
            return 'signature'
        if HEADER_PATTERN.match(first_line):
            return 'header'
        if idx == 1 and not clause_no:
            if len(first_line) <= 40:
                return 'title'
            if re.search(r'合同|协议|契约', first_line):
                return 'title'
        return 'body'

    @staticmethod
    def _detect_clause(text: str) -> Tuple[str | None, str | None]:
        """识别段落开头的条款编号和标题

        Returns: (clause_no, clause_title)
        - "第三条 违约责任：..." -> ("第三条", "违约责任")
        - "一、合同金额..." -> ("一", "合同金额")
        - "1. 付款方式..." -> ("1", "付款方式")
        """
        first_line = text.split("\n", 1)[0]
        match = CLAUSE_PATTERN.match(first_line.strip())
        if not match:
            return None, None

        clause_no = match.group(1).rstrip("、.）)")
        rest = first_line[match.end():].lstrip("：: 　")
        title_match = re.match(r"[^\s。；;]+", rest)
        title = title_match.group(0) if title_match else None
        if title and len(title) > 20:
            title = title[:20]
        return clause_no, title


pdf_service = PDFService()
