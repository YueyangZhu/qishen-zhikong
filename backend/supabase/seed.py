"""种子数据迁移脚本：把演示数据写入 Supabase 数据库

用法：
    cd backend
    .\\venv\\Scripts\\python.exe supabase\\seed.py

数据来源：src/mock/seedData.ts + src/mock/contractText.ts
表结构：  backend/supabase/schema.sql
字段映射：camelCase → snake_case（与 app/routers/data.py 的 _to_db_row 一致）

幂等性：每次运行先清空表再插入。
"""
import os
import re
import sys
import json
from pathlib import Path
from typing import Optional

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

from dotenv import load_dotenv
from supabase import create_client

# ===== 加载 .env（位于 backend/.env）=====
BACKEND_DIR = Path(__file__).resolve().parent.parent
load_dotenv(BACKEND_DIR / ".env")

SUPABASE_URL = os.getenv("SUPABASE_URL", "").strip()
SUPABASE_SERVICE_ROLE_KEY = os.getenv("SUPABASE_SERVICE_ROLE_KEY", "").strip()


# ============================================================
# 数据定义（与 src/mock/seedData.ts、contractText.ts 保持一致）
# 字段保持 camelCase，插入前由 to_snake_row 转为 snake_case
# ============================================================

# ===== 用户（3 个）=====
DEMO_USERS = [
    {
        "id": "U-PURCHASER",
        "name": "李明",
        "email": "purchaser@qszk.com",
        "role": "purchaser",
        "department": "采购部",
        "position": "采购经理",
        "avatarColor": "#1677ff",
    },
    {
        "id": "U-LEGAL",
        "name": "王律师",
        "email": "legal@qszk.com",
        "role": "legal",
        "department": "法务部",
        "position": "高级法务",
        "avatarColor": "#13c2c2",
    },
    {
        "id": "U-ADMIN",
        "name": "张管理员",
        "email": "admin@qszk.com",
        "role": "admin",
        "department": "信息技术部",
        "position": "系统管理员",
        "avatarColor": "#722ed1",
    },
]

# ===== 风险规则库（16 条）=====
DEMO_RULES = [
    {"id": "RR-001", "code": "RR-SUB-001", "name": "合同主体名称或统一社会信用代码缺失", "contractType": "采购合同", "riskType": "subject", "riskLevel": "medium", "method": "field", "triggerCondition": "甲乙方名称、统一社会信用代码、联系地址任一缺失", "reasonTemplate": "合同主体信息不完整，缺少{missing}，可能影响主体资格认定与履约追责。", "suggestionTemplate": "建议补充完整的主体信息，并与营业执照一致。", "status": "enabled", "version": 1, "description": "校验合同主体字段完整性", "updatedAt": "2026-06-01T09:00:00.000Z"},
    {"id": "RR-002", "code": "RR-AMT-001", "name": "合同大小写金额不一致", "contractType": "采购合同", "riskType": "amount", "riskLevel": "high", "method": "field", "triggerCondition": "数字金额与大写金额折算后不一致", "reasonTemplate": "合同金额数字与大写不一致，存在金额认定争议风险。", "suggestionTemplate": "以确认金额统一修改大小写表述。", "status": "enabled", "version": 1, "description": "金额大小写一致性校验", "updatedAt": "2026-06-01T09:00:00.000Z"},
    {"id": "RR-003", "code": "RR-PAY-001", "name": "预付款比例超过50%且无担保", "contractType": "采购合同", "riskType": "payment", "riskLevel": "high", "method": "field", "triggerCondition": "预付款比例>50% 且未约定履约保函/担保", "reasonTemplate": "预付款比例{ratio}过高，未约定履约保障，存在款项回收风险。", "suggestionTemplate": "降低预付款比例至30%以内，或要求乙方提供等额履约保函。", "status": "enabled", "version": 1, "description": "预付款比例与担保校验", "updatedAt": "2026-06-02T09:00:00.000Z"},
    {"id": "RR-004", "code": "RR-PAY-002", "name": "预付款缺少履约保障", "contractType": "采购合同", "riskType": "payment", "riskLevel": "high", "method": "ai", "triggerCondition": "存在预付款条款但未约定履约保函、保证金或担保措施", "reasonTemplate": "预付款无履约保障，乙方违约时甲方预付款难以追回。", "suggestionTemplate": "要求乙方提供与预付款等额的履约保函或保证金。", "status": "enabled", "version": 1, "description": "预付款履约保障校验", "updatedAt": "2026-06-02T09:00:00.000Z"},
    {"id": "RR-005", "code": "RR-PAY-003", "name": "未约定发票类型或开票时间", "contractType": "采购合同", "riskType": "payment", "riskLevel": "medium", "method": "field", "triggerCondition": "付款条款未约定发票类型或开票节点", "reasonTemplate": "未约定增值税发票类型及开票时间，影响进项抵扣与付款节奏。", "suggestionTemplate": "补充增值税专用发票类型，并明确开票节点（如付款前开具）。", "status": "enabled", "version": 1, "description": "发票条款校验", "updatedAt": "2026-06-02T09:00:00.000Z"},
    {"id": "RR-006", "code": "RR-DEL-001", "name": "未明确交付日期或仅写尽快", "contractType": "采购合同", "riskType": "delivery", "riskLevel": "high", "method": "ai", "triggerCondition": "交付日期缺失或使用\"尽快\"\"及时\"等模糊表述", "reasonTemplate": "交付日期不明确，使用\"{word}\"等模糊表述，延期责任难以主张。", "suggestionTemplate": "约定明确的交付日期（如2026年X月X日）及延期违约责任。", "status": "enabled", "version": 1, "description": "交付日期明确性校验", "updatedAt": "2026-06-03T09:00:00.000Z"},
    {"id": "RR-007", "code": "RR-ACC-001", "name": "验收标准无法量化", "contractType": "采购合同", "riskType": "acceptance", "riskLevel": "medium", "method": "ai", "triggerCondition": "验收标准使用\"符合甲方要求\"等无量化指标表述", "reasonTemplate": "验收标准模糊，缺乏量化指标，易引发验收争议。", "suggestionTemplate": "补充量化验收标准、测试用例及验收流程，作为合同附件。", "status": "enabled", "version": 1, "description": "验收标准量化校验", "updatedAt": "2026-06-03T09:00:00.000Z"},
    {"id": "RR-008", "code": "RR-ACC-002", "name": "验收期限缺失", "contractType": "采购合同", "riskType": "acceptance", "riskLevel": "medium", "method": "field", "triggerCondition": "未约定验收期限或异议期", "reasonTemplate": "未约定验收期限，可能导致视为验收合格的法律后果。", "suggestionTemplate": "约定交付后X个工作日内完成验收，逾期未提出视为合格。", "status": "enabled", "version": 1, "description": "验收期限校验", "updatedAt": "2026-06-03T09:00:00.000Z"},
    {"id": "RR-009", "code": "RR-WAR-001", "name": "未约定质保期和响应时限", "contractType": "采购合同", "riskType": "warranty", "riskLevel": "medium", "method": "field", "triggerCondition": "未约定质保期限或故障响应/修复时限", "reasonTemplate": "质保期限及响应时限缺失，售后责任不清。", "suggestionTemplate": "补充质保期限（如1年）、故障响应时限（如4小时）及修复时限。", "status": "enabled", "version": 1, "description": "质保条款校验", "updatedAt": "2026-06-04T09:00:00.000Z"},
    {"id": "RR-010", "code": "RR-BRCH-001", "name": "违约责任严重不对等", "contractType": "采购合同", "riskType": "breach", "riskLevel": "high", "method": "ai", "triggerCondition": "甲方违约金显著高于乙方违约金或乙方责任上限过低", "reasonTemplate": "违约责任不对等，甲方违约金{a}，乙方违约金{b}且上限过低，权利义务失衡。", "suggestionTemplate": "调整为对等违约责任，取消或提高乙方违约金上限。", "status": "enabled", "version": 1, "description": "违约责任对等性校验", "updatedAt": "2026-06-04T09:00:00.000Z"},
    {"id": "RR-011", "code": "RR-TERM-001", "name": "乙方享有单方解除权", "contractType": "采购合同", "riskType": "termination", "riskLevel": "high", "method": "ai", "triggerCondition": "乙方享有单方解除权而甲方无对应权利", "reasonTemplate": "乙方享有单方解除权，甲方无对应权利，存在被动终止风险。", "suggestionTemplate": "删除乙方单方解除权，或设置甲方对等的解除条件。", "status": "enabled", "version": 1, "description": "合同解除权对等性校验", "updatedAt": "2026-06-05T09:00:00.000Z"},
    {"id": "RR-012", "code": "RR-IP-001", "name": "定制成果知识产权归属不利", "contractType": "采购合同", "riskType": "ip", "riskLevel": "high", "method": "ai", "triggerCondition": "定制开发成果知识产权全部归乙方，甲方仅享有使用权", "reasonTemplate": "定制成果知识产权全部归乙方，甲方付费开发却无所有权，权益受损。", "suggestionTemplate": "明确甲方对定制成果的所有权，或至少享有永久、不可撤销的使用权。", "status": "enabled", "version": 1, "description": "知识产权归属校验", "updatedAt": "2026-06-05T09:00:00.000Z"},
    {"id": "RR-013", "code": "RR-CON-001", "name": "保密义务无期限", "contractType": "采购合同", "riskType": "confidentiality", "riskLevel": "medium", "method": "ai", "triggerCondition": "保密条款未约定保密期限或范围", "reasonTemplate": "保密义务无期限约定，可能被无限期扩大解释。", "suggestionTemplate": "补充保密信息定义、保密期限（如3年）及例外情形。", "status": "enabled", "version": 1, "description": "保密期限校验", "updatedAt": "2026-06-06T09:00:00.000Z"},
    {"id": "RR-014", "code": "RR-DS-001", "name": "数据安全责任划分不清", "contractType": "采购合同", "riskType": "data_security", "riskLevel": "high", "method": "ai", "triggerCondition": "未约定数据泄露责任、通知义务及终止后数据删除返还", "reasonTemplate": "数据安全责任划分不清，未约定泄露通知与数据删除机制。", "suggestionTemplate": "补充数据安全责任、泄露通知时限及合同终止后数据删除/返还机制。", "status": "enabled", "version": 1, "description": "数据安全责任校验", "updatedAt": "2026-06-06T09:00:00.000Z"},
    {"id": "RR-015", "code": "RR-DIS-001", "name": "争议管辖地对我方不利", "contractType": "采购合同", "riskType": "dispute", "riskLevel": "high", "method": "ai", "triggerCondition": "约定由乙方所在地法院管辖", "reasonTemplate": "争议由乙方所在地法院管辖，异地诉讼增加维权成本。", "suggestionTemplate": "改为甲方所在地或合同履行地有管辖权的法院管辖。", "status": "enabled", "version": 1, "description": "争议管辖校验", "updatedAt": "2026-06-07T09:00:00.000Z"},
    {"id": "RR-016", "code": "RR-TM-001", "name": "自动续期未设置提前通知", "contractType": "采购合同", "riskType": "term", "riskLevel": "medium", "method": "ai", "triggerCondition": "约定自动续期但未设置提前通知或确认", "reasonTemplate": "自动续期未约定提前通知，可能导致非预期续约。", "suggestionTemplate": "增加续期前30日书面确认或提前通知条款。", "status": "enabled", "version": 1, "description": "合同期限与续期校验", "updatedAt": "2026-06-07T09:00:00.000Z"},
    # 新增规则（关键词匹配增强版，供测试合同验证）
    {"id": "RR-017", "code": "RR-SUB-002", "name": "甲方主体信息不完整", "contractType": "采购合同", "riskType": "subject", "riskLevel": "medium", "method": "field", "triggerCondition": "甲方名称或统一社会信用代码缺失", "reasonTemplate": "甲方主体信息不完整，缺少{missing}，可能影响合同效力。", "suggestionTemplate": "补充甲方完整主体信息，包括名称、信用代码及联系地址。", "status": "enabled", "version": 1, "description": "甲方主体字段完整性校验", "updatedAt": "2026-06-08T09:00:00.000Z"},
    {"id": "RR-018", "code": "RR-DOC-001", "name": "合同附件清单缺失", "contractType": "采购合同", "riskType": "subject", "riskLevel": "low", "method": "keyword", "triggerCondition": "合同未列明附件清单或附件内容", "reasonTemplate": "未列明合同附件清单，附件法律效力不清。", "suggestionTemplate": "补充合同附件清单，明确附件内容与法律效力。", "status": "enabled", "version": 1, "description": "附件完整性校验", "updatedAt": "2026-06-08T09:00:00.000Z"},
    {"id": "RR-019", "code": "RR-PAY-004", "name": "付款节点与交付节点脱钩", "contractType": "采购合同", "riskType": "payment", "riskLevel": "high", "method": "keyword", "triggerCondition": "付款节点与交付、验收节点不对应", "reasonTemplate": "付款节点与交付/验收节点未关联，可能导致先付款后收货的风险。", "suggestionTemplate": "将付款节点与交付验收节点挂钩，如验收合格后X日内付款。", "status": "enabled", "version": 1, "description": "付款节奏校验", "updatedAt": "2026-06-09T09:00:00.000Z"},
    {"id": "RR-020", "code": "RR-REN-001", "name": "续约条件不明确", "contractType": "采购合同", "riskType": "term", "riskLevel": "medium", "method": "keyword", "triggerCondition": "续约条件或续约价格调整机制未约定", "reasonTemplate": "续约条件未明确，续约时价格与服务可能产生争议。", "suggestionTemplate": "明确续约条件、价格调整机制及续约通知期限。", "status": "enabled", "version": 1, "description": "续约条款校验", "updatedAt": "2026-06-09T09:00:00.000Z"},
    {"id": "RR-021", "code": "RR-FORCE-001", "name": "不可抗力条款缺失", "contractType": "采购合同", "riskType": "term", "riskLevel": "medium", "method": "keyword", "triggerCondition": "未约定不可抗力条款", "reasonTemplate": "未约定不可抗力条款，遇不可抗力事件时责任划分不清。", "suggestionTemplate": "补充不可抗力条款，明确不可抗力事件范围、通知义务与免责范围。", "status": "enabled", "version": 1, "description": "不可抗力条款校验", "updatedAt": "2026-06-10T09:00:00.000Z"},
    {"id": "RR-022", "code": "RR-INDEM-001", "name": "赔偿上限不合理", "contractType": "采购合同", "riskType": "breach", "riskLevel": "medium", "method": "keyword", "triggerCondition": "乙方赔偿上限过低", "reasonTemplate": "乙方赔偿上限过低，不足以覆盖甲方可能遭受的损失。", "suggestionTemplate": "提高赔偿上限至合同金额的100%或实际损失额。", "status": "enabled", "version": 1, "description": "赔偿上限校验", "updatedAt": "2026-06-10T09:00:00.000Z"},
    {"id": "RR-023", "code": "RR-ASSIGN-001", "name": "合同转让未约定限制", "contractType": "采购合同", "riskType": "term", "riskLevel": "low", "method": "keyword", "triggerCondition": "未约定合同转让限制条件", "reasonTemplate": "未约定合同转让限制，任一方可将合同权利义务转让给第三方。", "suggestionTemplate": "增加合同转让限制条款，约定未经对方书面同意不得转让。", "status": "enabled", "version": 1, "description": "合同转让限制校验", "updatedAt": "2026-06-11T09:00:00.000Z"},
    {"id": "RR-024", "code": "RR-LAW-001", "name": "适用法律未约定", "contractType": "采购合同", "riskType": "dispute", "riskLevel": "medium", "method": "keyword", "triggerCondition": "未约定合同适用法律", "reasonTemplate": "未约定适用法律，争议时法律适用不明确。", "suggestionTemplate": "明确约定适用中华人民共和国法律。", "status": "enabled", "version": 1, "description": "适用法律校验", "updatedAt": "2026-06-11T09:00:00.000Z"},
    {"id": "RR-025", "code": "RR-COST-001", "name": "隐含费用未约定承担方", "contractType": "采购合同", "riskType": "amount", "riskLevel": "medium", "method": "keyword", "triggerCondition": "运输、保险、安装等附带费用承担未约定", "reasonTemplate": "运输、保险、安装等附带费用承担方未明确，可能产生额外成本。", "suggestionTemplate": "明确运输费、保险费、安装费等附带费用的承担方。", "status": "enabled", "version": 1, "description": "附带费用校验", "updatedAt": "2026-06-12T09:00:00.000Z"},
    {"id": "RR-026", "code": "RR-TRAIN-001", "name": "培训义务未约定", "contractType": "采购合同", "riskType": "delivery", "riskLevel": "low", "method": "keyword", "triggerCondition": "涉及系统交付但未约定培训义务", "reasonTemplate": "系统交付未约定培训义务，甲方人员可能无法正常使用系统。", "suggestionTemplate": "补充乙方培训义务，包括培训时间、内容与人数。", "status": "enabled", "version": 1, "description": "培训义务校验", "updatedAt": "2026-06-12T09:00:00.000Z"},
]

# ===== 风险模板（18 个，引用 RISK_SNIPPETS）=====
DEMO_RISK_TEMPLATES = [
    {"snippetKey": "R1", "title": "乙方主体信息不完整", "riskType": "subject", "riskLevel": "medium", "clauseNumber": "第一条", "clauseTitle": "合同主体", "riskReason": "乙方未提供统一社会信用代码与联系地址，主体信息不完整，可能影响主体资格认定与后续履约追责。", "reviewBasis": "《合同主体信息完整性规则 RR-SUB-001》：合同主体应包含名称、统一社会信用代码、法定代表人及联系地址。", "suggestion": "要求乙方补充完整的统一社会信用代码、注册地址及联系方式，并与营业执照核对一致。", "confidence": 0.88, "sourceType": "rule", "ruleId": "RR-001"},
    {"snippetKey": "R2", "title": "合同金额大小写不一致", "riskType": "amount", "riskLevel": "high", "clauseNumber": "第三条", "clauseTitle": "合同金额", "riskReason": "合同金额数字为580000元（伍拾捌万元），但大写写为\"伍拾捌万捌仟元整\"，多出捌仟元，大小写不一致，存在金额认定争议风险。", "reviewBasis": "《金额一致性规则 RR-AMT-001》：合同金额数字与大写折算后应完全一致。", "suggestion": "以580000元为准，将大写修改为\"伍拾捌万元整\"，确保大小写一致。", "confidence": 0.95, "sourceType": "rule", "ruleId": "RR-002"},
    {"snippetKey": "R17", "title": "发票类型与开票时间未约定", "riskType": "payment", "riskLevel": "medium", "clauseNumber": "第三条", "clauseTitle": "合同金额", "riskReason": "金额条款仅注明\"已含增值税\"，未约定发票类型（专票/普票）及开票时间，影响进项抵扣。", "reviewBasis": "《发票规则 RR-PAY-003》：应约定发票类型与开票节点。", "suggestion": "明确开具增值税专用发票（13%），并约定付款前开具。", "confidence": 0.65, "sourceType": "rule", "ruleId": "RR-005"},
    {"snippetKey": "R3", "title": "预付款比例过高", "riskType": "payment", "riskLevel": "high", "clauseNumber": "第四条", "clauseTitle": "付款方式", "riskReason": "预付款比例为80%，显著超过50%的行业惯例，若乙方履约能力不足，甲方预付款回收风险较高。", "reviewBasis": "《预付款规则 RR-PAY-001》：预付款比例超过50%且无担保时触发高风险。", "suggestion": "将预付款比例降至30%以内，或要求乙方提供等额履约保函。", "confidence": 0.93, "sourceType": "rule", "ruleId": "RR-003"},
    {"snippetKey": "R4", "title": "预付款缺少履约保障", "riskType": "payment", "riskLevel": "high", "clauseNumber": "第四条", "clauseTitle": "付款方式", "riskReason": "约定80%预付款但未要求乙方提供履约保函或保证金，一旦乙方违约，甲方预付款难以追回。", "reviewBasis": "《预付款履约保障规则 RR-PAY-002》：高额预付款应配套履约保函或担保。", "suggestion": "要求乙方提供与预付款等额的银行履约保函，或在付款前约定担保措施。", "confidence": 0.86, "sourceType": "ai", "ruleId": "RR-004"},
    {"snippetKey": "R5", "title": "交付日期不明确", "riskType": "delivery", "riskLevel": "high", "clauseNumber": "第五条", "clauseTitle": "交付安排", "riskReason": "交付时间仅表述为\"尽快\"，无明确日期，乙方延期交付时甲方难以主张违约责任。", "reviewBasis": "《交付日期规则 RR-DEL-001》：交付日期应明确，禁止使用\"尽快\"等模糊表述。", "suggestion": "约定明确的交付日期（如2026年8月31日），并约定延期交付的违约责任。", "confidence": 0.5, "sourceType": "ai", "ruleId": "RR-006"},
    {"snippetKey": "R6", "title": "验收标准无法量化", "riskType": "acceptance", "riskLevel": "medium", "clauseNumber": "第六条", "clauseTitle": "验收标准", "riskReason": "验收标准为\"符合甲方要求\"，缺乏量化指标与测试依据，易引发验收争议。", "reviewBasis": "《验收标准规则 RR-ACC-001》：验收标准应可量化、可测试。", "suggestion": "补充量化验收指标、测试用例及验收流程，作为合同附件。", "confidence": 0.82, "sourceType": "ai", "ruleId": "RR-007"},
    {"snippetKey": "R7", "title": "验收期限缺失", "riskType": "acceptance", "riskLevel": "medium", "clauseNumber": "第六条", "clauseTitle": "验收标准", "riskReason": "未约定验收期限与异议期，依据《民法典》可能产生\"交付后视为验收合格\"的法律后果。", "reviewBasis": "《验收期限规则 RR-ACC-002》：应约定验收期限与异议期。", "suggestion": "约定交付后10个工作日内完成验收，逾期未提出异议视为验收合格。", "confidence": 0.62, "sourceType": "rule", "ruleId": "RR-008"},
    {"snippetKey": "R8", "title": "知识产权归属对甲方不利", "riskType": "ip", "riskLevel": "high", "clauseNumber": "第七条", "clauseTitle": "知识产权", "riskReason": "甲方付费定制开发的成果知识产权全部归乙方，甲方仅享有非独占使用权，权益严重受损。", "reviewBasis": "《知识产权规则 RR-IP-001》：定制开发成果归属应保障甲方权益。", "suggestion": "明确甲方对定制成果的所有权，或至少享有永久、不可撤销、可转授权的使用权。", "confidence": 0.91, "sourceType": "ai", "ruleId": "RR-012"},
    {"snippetKey": "R9", "title": "质保期限缺失", "riskType": "warranty", "riskLevel": "medium", "clauseNumber": "第八条", "clauseTitle": "质保服务", "riskReason": "质保期限与响应时限\"另行约定\"，实质未明确，售后责任不清。", "reviewBasis": "《质保规则 RR-WAR-001》：应明确质保期限与故障响应时限。", "suggestion": "约定不少于1年的质保期，故障4小时响应、24小时修复或提供备机。", "confidence": 0.55, "sourceType": "rule", "ruleId": "RR-009"},
    {"snippetKey": "R10", "title": "保密期限缺失", "riskType": "confidentiality", "riskLevel": "medium", "clauseNumber": "第九条", "clauseTitle": "保密条款", "riskReason": "保密义务未约定期限与范围，可能被无限期扩大解释。", "reviewBasis": "《保密规则 RR-CON-001》：保密条款应明确期限、范围与例外。", "suggestion": "补充保密信息定义、3年保密期限及法定例外情形。", "confidence": 0.79, "sourceType": "ai", "ruleId": "RR-013"},
    {"snippetKey": "R11", "title": "数据安全责任划分不清", "riskType": "data_security", "riskLevel": "high", "clauseNumber": "第十条", "clauseTitle": "数据安全", "riskReason": "数据泄露损失\"由双方共同承担\"，责任比例不清，且未约定泄露通知与终止后数据删除。", "reviewBasis": "《数据安全规则 RR-DS-001》：应明确数据泄露责任、通知义务与数据删除机制。", "suggestion": "明确乙方为主要责任方，约定泄露24小时内通知、终止后数据删除/返还。", "confidence": 0.84, "sourceType": "ai", "ruleId": "RR-014"},
    {"snippetKey": "R12", "title": "自动续期未约定提前通知", "riskType": "term", "riskLevel": "notice", "clauseNumber": "第十一条", "clauseTitle": "合同期限", "riskReason": "合同期满自动续期，未约定提前通知或确认，可能导致非预期续约。", "reviewBasis": "《期限规则 RR-TM-001》：自动续期应设置提前通知。", "suggestion": "增加续期前30日书面确认条款，任一方不同意则不续期。", "confidence": 0.9, "sourceType": "ai", "ruleId": "RR-016"},
    {"snippetKey": "R13", "title": "甲方逾期付款违约金过高", "riskType": "breach", "riskLevel": "high", "clauseNumber": "第十二条", "clauseTitle": "违约责任", "riskReason": "甲方逾期付款违约金为日千分之五（年化约182%），显著过高，可能被法院调整且显失公平。", "reviewBasis": "《违约责任规则 RR-BRCH-001》：违约金应合理且对等。", "suggestion": "将甲方违约金调整为日万分之五（与乙方对等），或日千分之一以内。", "confidence": 0.92, "sourceType": "ai", "ruleId": "RR-010"},
    {"snippetKey": "R14", "title": "乙方延期交付责任过低", "riskType": "breach", "riskLevel": "high", "clauseNumber": "第十二条", "clauseTitle": "违约责任", "riskReason": "乙方延期违约金仅日千分之一且上限1%，与甲方日千分之五严重不对等，乙方违约成本低。", "reviewBasis": "《违约责任对等规则 RR-BRCH-001》：双方违约责任应基本对等。", "suggestion": "提高乙方违约金至与甲方对等，取消1%上限或提高至合同总额的10%。", "confidence": 0.9, "sourceType": "ai", "ruleId": "RR-010"},
    {"snippetKey": "R15", "title": "乙方享有单方解除权", "riskType": "termination", "riskLevel": "high", "clauseNumber": "第十三条", "clauseTitle": "合同解除", "riskReason": "乙方在甲方逾期付款超15日时可单方解除合同，甲方无对应解除权，权利义务失衡。", "reviewBasis": "《解除权规则 RR-TERM-001》：解除权应双方对等。", "suggestion": "增加甲方对等解除权（如乙方延期交付超15日），或设置协商解除机制。", "confidence": 0.87, "sourceType": "ai", "ruleId": "RR-011"},
    {"snippetKey": "R16", "title": "争议管辖地对我方不利", "riskType": "dispute", "riskLevel": "high", "clauseNumber": "第十四条", "clauseTitle": "争议解决", "riskReason": "争议由乙方所在地法院管辖，异地诉讼显著增加甲方维权成本与沟通成本。", "reviewBasis": "《争议管辖规则 RR-DIS-001》：管辖地应有利于我方。", "suggestion": "改为由甲方所在地或合同履行地有管辖权的法院管辖。", "confidence": 0.94, "sourceType": "ai", "ruleId": "RR-015"},
    {"snippetKey": "R18", "title": "通知方式未明确约定", "riskType": "subject", "riskLevel": "low", "clauseNumber": "第十五条", "clauseTitle": "附则", "riskReason": "未约定双方通知送达方式与有效送达地址，争议时通知送达可能产生争议。", "reviewBasis": "常规合同条款完整性提示：建议约定通知送达条款。", "suggestion": "补充通知送达方式（书面/邮件）、有效送达地址及视为送达的情形。", "confidence": 0.7, "sourceType": "ai", "ruleId": ""},
]

# ===== 风险原文片段（key -> {paragraphId, originalText}）=====
RISK_SNIPPETS = {
    "R1": {"paragraphId": "p2", "originalText": "乙方（供应方）：星河软件有限公司\n统一社会信用代码：未提供\n法定代表人：刘星河\n联系地址：未提供"},
    "R2": {"paragraphId": "p4", "originalText": "580000元（大写：伍拾捌万捌仟元整）"},
    "R3": {"paragraphId": "p5", "originalText": "支付合同总额的80%作为预付款"},
    "R4": {"paragraphId": "p5", "originalText": "1.合同签订后7个工作日内，甲方应向乙方支付合同总额的80%作为预付款；"},
    "R5": {"paragraphId": "p6", "originalText": "尽快完成系统的交付与安装部署"},
    "R6": {"paragraphId": "p7", "originalText": "应符合甲方要求"},
    "R7": {"paragraphId": "p7", "originalText": "经甲方确认后签署验收报告，即视为验收合格"},
    "R8": {"paragraphId": "p8", "originalText": "全部归乙方所有，甲方仅享有非独占的使用权"},
    "R9": {"paragraphId": "p9", "originalText": "具体质保期限及响应时限由双方另行约定"},
    "R10": {"paragraphId": "p10", "originalText": "承担保密义务，未经对方书面同意不得向第三方披露"},
    "R11": {"paragraphId": "p11", "originalText": "因数据泄露造成的损失由双方共同承担"},
    "R12": {"paragraphId": "p12", "originalText": "期满后若需继续合作，自动续期"},
    "R13": {"paragraphId": "p13", "originalText": "每日按应付未付金额的千分之五支付违约金"},
    "R14": {"paragraphId": "p13", "originalText": "每日按合同总额的千分之一支付违约金，累计不超过合同总额的1%"},
    "R15": {"paragraphId": "p14", "originalText": "甲方逾期付款超过15日的，乙方有权单方解除本合同"},
    "R16": {"paragraphId": "p15", "originalText": "由乙方所在地有管辖权的人民法院管辖"},
    "R17": {"paragraphId": "p4", "originalText": "上述金额已含增值税"},
    "R18": {"paragraphId": "p16", "originalText": "未尽事宜由双方另行约定"},
}

# ===== 合同段落（用于计算风险原文位置）=====
# type 说明：title=标题 header=首部甲乙方信息 body=正文条款 signature=签署落款
DEMO_PARAGRAPHS = [
    {"id": "p1", "index": 1, "type": "title", "text": "软件系统采购合同"},
    {"id": "p2", "index": 2, "type": "body", "clauseNo": "第一条", "clauseTitle": "合同主体", "text": "第一条 合同主体\n甲方（采购方）：智远科技有限公司\n统一社会信用代码：91110108MA01ABC23X\n法定代表人：陈志远\n联系地址：北京市海淀区中关村大街1号\n乙方（供应方）：星河软件有限公司\n统一社会信用代码：未提供\n法定代表人：刘星河\n联系地址：未提供"},
    {"id": "p3", "index": 3, "type": "body", "clauseNo": "第二条", "clauseTitle": "采购标的", "text": "第二条 采购标的\n乙方为甲方提供“智远协同办公平台”软件系统一套，含系统授权许可、安装部署、基础培训及一年期技术支持服务。系统功能与技术规格以双方确认的附件为准。"},
    {"id": "p4", "index": 4, "type": "body", "clauseNo": "第三条", "clauseTitle": "合同金额", "text": "第三条 合同金额\n合同总金额为人民币580000元（大写：伍拾捌万捌仟元整）。上述金额已含增值税，不以其他费用另行收取。"},
    {"id": "p5", "index": 5, "type": "body", "clauseNo": "第四条", "clauseTitle": "付款方式", "text": "第四条 付款方式\n1.合同签订后7个工作日内，甲方应向乙方支付合同总额的80%作为预付款；\n2.系统验收合格后10个工作日内，甲方支付剩余20%尾款。"},
    {"id": "p6", "index": 6, "type": "body", "clauseNo": "第五条", "clauseTitle": "交付安排", "text": "第五条 交付安排\n乙方应在合同签订后尽快完成系统的交付与安装部署，并配合甲方完成上线准备。"},
    {"id": "p7", "index": 7, "type": "body", "clauseNo": "第六条", "clauseTitle": "验收标准", "text": "第六条 验收标准\n系统交付后应符合甲方要求，经甲方确认后签署验收报告，即视为验收合格。"},
    {"id": "p8", "index": 8, "type": "body", "clauseNo": "第七条", "clauseTitle": "知识产权", "text": "第七条 知识产权\n乙方为甲方定制开发的系统成果及相关知识产权，全部归乙方所有，甲方仅享有非独占的使用权。"},
    {"id": "p9", "index": 9, "type": "body", "clauseNo": "第八条", "clauseTitle": "质保服务", "text": "第八条 质保服务\n乙方对所提供的系统提供质保服务，具体质保期限及响应时限由双方另行约定。"},
    {"id": "p10", "index": 10, "type": "body", "clauseNo": "第九条", "clauseTitle": "保密条款", "text": "第九条 保密条款\n双方应对因履行本合同而知悉的对方商业信息承担保密义务，未经对方书面同意不得向第三方披露。"},
    {"id": "p11", "index": 11, "type": "body", "clauseNo": "第十条", "clauseTitle": "数据安全", "text": "第十条 数据安全\n双方应采取必要措施保障数据安全，因数据泄露造成的损失由双方共同承担。"},
    {"id": "p12", "index": 12, "type": "body", "clauseNo": "第十一条", "clauseTitle": "合同期限", "text": "第十一条 合同期限\n本合同自双方签字盖章之日起生效，有效期为2年。期满后若需继续合作，自动续期。"},
    {"id": "p13", "index": 13, "type": "body", "clauseNo": "第十二条", "clauseTitle": "违约责任", "text": "第十二条 违约责任\n1.甲方逾期付款的，每日按应付未付金额的千分之五支付违约金；\n2.乙方延期交付的，每日按合同总额的千分之一支付违约金，累计不超过合同总额的1%。"},
    {"id": "p14", "index": 14, "type": "body", "clauseNo": "第十三条", "clauseTitle": "合同解除", "text": "第十三条 合同解除\n甲方逾期付款超过15日的，乙方有权单方解除本合同，并要求甲方承担相应违约责任。"},
    {"id": "p15", "index": 15, "type": "body", "clauseNo": "第十四条", "clauseTitle": "争议解决", "text": "第十四条 争议解决\n因本合同产生的或与本合同有关的争议，由乙方所在地有管辖权的人民法院管辖。"},
    {"id": "p16", "index": 16, "type": "signature", "clauseNo": "第十五条", "clauseTitle": "附则", "text": "第十五条 附则\n本合同一式两份，双方各执一份，自双方签字盖章之日起生效。未尽事宜由双方另行约定。"},
]

DEMO_PARAGRAPHS_BY_ID = {p["id"]: p for p in DEMO_PARAGRAPHS}

# ===== 样例合同数据（来自 src/mock/sampleContracts.ts）=====
# 按 sampleId 匹配对应合同正文与风险模板；无 sampleId 的任务回退到 DEMO_PARAGRAPHS / DEMO_RISK_TEMPLATES
SAMPLE_CONTRACTS = {
    "sample-1": {
        "contractName": "办公设备采购合同",
        "paragraphs": [
            {"id": "p1", "index": 1, "type": "title", "text": "办公设备采购合同"},
            {"id": "p2", "index": 2, "type": "header", "text": "甲方（采购方）：智远科技有限公司"},
            {"id": "p3", "index": 3, "type": "header", "text": "乙方（供应方）：深圳市宏图办公设备有限公司，法定代表人：王某，注册地址：深圳市福田区某街道。"},
            {"id": "p4", "index": 4, "type": "body", "clauseNo": "第一条", "clauseTitle": "采购标的", "text": "第一条 采购标的：笔记本电脑60台，台式电脑15台，激光打印机5台，规格详见附件清单。"},
            {"id": "p5", "index": 5, "type": "body", "clauseNo": "第二条", "clauseTitle": "合同金额", "text": "第二条 合同金额：本合同总金额为人民币380000元（大写：人民币叁拾捌万元整），含增值税。"},
            {"id": "p6", "index": 6, "type": "body", "clauseNo": "第三条", "clauseTitle": "付款方式", "text": "第三条 付款方式：甲方在货到验收合格后7个工作日内一次性付清全款。"},
            {"id": "p7", "index": 7, "type": "body", "clauseNo": "第四条", "clauseTitle": "交付时间", "text": "第四条 交付时间：乙方应在合同签订后30日内完成交付并安装到位。"},
            {"id": "p8", "index": 8, "type": "body", "clauseNo": "第五条", "clauseTitle": "验收标准", "text": "第五条 验收标准：设备外观完好、功能正常，符合甲方使用要求。"},
            {"id": "p9", "index": 9, "type": "body", "clauseNo": "第六条", "clauseTitle": "质保期限", "text": "第六条 质保期限：乙方提供1年免费质保服务。"},
            {"id": "p10", "index": 10, "type": "body", "clauseNo": "第七条", "clauseTitle": "知识产权", "text": "第七条 知识产权：设备所附软件的知识产权归乙方所有，甲方仅享有使用权。"},
            {"id": "p11", "index": 11, "type": "body", "clauseNo": "第八条", "clauseTitle": "保密条款", "text": "第八条 保密条款：双方对合作中知悉的对方商业信息予以保密。"},
            {"id": "p12", "index": 12, "type": "body", "clauseNo": "第九条", "clauseTitle": "违约责任", "text": "第九条 违约责任：乙方逾期交付每日按合同总额的0.01%支付违约金。"},
            {"id": "p13", "index": 13, "type": "body", "clauseNo": "第十条", "clauseTitle": "合同解除", "text": "第十条 合同解除：任何一方需提前解除合同，应提前30日书面通知对方。"},
            {"id": "p14", "index": 14, "type": "body", "clauseNo": "第十一条", "clauseTitle": "争议解决", "text": "第十一条 争议解决：本合同争议由乙方所在地法院管辖。"},
            {"id": "p15", "index": 15, "type": "signature", "text": "甲方（盖章）：智远科技有限公司  乙方（盖章）：深圳市宏图办公设备有限公司  签订日期：2024年6月15日"},
        ],
        "riskTemplates": [
            {"title": "乙方逾期违约金比例过低", "riskType": "breach", "riskLevel": "high", "clauseNumber": "第九条", "clauseTitle": "违约责任", "originalText": "乙方逾期交付每日按合同总额的0.01%支付违约金", "paragraphId": "p12", "riskReason": "0.01%/日折年化仅 3.65%，远低于行业惯例 0.3%-0.5%/日，无法有效约束乙方按时交付。", "reviewBasis": "行业惯例：设备采购合同逾期违约金通常为 0.3%-0.5%/日。", "suggestion": "建议调整为每日 0.3%-0.5%，并设置违约金上限为合同总额的 10%。", "confidence": 0.92, "sourceType": "rule", "ruleId": "RR-010"},
            {"title": "争议管辖对甲方不利", "riskType": "dispute", "riskLevel": "high", "clauseNumber": "第十一条", "clauseTitle": "争议解决", "originalText": "本合同争议由乙方所在地法院管辖", "paragraphId": "p14", "riskReason": "管辖法院在乙方所在地（深圳），甲方发生争议时诉讼成本较高。", "reviewBasis": "《民事诉讼法》第 24 条规定合同纠纷可由被告所在地或合同履行地法院管辖。", "suggestion": "建议改为\"由甲方所在地或合同履行地法院管辖\"。", "confidence": 0.94, "sourceType": "rule", "ruleId": "RR-015"},
            {"title": "验收标准过于笼统", "riskType": "acceptance", "riskLevel": "medium", "clauseNumber": "第五条", "clauseTitle": "验收标准", "originalText": "设备外观完好、功能正常，符合甲方使用要求", "paragraphId": "p8", "riskReason": "验收标准缺乏量化指标，\"功能正常\"无明确判定依据，容易产生验收争议。", "reviewBasis": "《企业采购管理规范》第 5.2 条要求验收标准应可量化、可验证。", "suggestion": "建议明确验收指标，如设备开机自检通过率 100%、性能跑分达标、7 日内无故障率等，并约定复检流程。", "confidence": 0.85, "sourceType": "ai", "ruleId": None},
            {"title": "质保期限偏短", "riskType": "warranty", "riskLevel": "medium", "clauseNumber": "第六条", "clauseTitle": "质保期限", "originalText": "乙方提供1年免费质保服务", "paragraphId": "p9", "riskReason": "办公设备行业惯例质保期为 3 年，1 年质保期偏短，影响长期使用保障。", "reviewBasis": "《微型计算机商品修理更换退货责任规定》建议主要部件质保 2-3 年。", "suggestion": "建议延长至 3 年，或区分主要部件（3 年）与耗材（1 年）。", "confidence": 0.87, "sourceType": "rule", "ruleId": "RR-009"},
            {"title": "软件知识产权归乙方", "riskType": "ip", "riskLevel": "medium", "clauseNumber": "第七条", "clauseTitle": "知识产权", "originalText": "设备所附软件的知识产权归乙方所有", "paragraphId": "p10", "riskReason": "设备所附软件知识产权归乙方，甲方仅享有使用权，后续升级或迁移可能受限。", "reviewBasis": "《计算机软件保护条例》第 8 条软件著作权归属规则。", "suggestion": "建议明确\"甲方享有永久使用权及免费升级权，乙方不得限制甲方在自有设备上使用\"。", "confidence": 0.86, "sourceType": "ai", "ruleId": None},
            {"title": "交付期限偏长", "riskType": "delivery", "riskLevel": "low", "clauseNumber": "第四条", "clauseTitle": "交付时间", "originalText": "合同签订后30日内完成交付并安装到位", "paragraphId": "p7", "riskReason": "标准办公设备通常 7-15 日可交付，30 日交付期偏长，可能影响办公安排。", "reviewBasis": "行业惯例：现货办公设备交付期一般不超过 15 日。", "suggestion": "建议缩短至 15 日，并约定分批交付的时间节点。", "confidence": 0.78, "sourceType": "ai", "ruleId": None},
            {"title": "保密期限未约定", "riskType": "confidentiality", "riskLevel": "low", "clauseNumber": "第八条", "clauseTitle": "保密条款", "originalText": "双方对合作中知悉的对方商业信息予以保密", "paragraphId": "p11", "riskReason": "保密条款未约定保密期限与违约责任，约束力偏弱。", "reviewBasis": "《反不正当竞争法》第 9 条商业秘密保护要求采取合理保密措施。", "suggestion": "建议明确\"合同终止后 3 年内继续保密，违约赔偿 5 万元/次\"。", "confidence": 0.82, "sourceType": "ai", "ruleId": None},
            {"title": "解约通知期限偏短", "riskType": "termination", "riskLevel": "notice", "clauseNumber": "第十条", "clauseTitle": "合同解除", "originalText": "应提前30日书面通知对方", "paragraphId": "p13", "riskReason": "30 日解约通知期对设备类采购偏短，可能影响甲方备选方案落实。", "reviewBasis": "行业惯例：设备采购合同解约通知期通常为 60-90 日。", "suggestion": "建议延长至 60 日，并约定解约后的设备返还与费用结算流程。", "confidence": 0.75, "sourceType": "ai", "ruleId": None},
        ],
    },
    "sample-2": {
        "contractName": "IT运维服务合同",
        "paragraphs": [
            {"id": "p1", "index": 1, "type": "title", "text": "IT运维服务合同"},
            {"id": "p2", "index": 2, "type": "header", "text": "甲方：智远科技有限公司"},
            {"id": "p3", "index": 3, "type": "header", "text": "乙方：北京云图信息技术有限公司，法定代表人：李某，注册地址：北京市海淀区中关村大街。"},
            {"id": "p4", "index": 4, "type": "body", "clauseNo": "第一条", "clauseTitle": "服务内容", "text": "第一条 服务内容：乙方为甲方提供IT系统运维服务，包括服务器监控、故障处理、系统升级、安全防护等。"},
            {"id": "p5", "index": 5, "type": "body", "clauseNo": "第二条", "clauseTitle": "合同金额", "text": "第二条 合同金额：本合同总金额为人民币680000元（大写：陆拾捌万元整），含税。"},
            {"id": "p6", "index": 6, "type": "body", "clauseNo": "第三条", "clauseTitle": "付款方式", "text": "第三条 付款方式：甲方按季度付款，每季度末付25%，乙方开具增值税专用发票。"},
            {"id": "p7", "index": 7, "type": "body", "clauseNo": "第四条", "clauseTitle": "服务期限", "text": "第四条 服务期限：本合同服务期限为1年，自双方签字之日起计算。"},
            {"id": "p8", "index": 8, "type": "body", "clauseNo": "第五条", "clauseTitle": "服务等级", "text": "第五条 服务等级：乙方承诺核心系统可用性不低于99.5%，故障响应时间不超过4小时。"},
            {"id": "p9", "index": 9, "type": "body", "clauseNo": "第六条", "clauseTitle": "知识产权", "text": "第六条 知识产权：乙方在服务过程中产生的所有成果，知识产权归乙方所有。"},
            {"id": "p10", "index": 10, "type": "body", "clauseNo": "第七条", "clauseTitle": "保密条款", "text": "第七条 保密条款：双方应对在合作中知悉的对方商业信息予以保密。"},
            {"id": "p11", "index": 11, "type": "body", "clauseNo": "第八条", "clauseTitle": "违约责任", "text": "第八条 违约责任：乙方未按约定提供服务，每次扣减当季服务费的5%。"},
            {"id": "p12", "index": 12, "type": "body", "clauseNo": "第九条", "clauseTitle": "合同解除", "text": "第九条 合同解除：乙方有权根据自身业务情况随时解除本合同，提前7日通知甲方即可。"},
            {"id": "p13", "index": 13, "type": "body", "clauseNo": "第十条", "clauseTitle": "争议解决", "text": "第十条 争议解决：本合同争议提交北京仲裁委员会仲裁。"},
            {"id": "p14", "index": 14, "type": "signature", "text": "甲方（盖章）：智远科技有限公司  乙方（盖章）：北京云图信息技术有限公司  签订日期：2024年7月1日"},
        ],
        "riskTemplates": [
            {"title": "知识产权归属对甲方不利", "riskType": "ip", "riskLevel": "high", "clauseNumber": "第六条", "clauseTitle": "知识产权", "originalText": "知识产权归乙方所有", "paragraphId": "p9", "riskReason": "甲方付费委托的运维成果归乙方所有，甲方丧失使用权与改进权，存在二次付费风险。", "reviewBasis": "《著作权法》第 19 条规定委托作品归属由合同约定；行业惯例建议付费方应取得使用权或共有权。", "suggestion": "建议改为\"成果知识产权归甲方所有\"或\"双方共有，甲方享有永久免费使用权与修改权\"。", "confidence": 0.95, "sourceType": "ai", "ruleId": None},
            {"title": "乙方单方解除权过大", "riskType": "termination", "riskLevel": "high", "clauseNumber": "第九条", "clauseTitle": "合同解除", "originalText": "随时解除本合同，提前7日通知甲方即可", "paragraphId": "p12", "riskReason": "乙方享有过大的单方解除权，仅需 7 日通知，将严重影响甲方 IT 系统连续性。", "reviewBasis": "《民法典》第 563 条合同解除应基于根本违约或约定事由。", "suggestion": "建议删除\"随时解除\"，改为\"乙方需提前 90 日书面通知并配合完成服务交接，否则承担违约责任\"。", "confidence": 0.94, "sourceType": "ai", "ruleId": None},
            {"title": "按季度付款无预付款保障", "riskType": "payment", "riskLevel": "medium", "clauseNumber": "第三条", "clauseTitle": "付款方式", "originalText": "每季度末付25%", "paragraphId": "p6", "riskReason": "按季度付款无前期预付款约束，但乙方需先提供服务，可能存在服务质量下降风险。", "reviewBasis": "《企业财务内控指引》建议服务合同应明确付款节点与服务质量挂钩。", "suggestion": "建议增加\"每季度服务考核合格后付款，不合格扣减相应比例\"。", "confidence": 0.84, "sourceType": "ai", "ruleId": None},
            {"title": "保密期限未约定", "riskType": "confidentiality", "riskLevel": "medium", "clauseNumber": "第七条", "clauseTitle": "保密条款", "originalText": "双方应对在合作中知悉的对方商业信息予以保密", "paragraphId": "p10", "riskReason": "保密条款未约定保密期限，按法律默认为永久保密，实际难以执行。", "reviewBasis": "《反不正当竞争法》第 9 条商业秘密保护要求采取合理保密措施。", "suggestion": "建议明确\"合同终止后 5 年内继续保密\"或\"涉及核心技术的永久保密\"。", "confidence": 0.86, "sourceType": "ai", "ruleId": None},
            {"title": "违约扣费比例过低", "riskType": "breach", "riskLevel": "medium", "clauseNumber": "第八条", "clauseTitle": "违约责任", "originalText": "每次扣减当季服务费的5%", "paragraphId": "p11", "riskReason": "5% 扣费比例过低，对乙方服务质量约束力不足，难以覆盖甲方损失。", "reviewBasis": "行业惯例：服务违约扣费通常为当季服务费的 10%-20%。", "suggestion": "建议提高至 10%-20%，并约定累计违约 3 次甲方可解除合同。", "confidence": 0.85, "sourceType": "rule", "ruleId": "RR-010"},
            {"title": "仲裁地点在北京", "riskType": "dispute", "riskLevel": "medium", "clauseNumber": "第十条", "clauseTitle": "争议解决", "originalText": "提交北京仲裁委员会仲裁", "paragraphId": "p13", "riskReason": "仲裁机构在乙方所在地（北京），甲方异地仲裁成本较高，且仲裁一裁终局无上诉权。", "reviewBasis": "《仲裁法》第 4 条仲裁应双方自愿达成协议。", "suggestion": "建议改为\"由甲方所在地仲裁委员会仲裁\"或\"由合同履行地法院管辖\"。", "confidence": 0.88, "sourceType": "rule", "ruleId": "RR-015"},
            {"title": "服务期限仅1年偏短", "riskType": "term", "riskLevel": "low", "clauseNumber": "第四条", "clauseTitle": "服务期限", "originalText": "本合同服务期限为1年", "paragraphId": "p7", "riskReason": "1 年服务期偏短，频繁招标切换供应商增加管理成本，且影响系统运维连续性。", "reviewBasis": "行业惯例：IT 运维服务合同通常为 2-3 年。", "suggestion": "建议延长至 2-3 年，并约定续约条款。", "confidence": 0.8, "sourceType": "ai", "ruleId": None},
            {"title": "服务范围描述过宽", "riskType": "subject", "riskLevel": "low", "clauseNumber": "第一条", "clauseTitle": "服务内容", "originalText": "包括服务器监控、故障处理、系统升级、安全防护等", "paragraphId": "p4", "riskReason": "\"等\"字表述使服务范围具有开放性，可能产生服务边界争议。", "reviewBasis": "《民法典》第 470 条合同标的应明确具体。", "suggestion": "建议删除\"等\"字，明确列出全部服务项并附服务清单。", "confidence": 0.78, "sourceType": "ai", "ruleId": None},
            {"title": "税率未明确", "riskType": "warranty", "riskLevel": "notice", "clauseNumber": "第二条", "clauseTitle": "合同金额", "originalText": "含税", "paragraphId": "p5", "riskReason": "仅约定\"含税\"未明确税率，开票时可能产生税额争议。", "reviewBasis": "《增值税暂行条例》建议合同明确税率。", "suggestion": "建议明确\"含 6% 增值税\"或\"含 13% 增值税\"。", "confidence": 0.7, "sourceType": "ai", "ruleId": None},
        ],
    },
    "sample-3": {
        "contractName": "云服务器租赁合同",
        "paragraphs": [
            {"id": "p1", "index": 1, "type": "title", "text": "云服务器租赁合同"},
            {"id": "p2", "index": 2, "type": "header", "text": "甲方：智远科技有限公司"},
            {"id": "p3", "index": 3, "type": "header", "text": "乙方：杭州数云科技有限公司，法定代表人：张某，注册地址：杭州市余杭区文一西路。"},
            {"id": "p4", "index": 4, "type": "body", "clauseNo": "第一条", "clauseTitle": "租赁标的", "text": "第一条 租赁标的：乙方为甲方提供云服务器租赁服务，规格为8核16G配置，存储500G。"},
            {"id": "p5", "index": 5, "type": "body", "clauseNo": "第二条", "clauseTitle": "合同金额", "text": "第二条 合同金额：人民币45万元整，租赁期限为3年，按年付费。"},
            {"id": "p6", "index": 6, "type": "body", "clauseNo": "第三条", "clauseTitle": "数据安全", "text": "第三条 数据安全：乙方应保障甲方数据安全，发生安全事件应及时处理。"},
            {"id": "p7", "index": 7, "type": "body", "clauseNo": "第四条", "clauseTitle": "服务等级", "text": "第四条 服务等级：乙方承诺服务可用性不低于99.9%，月度故障时长不超过30分钟。"},
            {"id": "p8", "index": 8, "type": "body", "clauseNo": "第五条", "clauseTitle": "保密条款", "text": "第五条 保密条款：双方对合作中知悉的对方商业秘密和技术信息予以保密。"},
            {"id": "p9", "index": 9, "type": "body", "clauseNo": "第六条", "clauseTitle": "知识产权", "text": "第六条 知识产权：甲方在云服务器上存储的数据和应用，知识产权归甲方所有。"},
            {"id": "p10", "index": 10, "type": "body", "clauseNo": "第七条", "clauseTitle": "违约责任", "text": "第七条 违约责任：乙方服务中断超过4小时，按月租金的10%支付违约金。"},
            {"id": "p11", "index": 11, "type": "body", "clauseNo": "第八条", "clauseTitle": "合同终止", "text": "第八条 合同终止：合同期内甲方不得提前终止，否则需支付剩余全部租金。"},
            {"id": "p12", "index": 12, "type": "body", "clauseNo": "第九条", "clauseTitle": "数据返还", "text": "第九条 数据返还：合同终止后乙方应在7日内返还甲方全部数据并销毁备份。"},
            {"id": "p13", "index": 13, "type": "body", "clauseNo": "第十条", "clauseTitle": "争议解决", "text": "第十条 争议解决：本合同争议由乙方所在地法院管辖。"},
            {"id": "p14", "index": 14, "type": "signature", "text": "甲方（盖章）：智远科技有限公司  乙方（盖章）：杭州数云科技有限公司  签订日期：2024年5月20日"},
        ],
        "riskTemplates": [
            {"title": "数据安全责任划分不清", "riskType": "data_security", "riskLevel": "high", "clauseNumber": "第三条", "clauseTitle": "数据安全", "originalText": "发生安全事件应及时处理", "paragraphId": "p6", "riskReason": "仅约定\"及时处理\"，未明确赔偿标准、数据泄露通知义务、第三方审计权利。", "reviewBasis": "《数据安全法》第 27 条要求建立数据安全管理制度；《个人信息保护法》第 57 条要求泄露时通知。", "suggestion": "建议明确：乙方承担全额损失赔偿、24 小时内书面通知、配合甲方或第三方审计。", "confidence": 0.92, "sourceType": "ai", "ruleId": None},
            {"title": "甲方提前终止违约金过高", "riskType": "breach", "riskLevel": "high", "clauseNumber": "第八条", "clauseTitle": "合同终止", "originalText": "需支付剩余全部租金", "paragraphId": "p11", "riskReason": "要求甲方支付剩余全部租金（最长 2 年费用），违约金明显过高，可能被法院调减。", "reviewBasis": "《民法典》第 585 条违约金过分高于损失可请求调整。", "suggestion": "建议改为\"提前 30 日通知，支付 3 个月租金作为违约金\"。", "confidence": 0.9, "sourceType": "rule", "ruleId": "RR-010"},
            {"title": "争议管辖对甲方不利", "riskType": "dispute", "riskLevel": "high", "clauseNumber": "第十条", "clauseTitle": "争议解决", "originalText": "本合同争议由乙方所在地法院管辖", "paragraphId": "p13", "riskReason": "管辖法院在乙方所在地（杭州），甲方异地诉讼成本较高。", "reviewBasis": "《民事诉讼法》第 24 条合同纠纷管辖规则。", "suggestion": "建议改为\"由甲方所在地或合同履行地法院管辖\"。", "confidence": 0.93, "sourceType": "rule", "ruleId": "RR-015"},
            {"title": "按年付费无中期解约条款", "riskType": "payment", "riskLevel": "medium", "clauseNumber": "第二条", "clauseTitle": "合同金额", "originalText": "按年付费", "paragraphId": "p5", "riskReason": "按年付费且无中期解约条款，服务质量不达标时甲方难以调整供应商。", "reviewBasis": "《企业财务内控指引》建议长周期服务合同设置中期评估与解约机制。", "suggestion": "建议改为\"按季付费\"或增加\"年度服务质量评估不合格甲方可解约\"。", "confidence": 0.84, "sourceType": "ai", "ruleId": None},
            {"title": "故障时长标准偏宽", "riskType": "acceptance", "riskLevel": "medium", "clauseNumber": "第四条", "clauseTitle": "服务等级", "originalText": "月度故障时长不超过30分钟", "paragraphId": "p7", "riskReason": "月度 30 分钟故障时长对核心业务系统偏宽，可能导致业务中断损失。", "reviewBasis": "行业惯例：核心云服务月度故障时长通常不超过 5 分钟。", "suggestion": "建议缩短至 5-10 分钟，并约定超出时按比例减免月租。", "confidence": 0.86, "sourceType": "ai", "ruleId": None},
            {"title": "保密期限未约定", "riskType": "confidentiality", "riskLevel": "medium", "clauseNumber": "第五条", "clauseTitle": "保密条款", "originalText": "双方对合作中知悉的对方商业秘密和技术信息予以保密", "paragraphId": "p8", "riskReason": "保密条款未约定保密期限与违约责任，约束力偏弱。", "reviewBasis": "《反不正当竞争法》第 9 条商业秘密保护要求。", "suggestion": "建议明确\"合同终止后 5 年内继续保密，违约赔偿 10 万元/次\"。", "confidence": 0.85, "sourceType": "ai", "ruleId": None},
            {"title": "知识产权表述笼统", "riskType": "ip", "riskLevel": "medium", "clauseNumber": "第六条", "clauseTitle": "知识产权", "originalText": "知识产权归甲方所有", "paragraphId": "p9", "riskReason": "虽约定归甲方所有，但未明确乙方运营平台的底层技术归属，可能产生混淆。", "reviewBasis": "《著作权法》第 19 条委托作品归属规则。", "suggestion": "建议区分\"甲方数据与应用归甲方，乙方平台底层技术归乙方\"。", "confidence": 0.83, "sourceType": "ai", "ruleId": None},
            {"title": "配置规格单一无弹性", "riskType": "subject", "riskLevel": "low", "clauseNumber": "第一条", "clauseTitle": "租赁标的", "originalText": "规格为8核16G配置，存储500G", "paragraphId": "p4", "riskReason": "配置固定无弹性升降级机制，业务增长时无法灵活调整资源。", "reviewBasis": "行业惯例：云服务应支持弹性伸缩。", "suggestion": "建议增加\"甲方可在合同期内申请升降级配置，费用按差价结算\"。", "confidence": 0.78, "sourceType": "ai", "ruleId": None},
            {"title": "数据返还期限偏短", "riskType": "warranty", "riskLevel": "low", "clauseNumber": "第九条", "clauseTitle": "数据返还", "originalText": "7日内返还甲方全部数据并销毁备份", "paragraphId": "p12", "riskReason": "7 日返还期偏短，大数据量迁移可能无法完成，且销毁备份缺乏第三方验证。", "reviewBasis": "行业惯例：数据迁移周期通常为 15-30 日。", "suggestion": "建议延长至 30 日，并约定销毁后出具书面证明。", "confidence": 0.8, "sourceType": "rule", "ruleId": "RR-009"},
            {"title": "违约金计算基准不明确", "riskType": "term", "riskLevel": "notice", "clauseNumber": "第七条", "clauseTitle": "违约责任", "originalText": "按月租金的10%支付违约金", "paragraphId": "p10", "riskReason": "月租金为合同总额/36，10% 仅约 1250 元，违约金偏低缺乏威慑力。", "reviewBasis": "《民法典》第 585 条违约金应与损失相当。", "suggestion": "建议改为\"按合同总额的 1% 支付违约金，或实际损失较高时按实际损失赔偿\"。", "confidence": 0.72, "sourceType": "ai", "ruleId": None},
        ],
    },
    "sample-4": {
        "contractName": "人力资源外包合同",
        "paragraphs": [
            {"id": "p1", "index": 1, "type": "title", "text": "人力资源外包合同"},
            {"id": "p2", "index": 2, "type": "header", "text": "甲方：智远科技有限公司"},
            {"id": "p3", "index": 3, "type": "header", "text": "乙方：广州人才之星人力资源服务有限公司，法定代表人：陈某，注册地址：广州市天河区珠江新城。"},
            {"id": "p4", "index": 4, "type": "body", "clauseNo": "第一条", "clauseTitle": "服务内容", "text": "第一条 服务内容：乙方为甲方提供招聘、薪酬核算、社保代缴、员工关系管理等人力资源外包服务。"},
            {"id": "p5", "index": 5, "type": "body", "clauseNo": "第二条", "clauseTitle": "合同金额", "text": "第二条 合同金额：人民币32万元整，按月支付，每月末付当月服务费。"},
            {"id": "p6", "index": 6, "type": "body", "clauseNo": "第三条", "clauseTitle": "服务期限", "text": "第三条 服务期限：本合同有效期为2年，自双方签字之日起计算。"},
            {"id": "p7", "index": 7, "type": "body", "clauseNo": "第四条", "clauseTitle": "保密条款", "text": "第四条 保密条款：乙方对甲方员工个人信息负有保密义务。"},
            {"id": "p8", "index": 8, "type": "body", "clauseNo": "第五条", "clauseTitle": "数据保护", "text": "第五条 数据保护：乙方处理员工个人信息应遵守《个人信息保护法》相关规定。"},
            {"id": "p9", "index": 9, "type": "body", "clauseNo": "第六条", "clauseTitle": "验收方式", "text": "第六条 验收方式：乙方每月提交服务报告，甲方应在3个工作日内确认。"},
            {"id": "p10", "index": 10, "type": "body", "clauseNo": "第七条", "clauseTitle": "违约责任", "text": "第七条 违约责任：乙方泄露员工信息，每次支付违约金5万元。"},
            {"id": "p11", "index": 11, "type": "body", "clauseNo": "第八条", "clauseTitle": "合同解除", "text": "第八条 合同解除：甲方提前30日书面通知可解除合同，已付费用不予退还。"},
            {"id": "p12", "index": 12, "type": "body", "clauseNo": "第九条", "clauseTitle": "知识产权", "text": "第九条 知识产权：乙方在服务中形成的薪酬体系、流程文档归乙方所有。"},
            {"id": "p13", "index": 13, "type": "body", "clauseNo": "第十条", "clauseTitle": "争议解决", "text": "第十条 争议解决：本合同争议由乙方所在地法院管辖。"},
            {"id": "p14", "index": 14, "type": "signature", "text": "甲方（盖章）：智远科技有限公司  乙方（盖章）：广州人才之星人力资源服务有限公司  签订日期：2024年3月10日"},
        ],
        "riskTemplates": [
            {"title": "保密条款缺乏约束力", "riskType": "confidentiality", "riskLevel": "high", "clauseNumber": "第四条", "clauseTitle": "保密条款", "originalText": "乙方对甲方员工个人信息负有保密义务", "paragraphId": "p7", "riskReason": "仅约定\"保密义务\"，未明确违约金、保密期限、违约后果，约束力弱。", "reviewBasis": "《个人信息保护法》第 21 条要求委托处理个人信息应约定双方权利义务。", "suggestion": "建议明确\"保密期限为合同终止后 5 年，违约赔偿 10 万元/次\"。", "confidence": 0.9, "sourceType": "ai", "ruleId": None},
            {"title": "数据保护条款笼统", "riskType": "data_security", "riskLevel": "high", "clauseNumber": "第五条", "clauseTitle": "数据保护", "originalText": "应遵守《个人信息保护法》相关规定", "paragraphId": "p8", "riskReason": "仅引用法律名称，未约定具体保护措施、数据泄露通知义务、赔偿标准。", "reviewBasis": "《个人信息保护法》第 51 条要求采取相应技术措施和必要措施保障安全。", "suggestion": "建议明确：数据加密存储、24 小时内通知泄露、按实际损失全额赔偿。", "confidence": 0.89, "sourceType": "ai", "ruleId": None},
            {"title": "争议管辖对甲方不利", "riskType": "dispute", "riskLevel": "high", "clauseNumber": "第十条", "clauseTitle": "争议解决", "originalText": "本合同争议由乙方所在地法院管辖", "paragraphId": "p13", "riskReason": "乙方所在地（广州）距甲方较远，诉讼成本高。", "reviewBasis": "《民事诉讼法》第 24 条合同纠纷管辖规则。", "suggestion": "建议改为\"由甲方所在地或合同履行地法院管辖\"。", "confidence": 0.92, "sourceType": "rule", "ruleId": "RR-015"},
            {"title": "付款方式无对赌条款", "riskType": "payment", "riskLevel": "medium", "clauseNumber": "第二条", "clauseTitle": "合同金额", "originalText": "每月末付当月服务费", "paragraphId": "p5", "riskReason": "按月付费无服务质量考核挂钩，乙方服务质量下降时甲方缺乏制约。", "reviewBasis": "《企业财务内控指引》建议服务合同应与服务质量挂钩。", "suggestion": "建议增加\"月度服务考核合格后付款，不合格扣减相应比例\"。", "confidence": 0.84, "sourceType": "ai", "ruleId": None},
            {"title": "知识产权归乙方", "riskType": "ip", "riskLevel": "medium", "clauseNumber": "第九条", "clauseTitle": "知识产权", "originalText": "薪酬体系、流程文档归乙方所有", "paragraphId": "p12", "riskReason": "甲方付费委托形成的薪酬体系归乙方，合同终止后甲方无法继续使用。", "reviewBasis": "《著作权法》第 19 条委托作品归属规则。", "suggestion": "建议改为\"归甲方所有\"或\"甲方享有永久免费使用权\"。", "confidence": 0.87, "sourceType": "ai", "ruleId": None},
            {"title": "违约金偏低", "riskType": "breach", "riskLevel": "medium", "clauseNumber": "第七条", "clauseTitle": "违约责任", "originalText": "每次支付违约金5万元", "paragraphId": "p10", "riskReason": "5 万元违约金相对员工信息泄露可能造成的损失偏低，威慑力不足。", "reviewBasis": "《个人信息保护法》第 66 条罚款上限可达 100 万元。", "suggestion": "建议提高至 20 万元/次，并约定承担甲方因此遭受的全部损失。", "confidence": 0.86, "sourceType": "rule", "ruleId": "RR-010"},
            {"title": "验收期限过短", "riskType": "acceptance", "riskLevel": "medium", "clauseNumber": "第六条", "clauseTitle": "验收方式", "originalText": "甲方应在3个工作日内确认", "paragraphId": "p9", "riskReason": "3 个工作日确认期偏短，HR 部门难以全面核对薪酬社保数据。", "reviewBasis": "行业惯例：人力资源服务验收期通常为 5-10 个工作日。", "suggestion": "建议延长至 7-10 个工作日，并约定\"期内未提出异议视为确认\"。", "confidence": 0.83, "sourceType": "ai", "ruleId": None},
            {"title": "服务范围描述模糊", "riskType": "subject", "riskLevel": "low", "clauseNumber": "第一条", "clauseTitle": "服务内容", "originalText": "员工关系管理等人力资源外包服务", "paragraphId": "p4", "riskReason": "\"等\"字表述使服务范围具有开放性，可能产生额外服务争议。", "reviewBasis": "《民法典》第 470 条合同标的应明确具体。", "suggestion": "建议删除\"等\"字，明确列出全部服务项并附服务清单。", "confidence": 0.78, "sourceType": "ai", "ruleId": None},
            {"title": "已付费用不退显失公平", "riskType": "termination", "riskLevel": "low", "clauseNumber": "第八条", "clauseTitle": "合同解除", "originalText": "已付费用不予退还", "paragraphId": "p11", "riskReason": "甲方提前终止即损失全部已付费用，可能被认定为显失公平条款。", "reviewBasis": "《民法典》第 496-498 条格式条款效力规则。", "suggestion": "建议改为\"按已服务月份比例结算，未服务部分予以退还\"。", "confidence": 0.82, "sourceType": "ai", "ruleId": None},
            {"title": "服务期限2年偏长", "riskType": "term", "riskLevel": "notice", "clauseNumber": "第三条", "clauseTitle": "服务期限", "originalText": "本合同有效期为2年", "paragraphId": "p6", "riskReason": "人力资源外包 2 年期限偏长，缺乏中期评估机制，服务质量下降时难以调整。", "reviewBasis": "行业惯例：人力资源外包合同通常为 1 年并约定续约评估。", "suggestion": "建议缩短至 1 年，并约定年度评估合格后自动续约。", "confidence": 0.75, "sourceType": "ai", "ruleId": None},
            {"title": "合同金额未明确税费承担", "riskType": "amount", "riskLevel": "low", "clauseNumber": "第二条", "clauseTitle": "合同金额", "originalText": "人民币32万元整", "paragraphId": "p5", "riskReason": "合同金额未明确是否含税及税率，按月支付时可能产生税额争议与开票纠纷。", "reviewBasis": "《增值税暂行条例》建议服务合同明确税率与含税口径。", "suggestion": "建议明确\"含 6% 增值税，乙方开具增值税专用发票\"。", "confidence": 0.78, "sourceType": "ai", "ruleId": None},
        ],
    },
}

# ===== 预设审核任务（15 个，类型多样，日期跨6个月）=====
DEMO_TASKS = [
    {
        "id": "RVT-DEMO-001", "contractId": "C-001", "contractName": "软件系统采购合同",
        "contractNo": "HT-CG-2026-001", "counterparty": "星河软件有限公司",
        "amount": 580000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "采购部",
        "reviewFocus": ["subject", "payment", "delivery", "breach", "ip", "confidentiality", "termination", "dispute"],
        "reviewNote": "本次采购为协同办公平台，需重点关注预付款保障与知识产权归属。",
        "fileName": "软件系统采购合同.docx", "fileSize": 248 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_business", "riskLevelMax": "high",
        "riskCount": {"high": 10, "medium": 6, "low": 1, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-06-20T09:30:00.000Z", "updatedAt": "2026-06-20T09:42:00.000Z",
    },
    {
        "id": "RVT-DEMO-002", "contractId": "C-002", "contractName": "办公设备批量采购合同",
        "contractNo": "HT-CG-2026-002", "counterparty": "华盛办公设备有限公司",
        "amount": 320000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "采购部",
        "reviewFocus": ["payment", "delivery", "breach"],
        "reviewNote": "办公设备批量采购，关注交付与验收。",
        "fileName": "办公设备采购合同.pdf", "fileSize": 512 * 1024, "sampleId": "sample-1",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_legal", "riskLevelMax": "high",
        "riskCount": {"high": 3, "medium": 4, "low": 1, "notice": 0},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": "2026-06-18T14:00:00.000Z", "completedAt": None,
        "createdAt": "2026-06-18T10:00:00.000Z", "updatedAt": "2026-06-18T14:00:00.000Z",
    },
    {
        "id": "RVT-DEMO-003", "contractId": "C-003", "contractName": "云服务采购合同",
        "contractNo": "HT-CG-2026-003", "counterparty": "云栈科技有限公司",
        "amount": 860000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "信息技术部",
        "reviewFocus": ["payment", "confidentiality", "ip"],
        "reviewNote": "云服务采购，关注数据安全与保密。",
        "fileName": "云服务采购合同.docx", "fileSize": 320 * 1024, "sampleId": "sample-4",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "completed", "riskLevelMax": "medium",
        "riskCount": {"high": 0, "medium": 4, "low": 2, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": "已按建议修改付款节点与数据安全条款，风险可控。",
        "legalConclusion": "sign_after_modify",
        "legalReviewerId": "U-LEGAL", "legalReviewerName": "王律师",
        "submittedAt": "2026-06-10T11:00:00.000Z", "completedAt": "2026-06-11T15:30:00.000Z",
        "createdAt": "2026-06-09T09:00:00.000Z", "updatedAt": "2026-06-11T15:30:00.000Z",
    },
    {
        "id": "RVT-DEMO-004", "contractId": "C-004", "contractName": "办公耗材采购合同",
        "contractNo": "HT-CG-2026-004", "counterparty": "日常文具商行",
        "amount": 28000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "采购部",
        "reviewFocus": ["payment"],
        "reviewNote": "",
        "fileName": "耗材采购合同.pdf", "fileSize": 180 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "failed", "riskLevelMax": None,
        "riskCount": {"high": 0, "medium": 0, "low": 0, "notice": 0},
        "progress": 35, "currentStage": "parse",
        "errorCode": "PARSE_FAILED", "errorMsg": "文档解析失败：检测到加密PDF，无法提取文本内容。",
        "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-06-22T16:00:00.000Z", "updatedAt": "2026-06-22T16:05:00.000Z",
    },
    {"id": "RVT-DEMO-005", "contractId": "C-005", "contractName": "IT运维服务外包合同",
        "contractNo": "HT-CG-2026-005", "counterparty": "运维技术服务部",
        "amount": 150000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "信息技术部",
        "reviewFocus": ["subject", "breach", "dispute"],
        "reviewNote": "运维外包，待补充附件。",
        "fileName": "IT运维服务外包合同.docx", "fileSize": 96 * 1024, "sampleId": "sample-2",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "draft", "riskLevelMax": None,
        "riskCount": {"high": 0, "medium": 0, "low": 0, "notice": 0},
        "progress": 0, "currentStage": "upload",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-06-25T10:00:00.000Z", "updatedAt": "2026-06-25T10:00:00.000Z",
    },
    {"id": "RVT-DEMO-006", "contractId": "C-006", "contractName": "生产设备采购合同",
        "contractNo": "HT-CG-2026-006", "counterparty": "精工机械设备有限公司",
        "amount": 450000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "生产部",
        "reviewFocus": ["subject", "payment", "delivery", "acceptance", "breach"],
        "reviewNote": "生产设备采购，关注交付周期与验收节点。",
        "fileName": "生产设备采购合同.pdf", "fileSize": 280 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_business", "riskLevelMax": "high",
        "riskCount": {"high": 5, "medium": 4, "low": 2, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-05-15T14:00:00.000Z", "updatedAt": "2026-05-15T14:20:00.000Z",
    },
    {"id": "RVT-DEMO-007", "contractId": "C-007", "contractName": "市场宣传设计服务合同",
        "contractNo": "HT-CG-2026-007", "counterparty": "创艺广告设计有限公司",
        "amount": 180000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "市场部",
        "reviewFocus": ["subject", "payment", "ip", "dispute"],
        "reviewNote": "公司年度宣传物料设计，需明确知识产权归属。",
        "fileName": "宣传设计服务合同.docx", "fileSize": 160 * 1024, "sampleId": "sample-3",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_business", "riskLevelMax": "medium",
        "riskCount": {"high": 0, "medium": 6, "low": 3, "notice": 0},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-05-20T10:30:00.000Z", "updatedAt": "2026-05-20T11:00:00.000Z",
    },
    {"id": "RVT-DEMO-008", "contractId": "C-008", "contractName": "物流运输服务合同",
        "contractNo": "HT-CG-2026-008", "counterparty": "速达物流有限公司",
        "amount": 250000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "物流部",
        "reviewFocus": ["payment", "delivery", "breach", "confidentiality"],
        "reviewNote": "年度物流运输框架协议，关注赔付标准及保密条款。",
        "fileName": "物流运输服务合同.pdf", "fileSize": 200 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "completed", "riskLevelMax": "medium",
        "riskCount": {"high": 0, "medium": 5, "low": 2, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": "物流赔付标准已调整至行业合理水平，同意签署。",
        "legalConclusion": "sign",
        "legalReviewerId": "U-LEGAL", "legalReviewerName": "王律师",
        "submittedAt": "2026-04-11T09:00:00.000Z", "completedAt": "2026-04-12T16:00:00.000Z",
        "createdAt": "2026-04-10T08:30:00.000Z", "updatedAt": "2026-04-12T16:00:00.000Z",
    },
    {"id": "RVT-DEMO-009", "contractId": "C-009", "contractName": "年度IT维保服务合同",
        "contractNo": "HT-CG-2026-009", "counterparty": "鼎新信息技术有限公司",
        "amount": 680000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "信息技术部",
        "reviewFocus": ["subject", "payment", "warranty", "confidentiality", "data_security", "dispute"],
        "reviewNote": "公司年度IT基础设施维保，重点关注服务SLA与数据安全保障。",
        "fileName": "IT维保服务合同.docx", "fileSize": 380 * 1024, "sampleId": "sample-4",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "completed", "riskLevelMax": "medium",
        "riskCount": {"high": 0, "medium": 7, "low": 3, "notice": 2},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": "SLA条款已明确，数据安全条款完善，建议签署。",
        "legalConclusion": "sign",
        "legalReviewerId": "U-LEGAL", "legalReviewerName": "王律师",
        "submittedAt": "2026-04-09T10:00:00.000Z", "completedAt": "2026-04-08T17:30:00.000Z",
        "createdAt": "2026-04-08T09:00:00.000Z", "updatedAt": "2026-04-08T17:30:00.000Z",
    },
    {"id": "RVT-DEMO-010", "contractId": "C-010", "contractName": "办公家具采购合同",
        "contractNo": "HT-CG-2026-010", "counterparty": "雅致办公家具有限公司",
        "amount": 95000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "行政部",
        "reviewFocus": ["payment", "delivery", "acceptance"],
        "reviewNote": "新办公区家具采购，已在前期考察确认。",
        "fileName": "办公家具采购合同.pdf", "fileSize": 140 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_legal", "riskLevelMax": "low",
        "riskCount": {"high": 0, "medium": 0, "low": 3, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": "2026-04-25T16:30:00.000Z", "completedAt": None,
        "createdAt": "2026-04-25T14:00:00.000Z", "updatedAt": "2026-04-25T16:30:00.000Z",
    },
    {"id": "RVT-DEMO-011", "contractId": "C-011", "contractName": "网络设备采购合同",
        "contractNo": "HT-CG-2026-011", "counterparty": "网域科技股份有限公司",
        "amount": 420000, "currency": "CNY", "contractType": "采购合同", "myRole": "buyer",
        "department": "信息技术部",
        "reviewFocus": ["subject", "payment", "delivery", "acceptance", "breach"],
        "reviewNote": "数据中心网络设备升级，关注交付时间和技术验收标准。",
        "fileName": "网络设备采购合同.docx", "fileSize": 260 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_business", "riskLevelMax": "high",
        "riskCount": {"high": 4, "medium": 5, "low": 1, "notice": 0},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-03-18T11:00:00.000Z", "updatedAt": "2026-03-18T11:30:00.000Z",
    },
    {"id": "RVT-DEMO-012", "contractId": "C-012", "contractName": "广告投放代理合同",
        "contractNo": "HT-CG-2026-012", "counterparty": "明锐传媒有限公司",
        "amount": 350000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "市场部",
        "reviewFocus": ["subject", "payment", "breach", "termination", "dispute"],
        "reviewNote": "Q2线上广告投放代理，关注投放效果KPI与违约责任。",
        "fileName": "广告投放代理合同.pdf", "fileSize": 220 * 1024, "sampleId": "sample-3",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "completed", "riskLevelMax": "high",
        "riskCount": {"high": 2, "medium": 5, "low": 1, "notice": 0},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": "KPI考核条款已调整，建议签署。",
        "legalConclusion": "sign_after_modify",
        "legalReviewerId": "U-LEGAL", "legalReviewerName": "王律师",
        "submittedAt": "2026-03-11T14:00:00.000Z", "completedAt": "2026-03-10T17:00:00.000Z",
        "createdAt": "2026-03-10T09:00:00.000Z", "updatedAt": "2026-03-10T17:00:00.000Z",
    },
    {"id": "RVT-DEMO-013", "contractId": "C-013", "contractName": "员工培训服务合同",
        "contractNo": "HT-CG-2026-013", "counterparty": "知行管理咨询有限公司",
        "amount": 120000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "人力资源部",
        "reviewFocus": ["subject", "payment", "ip"],
        "reviewNote": "中高层管理技能提升培训，待与供应商确认课程大纲。",
        "fileName": "培训服务合同.docx", "fileSize": 88 * 1024, "sampleId": "sample-3",
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "draft", "riskLevelMax": None,
        "riskCount": {"high": 0, "medium": 0, "low": 0, "notice": 0},
        "progress": 0, "currentStage": "upload",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-02-20T15:00:00.000Z", "updatedAt": "2026-02-20T15:00:00.000Z",
    },
    {"id": "RVT-DEMO-014", "contractId": "C-014", "contractName": "法律顾问服务合同",
        "contractNo": "HT-CG-2026-014", "counterparty": "正和法律事务所",
        "amount": 200000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "法务部",
        "reviewFocus": ["subject", "payment", "confidentiality", "dispute"],
        "reviewNote": "年度法律顾问服务，关注服务范围及保密义务。",
        "fileName": "法律顾问服务合同.docx", "fileSize": 120 * 1024, "sampleId": None,
        "creatorId": "U-LEGAL", "creatorName": "王律师",
        "status": "completed", "riskLevelMax": "low",
        "riskCount": {"high": 0, "medium": 0, "low": 2, "notice": 1},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": True,
        "legalOpinion": "服务范围与费用标准已明确，建议签署。",
        "legalConclusion": "sign",
        "legalReviewerId": "U-LEGAL", "legalReviewerName": "王律师",
        "submittedAt": "2026-02-17T11:00:00.000Z", "completedAt": "2026-02-15T16:00:00.000Z",
        "createdAt": "2026-02-15T09:00:00.000Z", "updatedAt": "2026-02-15T16:00:00.000Z",
    },
    {"id": "RVT-DEMO-015", "contractId": "C-015", "contractName": "保洁服务外包合同",
        "contractNo": "HT-CG-2026-015", "counterparty": "洁美保洁服务有限公司",
        "amount": 86000, "currency": "CNY", "contractType": "服务合同", "myRole": "buyer",
        "department": "行政部",
        "reviewFocus": ["subject", "payment", "breach"],
        "reviewNote": "年度办公区保洁外包服务合同。",
        "fileName": "保洁外包合同.pdf", "fileSize": 72 * 1024, "sampleId": None,
        "creatorId": "U-PURCHASER", "creatorName": "李明",
        "status": "pending_business", "riskLevelMax": "low",
        "riskCount": {"high": 0, "medium": 0, "low": 2, "notice": 0},
        "progress": 100, "currentStage": "result",
        "errorCode": None, "errorMsg": None, "fieldsConfirmed": False,
        "legalOpinion": None, "legalConclusion": None,
        "legalReviewerId": None, "legalReviewerName": None,
        "submittedAt": None, "completedAt": None,
        "createdAt": "2026-01-12T10:00:00.000Z", "updatedAt": "2026-01-12T10:30:00.000Z",
    },
]

# ===== 演示审计日志（操作时间轴数据）=====
# 按任务状态还原操作流程
# T1: 创建 → AI审核完成（待人工确认）
# T2: 创建 → AI审核完成 → 提交法务复核
# T3: 创建 → AI审核完成 → 提交法务复核 → 法务审核通过
# T4: 创建 → 解析失败
# T5: 创建（草稿，未开始审核）
DEMO_AUDIT_LOGS = [
    # --- T1: RVT-DEMO-001 ---
    {"id": "AL-001-001", "reviewTaskId": "RVT-DEMO-001", "objectType": "task",
     "objectId": "RVT-DEMO-001", "action": "创建审核任务",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": None, "afterState": "草稿",
     "remark": "合同名称：软件系统采购合同\n相对方：星河软件有限公司\n合同金额：580000 元",
     "createdAt": "2026-06-20T09:30:00.000Z"},
    {"id": "AL-001-002", "reviewTaskId": "RVT-DEMO-001", "objectType": "task",
     "objectId": "RVT-DEMO-001", "action": "开始AI审核",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": "草稿", "afterState": "解析中",
     "remark": "启动文档解析与AI审核流程",
     "createdAt": "2026-06-20T09:30:05.000Z"},
    {"id": "AL-001-003", "reviewTaskId": "RVT-DEMO-001", "objectType": "task",
     "objectId": "RVT-DEMO-001", "action": "AI审核完成",
     "operatorId": "system", "operatorName": "AI 系统（DeepSeek）",
     "beforeState": "AI审核中", "afterState": "待人工确认",
     "remark": "本次共识别 18 项风险，其中高风险 10 项、中风险 6 项、低风险 1 项、提示 1 项。规则库命中 8 项。",
     "createdAt": "2026-06-20T09:42:00.000Z"},

    # --- T2: RVT-DEMO-002 ---
    {"id": "AL-002-001", "reviewTaskId": "RVT-DEMO-002", "objectType": "task",
     "objectId": "RVT-DEMO-002", "action": "创建审核任务",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": None, "afterState": "草稿",
     "remark": "合同名称：办公设备批量采购合同\n相对方：华盛办公设备有限公司\n合同金额：320000 元",
     "createdAt": "2026-06-18T10:00:00.000Z"},
    {"id": "AL-002-002", "reviewTaskId": "RVT-DEMO-002", "objectType": "task",
     "objectId": "RVT-DEMO-002", "action": "开始AI审核",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": "草稿", "afterState": "解析中",
     "remark": "启动文档解析与AI审核流程",
     "createdAt": "2026-06-18T10:00:05.000Z"},
    {"id": "AL-002-003", "reviewTaskId": "RVT-DEMO-002", "objectType": "task",
     "objectId": "RVT-DEMO-002", "action": "AI审核完成",
     "operatorId": "system", "operatorName": "AI 系统（DeepSeek）",
     "beforeState": "AI审核中", "afterState": "待人工确认",
     "remark": "本次共识别 8 项风险，其中高风险 3 项、中风险 4 项、低风险 1 项。",
     "createdAt": "2026-06-18T11:00:00.000Z"},
    {"id": "AL-002-004", "reviewTaskId": "RVT-DEMO-002", "objectType": "task",
     "objectId": "RVT-DEMO-002", "action": "提交法务复核",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": "待人工确认", "afterState": "待法务复核",
     "remark": "风险处理汇总：共 8 项，已接受 3 项、已编辑 2 项、已忽略 2 项、转人工 1 项",
     "createdAt": "2026-06-18T14:00:00.000Z"},

    # --- T3: RVT-DEMO-003 ---
    {"id": "AL-003-001", "reviewTaskId": "RVT-DEMO-003", "objectType": "task",
     "objectId": "RVT-DEMO-003", "action": "创建审核任务",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": None, "afterState": "草稿",
     "remark": "合同名称：云服务采购合同\n相对方：云栈科技有限公司\n合同金额：860000 元",
     "createdAt": "2026-06-09T09:00:00.000Z"},
    {"id": "AL-003-002", "reviewTaskId": "RVT-DEMO-003", "objectType": "task",
     "objectId": "RVT-DEMO-003", "action": "开始AI审核",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": "草稿", "afterState": "解析中",
     "remark": "启动文档解析与AI审核流程",
     "createdAt": "2026-06-09T09:00:05.000Z"},
    {"id": "AL-003-003", "reviewTaskId": "RVT-DEMO-003", "objectType": "task",
     "objectId": "RVT-DEMO-003", "action": "AI审核完成",
     "operatorId": "system", "operatorName": "AI 系统（DeepSeek）",
     "beforeState": "AI审核中", "afterState": "待人工确认",
     "remark": "本次共识别 7 项风险，其中高风险 0 项、中风险 4 项、低风险 2 项、提示 1 项。",
     "createdAt": "2026-06-09T10:00:00.000Z"},
    {"id": "AL-003-004", "reviewTaskId": "RVT-DEMO-003", "objectType": "task",
     "objectId": "RVT-DEMO-003", "action": "提交法务复核",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": "待人工确认", "afterState": "待法务复核",
     "remark": "风险处理汇总：共 7 项，已确认 7 项",
     "createdAt": "2026-06-10T11:00:00.000Z"},
    {"id": "AL-003-005", "reviewTaskId": "RVT-DEMO-003", "objectType": "task",
     "objectId": "RVT-DEMO-003", "action": "法务审核通过",
     "operatorId": "U-LEGAL", "operatorName": "王律师",
     "beforeState": "待法务复核", "afterState": "已完成",
     "remark": "审核结论：建议修改后签署\n法务意见：已按建议修改付款节点与数据安全条款，风险可控。\n法务审核人：王律师（法务部）\n风险处理汇总：共 7 项，已确认 7 项",
     "createdAt": "2026-06-11T15:30:00.000Z"},

    # --- T4: RVT-DEMO-004 ---
    {"id": "AL-004-001", "reviewTaskId": "RVT-DEMO-004", "objectType": "task",
     "objectId": "RVT-DEMO-004", "action": "创建审核任务",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": None, "afterState": "草稿",
     "remark": "合同名称：生产设备采购合同\n相对方：精密机械有限公司\n合同金额：450000 元",
     "createdAt": "2026-06-22T16:00:00.000Z"},
    {"id": "AL-004-002", "reviewTaskId": "RVT-DEMO-004", "objectType": "task",
     "objectId": "RVT-DEMO-004", "action": "解析失败",
     "operatorId": "system", "operatorName": "系统",
     "beforeState": "解析中", "afterState": "失败",
     "remark": "文档解析失败：检测到加密PDF，无法提取文本内容。",
     "createdAt": "2026-06-22T16:05:00.000Z"},

    # --- T5: RVT-DEMO-005（仅创建） ---
    {"id": "AL-005-001", "reviewTaskId": "RVT-DEMO-005", "objectType": "task",
     "objectId": "RVT-DEMO-005", "action": "创建审核任务",
     "operatorId": "U-PURCHASER", "operatorName": "李明",
     "beforeState": None, "afterState": "草稿",
     "remark": "合同名称：IT运维服务外包合同\n相对方：运维技术服务部\n合同金额：150000 元",
     "createdAt": "2026-06-25T10:00:00.000Z"},
]

# ===== 抽取字段模板（T1/T2/T3 各一套）=====
DEMO_EXTRACTED_FIELDS = [
    {"fieldKey": "contractName", "fieldLabel": "合同名称", "fieldValue": "软件系统采购合同", "confidence": 0.98, "lowConfidence": False, "sourceText": "软件系统采购合同"},
    {"fieldKey": "buyer", "fieldLabel": "甲方", "fieldValue": "智远科技有限公司", "confidence": 0.97, "lowConfidence": False, "sourceText": "甲方（采购方）：智远科技有限公司"},
    {"fieldKey": "seller", "fieldLabel": "乙方", "fieldValue": "星河软件有限公司", "confidence": 0.88, "lowConfidence": False, "sourceText": "乙方（供应方）：星河软件有限公司"},
    {"fieldKey": "contractNo", "fieldLabel": "合同编号", "fieldValue": "RVT-2026-001", "confidence": 0.55, "lowConfidence": True, "sourceText": "（系统生成编号）"},
    {"fieldKey": "amount", "fieldLabel": "合同金额", "fieldValue": "580000", "confidence": 0.95, "lowConfidence": False, "sourceText": "合同总金额为人民币580000元"},
    {"fieldKey": "currency", "fieldLabel": "币种", "fieldValue": "CNY", "confidence": 0.99, "lowConfidence": False, "sourceText": "人民币"},
    {"fieldKey": "taxRate", "fieldLabel": "税率", "fieldValue": "13%", "confidence": 0.72, "lowConfidence": True, "sourceText": "上述金额已含增值税"},
    {"fieldKey": "signDate", "fieldLabel": "签约日期", "fieldValue": "2026-06-15", "confidence": 0.65, "lowConfidence": True, "sourceText": "（未在正文明确，依据创建时间推断）"},
    {"fieldKey": "effectiveDate", "fieldLabel": "生效日期", "fieldValue": "2026-06-15", "confidence": 0.9, "lowConfidence": False, "sourceText": "本合同自双方签字盖章之日起生效"},
    {"fieldKey": "term", "fieldLabel": "合同期限", "fieldValue": "2年", "confidence": 0.96, "lowConfidence": False, "sourceText": "有效期为2年"},
    {"fieldKey": "paymentMethod", "fieldLabel": "付款方式", "fieldValue": "签约后付80%，验收后付20%", "confidence": 0.94, "lowConfidence": False, "sourceText": "支付合同总额的80%作为预付款"},
    {"fieldKey": "deliveryDate", "fieldLabel": "交付时间", "fieldValue": "尽快", "confidence": 0.4, "lowConfidence": True, "sourceText": "尽快完成系统的交付与安装部署"},
    {"fieldKey": "acceptanceMethod", "fieldLabel": "验收方式", "fieldValue": "符合甲方要求并签署验收报告", "confidence": 0.6, "lowConfidence": True, "sourceText": "应符合甲方要求，经甲方确认后签署验收报告"},
    {"fieldKey": "warrantyPeriod", "fieldLabel": "质保期限", "fieldValue": "未约定", "confidence": 0.3, "lowConfidence": True, "sourceText": "具体质保期限及响应时限由双方另行约定"},
    {"fieldKey": "jurisdiction", "fieldLabel": "争议管辖", "fieldValue": "乙方所在地法院", "confidence": 0.93, "lowConfidence": False, "sourceText": "由乙方所在地有管辖权的人民法院管辖"},
]

# ===== 预设报告（3 个，关联已完成的 T3/T8/T9）=====
DEMO_REPORTS = [
    {
        "id": "RPT-DEMO-001",
        "reviewTaskId": "RVT-DEMO-003",
        "reportNo": "QSZK-RPT-2026-001",
        "versionNo": 1,
        "snapshot": None,
        "status": "generated",
        "errorMsg": None,
        "createdAt": "2026-06-11T15:30:00.000Z",
    },
    {
        "id": "RPT-DEMO-002",
        "reviewTaskId": "RVT-DEMO-008",
        "reportNo": "QSZK-RPT-2026-002",
        "versionNo": 1,
        "snapshot": None,
        "status": "generated",
        "errorMsg": None,
        "createdAt": "2026-04-12T16:00:00.000Z",
    },
    {
        "id": "RPT-DEMO-003",
        "reviewTaskId": "RVT-DEMO-009",
        "reportNo": "QSZK-RPT-2026-003",
        "versionNo": 1,
        "snapshot": None,
        "status": "generated",
        "errorMsg": None,
        "createdAt": "2026-04-08T17:30:00.000Z",
    },
]


# ============================================================
# 工具函数
# ============================================================

def camel_to_snake(name: str) -> str:
    """camelCase → snake_case"""
    s1 = re.sub(r"(.)([A-Z][a-z]+)", r"\1_\2", name)
    return re.sub(r"([a-z0-9])([A-Z])", r"\1_\2", s1).lower()


def to_snake_row(data: dict) -> dict:
    """转换 dict 的所有 key 为 snake_case（与 data.py _to_db_row 一致）"""
    return {camel_to_snake(k): v for k, v in data.items()}


def calc_risk_positions(risks: list, paragraphs_by_id: dict = None) -> list:
    """计算 startPosition/endPosition（与 db.ts calcRiskPositions 一致）

    用段落原文 indexOf 计算，定位不到则保持 0。
    paragraphs_by_id 允许传入任务对应的样例段落索引；不传则使用默认 DEMO_PARAGRAPHS_BY_ID。
    """
    if paragraphs_by_id is None:
        paragraphs_by_id = DEMO_PARAGRAPHS_BY_ID
    result = []
    for r in risks:
        para = paragraphs_by_id.get(r["paragraphId"])
        if not para:
            result.append(r)
            continue
        text = para["text"]
        start = text.find(r["originalText"])
        if start < 0:
            result.append(r)
            continue
        r2 = dict(r)
        r2["startPosition"] = start
        r2["endPosition"] = start + len(r["originalText"])
        result.append(r2)
    return result


def build_risks_for_task(review_task_id: str, status_config, created_at: str, sample_id: str = None) -> list:
    """为任务生成风险项（复刻 seedData.ts buildRisksForTask）

    如果提供 sample_id 且存在于 SAMPLE_CONTRACTS，使用该样例的 riskTemplates；
    否则使用默认 DEMO_RISK_TEMPLATES。
    """
    if sample_id and sample_id in SAMPLE_CONTRACTS:
        templates = SAMPLE_CONTRACTS[sample_id]["riskTemplates"]
    else:
        templates = DEMO_RISK_TEMPLATES

    risks = []
    for idx, tpl in enumerate(templates):
        # 默认模板通过 snippetKey 引用原文；样例模板直接内置 originalText/paragraphId
        snippet = RISK_SNIPPETS.get(tpl.get("snippetKey"), {}) if "snippetKey" in tpl else {}
        original_text = tpl.get("originalText") or snippet.get("originalText", "")
        paragraph_id = tpl.get("paragraphId") or snippet.get("paragraphId", "")
        status = status_config(idx, tpl)
        risk = {
            "id": f"RISK-{review_task_id}-{str(idx + 1).zfill(3)}",
            "reviewTaskId": review_task_id,
            "title": tpl["title"],
            "riskType": tpl["riskType"],
            "riskLevel": tpl["riskLevel"],
            "clauseNumber": tpl["clauseNumber"],
            "clauseTitle": tpl["clauseTitle"],
            "originalText": original_text,
            "paragraphId": paragraph_id,
            "startPosition": 0,  # 运行时由 calc_risk_positions 计算
            "endPosition": 0,
            "riskReason": tpl["riskReason"],
            "reviewBasis": tpl["reviewBasis"],
            "suggestion": tpl["suggestion"],
            "editedSuggestion": None,
            "confidence": tpl["confidence"],
            "sourceType": tpl["sourceType"],
            "ruleId": tpl["ruleId"] if tpl.get("ruleId") else None,
            "status": status,
            "handler": None if status == "pending" else "李明",
            "handleComment": None,
            "ignoreReason": None,
            "version": 1,
            "createdAt": created_at,
            "updatedAt": created_at,
        }
        risks.append(risk)
    return risks


def build_risks_by_count(review_task_id: str, risk_count: dict,
                          status_config, created_at: str, sample_id: str = None) -> list:
    """根据 riskCount 配比选择 risks（用于补全缺失任务的 risks 数据）

    按 high/medium/low/notice 等级从对应模板中选择，仅使用本任务未使用过的模板，
    严格避免同一任务出现完全重复的风险。当某等级唯一模板数量不足时，按实际可用
    数量生成，不再循环复用重复模板。
    返回的 risks 数量可能小于 risk_count 总和（由 _sync_task_risk_count 修正任务统计）。

    如果提供 sample_id 且存在于 SAMPLE_CONTRACTS，使用该样例的 riskTemplates；
    否则使用默认 DEMO_RISK_TEMPLATES。
    """
    if sample_id and sample_id in SAMPLE_CONTRACTS:
        templates = SAMPLE_CONTRACTS[sample_id]["riskTemplates"]
    else:
        templates = DEMO_RISK_TEMPLATES

    def _template_key(tpl):
        snippet = RISK_SNIPPETS.get(tpl.get("snippetKey"), {}) if "snippetKey" in tpl else {}
        original_text = tpl.get("originalText") or snippet.get("originalText", "")
        return (tpl["title"], original_text)

    by_level = {"high": [], "medium": [], "low": [], "notice": []}
    for tpl in templates:
        by_level[tpl["riskLevel"]].append(tpl)

    selected = []
    task_used_keys = set()
    for level in ["high", "medium", "low", "notice"]:
        n = risk_count.get(level, 0)
        level_templates = by_level[level]
        if not level_templates:
            continue
        count = 0
        # 仅选择本任务未使用过的唯一模板，避免完全重复
        for tpl in level_templates:
            if count >= n:
                break
            key = _template_key(tpl)
            if key in task_used_keys:
                continue
            selected.append(tpl)
            task_used_keys.add(key)
            count += 1
        # 唯一模板耗尽后不再补充，确保同任务内无重复风险


    risks = []
    for seq, tpl in enumerate(selected):
        snippet = RISK_SNIPPETS.get(tpl.get("snippetKey"), {}) if "snippetKey" in tpl else {}
        original_text = tpl.get("originalText") or snippet.get("originalText", "")
        paragraph_id = tpl.get("paragraphId") or snippet.get("paragraphId", "")
        status = status_config(seq, tpl)
        risk = {
            "id": f"RISK-{review_task_id}-{str(seq + 1).zfill(3)}",
            "reviewTaskId": review_task_id,
            "title": tpl["title"],
            "riskType": tpl["riskType"],
            "riskLevel": tpl["riskLevel"],
            "clauseNumber": tpl["clauseNumber"],
            "clauseTitle": tpl["clauseTitle"],
            "originalText": original_text,
            "paragraphId": paragraph_id,
            "startPosition": 0,
            "endPosition": 0,
            "riskReason": tpl["riskReason"],
            "reviewBasis": tpl["reviewBasis"],
            "suggestion": tpl["suggestion"],
            "editedSuggestion": None,
            "confidence": tpl["confidence"],
            "sourceType": tpl["sourceType"],
            "ruleId": tpl["ruleId"] if tpl.get("ruleId") else None,
            "status": status,
            "handler": None if status == "pending" else "李明",
            "handleComment": None,
            "ignoreReason": None,
            "version": 1,
            "createdAt": created_at,
            "updatedAt": created_at,
        }
        risks.append(risk)
    return risks


def build_all_demo_risks() -> list:
    """生成所有演示任务的 risks

    为主演示任务 T1/T2/T3/T8/T9 复刻 seedData.ts 配置；
    为补全任务 T6/T7/T10/T11/T12/T14/T15 按 riskCount 配比生成 risks，
    确保 review_tasks 表的 riskCount 字段与 risks 表实际数据一致。

    有 sampleId 的任务使用对应样例合同的风险模板，无 sampleId 的任务使用默认模板。
    """
    risks = []
    task_sample_map = {t["id"]: t.get("sampleId") for t in DEMO_TASKS}

    # T1：18 个，全部 pending（主演示任务）
    risks.extend(build_risks_for_task(
        "RVT-DEMO-001", lambda idx, tpl: "pending", "2026-06-20T09:40:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-001"),
    ))
    # T2：8 个，混合已处理（已提交法务复核）
    t2_statuses = ["accepted", "edited", "accepted", "ignored",
                   "accepted", "edited", "accepted", "manual_review"]

    def t2_config(idx, tpl):
        if idx >= 8:
            return "ignored"
        return t2_statuses[idx] if idx < len(t2_statuses) else "accepted"

    risks.extend(build_risks_for_task(
        "RVT-DEMO-002", t2_config, "2026-06-18T11:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-002"),
    )[:8])
    # T3：7 个，全部 confirmed（已完成）
    risks.extend(build_risks_for_task(
        "RVT-DEMO-003", lambda idx, tpl: "confirmed", "2026-06-09T11:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-003"),
    )[:7])
    # T8：8 个，全部 confirmed
    risks.extend(build_risks_for_task(
        "RVT-DEMO-008", lambda idx, tpl: "confirmed", "2026-04-10T10:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-008"),
    )[:8])
    # T9：12 个，全部 confirmed
    risks.extend(build_risks_for_task(
        "RVT-DEMO-009", lambda idx, tpl: "confirmed", "2026-04-08T10:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-009"),
    )[:12])

    # ===== 补全缺失任务的 risks（按 riskCount 配比生成） =====
    # T6：12 个（5 high + 4 medium + 2 low + 1 notice），全部 pending
    risks.extend(build_risks_by_count(
        "RVT-DEMO-006",
        {"high": 5, "medium": 4, "low": 2, "notice": 1},
        lambda idx, tpl: "pending", "2026-05-15T14:20:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-006"),
    ))
    # T7：9 个（0 high + 6 medium + 3 low + 0 notice），全部 pending
    risks.extend(build_risks_by_count(
        "RVT-DEMO-007",
        {"high": 0, "medium": 6, "low": 3, "notice": 0},
        lambda idx, tpl: "pending", "2026-05-20T11:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-007"),
    ))
    # T10：4 个（0 high + 0 medium + 3 low + 1 notice），已提交法务复核
    t10_statuses = ["accepted", "ignored", "accepted", "confirmed"]

    def t10_config(idx, tpl):
        return t10_statuses[idx] if idx < len(t10_statuses) else "accepted"

    risks.extend(build_risks_by_count(
        "RVT-DEMO-010",
        {"high": 0, "medium": 0, "low": 3, "notice": 1},
        t10_config, "2026-04-25T16:30:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-010"),
    ))
    # T11：10 个（4 high + 5 medium + 1 low + 0 notice），全部 pending
    risks.extend(build_risks_by_count(
        "RVT-DEMO-011",
        {"high": 4, "medium": 5, "low": 1, "notice": 0},
        lambda idx, tpl: "pending", "2026-03-18T11:30:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-011"),
    ))
    # T12：8 个（2 high + 5 medium + 1 low + 0 notice），全部 confirmed
    risks.extend(build_risks_by_count(
        "RVT-DEMO-012",
        {"high": 2, "medium": 5, "low": 1, "notice": 0},
        lambda idx, tpl: "confirmed", "2026-03-10T17:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-012"),
    ))
    # T14：3 个（0 high + 0 medium + 2 low + 1 notice），全部 confirmed
    risks.extend(build_risks_by_count(
        "RVT-DEMO-014",
        {"high": 0, "medium": 0, "low": 2, "notice": 1},
        lambda idx, tpl: "confirmed", "2026-02-15T16:00:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-014"),
    ))
    # T15：2 个（0 high + 0 medium + 2 low + 0 notice），全部 pending
    risks.extend(build_risks_by_count(
        "RVT-DEMO-015",
        {"high": 0, "medium": 0, "low": 2, "notice": 0},
        lambda idx, tpl: "pending", "2026-01-12T10:30:00.000Z",
        sample_id=task_sample_map.get("RVT-DEMO-015"),
    ))
    return risks


def build_rule_versions() -> list:
    """为每条规则生成初始版本记录 v1（复刻 db.ts initDB 中 rule_versions 逻辑）"""
    versions = []
    for rule in DEMO_RULES:
        versions.append({
            "id": f"RV-{rule['id']}",  # 稳定 ID，保证幂等
            "ruleId": rule["id"],
            "version": rule["version"],
            "snapshot": dict(rule),  # 完整规则快照
            "changeNote": "初始版本",
            "operatorName": "系统",
            "createdAt": rule["updatedAt"],
        })
    return versions


def build_all_fields() -> list:
    """生成所有抽取字段（T1/T2/T3/T8/T9 各一套）

    复刻 db.ts initDB：T1 未确认（confirmed=false），其余已确认。
    """
    fields = []
    for tid in ["RVT-DEMO-001", "RVT-DEMO-002", "RVT-DEMO-003", "RVT-DEMO-008", "RVT-DEMO-009"]:
        confirmed = tid != "RVT-DEMO-001"
        for f in DEMO_EXTRACTED_FIELDS:
            fields.append({
                "id": f"EF-{tid}-{f['fieldKey']}",
                "reviewTaskId": tid,
                "fieldKey": f["fieldKey"],
                "fieldLabel": f["fieldLabel"],
                "fieldValue": f["fieldValue"],
                "confidence": f["confidence"],
                "confirmedValue": f["fieldValue"] if confirmed else None,
                "lowConfidence": f["lowConfidence"],
                "sourceText": f["sourceText"],
                "confirmed": confirmed,
            })
    return fields


def _delete_all(sb, table: str, id_col: str = "id"):
    """删除表内所有数据（supabase-py 需要一个 filter）

    用 neq 匹配所有非空 id（业务表 id 均为非空文本）。
    """
    sb.table(table).delete().neq(id_col, "").execute()


# ============================================================
# seed 函数
# ============================================================

def seed_users(sb):
    """写入 3 个演示用户（auth_uid 留空，登录时自动补写）"""
    _delete_all(sb, "users")
    rows = []
    for u in DEMO_USERS:
        row = to_snake_row(u)
        row["auth_uid"] = None  # 显式置空，登录时由 auth 流程补写
        rows.append(row)
    sb.table("users").insert(rows).execute()
    print(f"  ✓ users: 写入 {len(rows)} 条")


def seed_rules(sb):
    """写入 16 条规则 + 每条规则的初始版本记录（v1）"""
    # 先删子表 rule_versions，再删 rules（FK 级联也会处理，这里显式删除保证干净）
    _delete_all(sb, "rule_versions", "id")
    _delete_all(sb, "rules")

    rule_rows = [to_snake_row(r) for r in DEMO_RULES]
    sb.table("rules").insert(rule_rows).execute()

    version_rows = [to_snake_row(v) for v in build_rule_versions()]
    sb.table("rule_versions").insert(version_rows).execute()
    print(f"  ✓ rules: 写入 {len(rule_rows)} 条 + rule_versions: 写入 {len(version_rows)} 条")


def seed_tasks(sb):
    """写入 5 个演示任务"""
    # review_tasks 的删除会级联删除 risks/fields/reports，但为了 seed_risks 等函数可独立运行，
    # 这里只删 review_tasks，子表由各自 seed 函数负责清理
    _delete_all(sb, "review_tasks")
    rows = [to_snake_row(t) for t in DEMO_TASKS]
    sb.table("review_tasks").insert(rows).execute()
    print(f"  ✓ review_tasks: 写入 {len(rows)} 条")


def seed_risks(sb):
    """写入所有风险（T1 18 个 pending、T2 8 个混合状态、T3 7 个 confirmed）"""
    _delete_all(sb, "risks")
    raw_risks = build_all_demo_risks()

    # 按任务 sampleId 构建对应段落索引，确保风险原文定位准确
    task_sample_map = {t["id"]: t.get("sampleId") for t in DEMO_TASKS}
    paragraphs_by_task_id = {}
    for tid, sample_id in task_sample_map.items():
        if sample_id and sample_id in SAMPLE_CONTRACTS:
            paragraphs = SAMPLE_CONTRACTS[sample_id]["paragraphs"]
        else:
            paragraphs = DEMO_PARAGRAPHS
        paragraphs_by_task_id[tid] = {p["id"]: p for p in paragraphs}

    # 按任务分组计算风险位置
    risks_by_task = {}
    for r in raw_risks:
        risks_by_task.setdefault(r["reviewTaskId"], []).append(r)

    all_risks = []
    for tid, task_risks in risks_by_task.items():
        all_risks.extend(calc_risk_positions(task_risks, paragraphs_by_task_id.get(tid, DEMO_PARAGRAPHS_BY_ID)))

    rows = [to_snake_row(r) for r in all_risks]
    if rows:
        sb.table("risks").insert(rows).execute()
    print(f"  ✓ risks: 写入 {len(rows)} 条")

    # 根据 risks 表实际数据，同步更新 review_tasks 的 risk_count 字段
    # 确保 task 列表显示的 riskCount 与 risks 表完全一致
    _sync_task_risk_count(sb, all_risks)


def _sync_task_risk_count(sb, all_risks: list):
    """根据实际 risks 数据，更新 review_tasks 的 risk_count 和 risk_level_max 字段

    避免出现 task.riskCount 与 risks 表等级分布不一致的情况
    （例如 T2 原预设 3 high，但实际 risks 切片后是 4 high）。
    """
    count_map = {}  # {task_id: {high:N, medium:N, low:N, notice:N}}
    for r in all_risks:
        tid = r.get("reviewTaskId")
        lv = r.get("riskLevel", "low")
        if tid not in count_map:
            count_map[tid] = {"high": 0, "medium": 0, "low": 0, "notice": 0}
        if lv in count_map[tid]:
            count_map[tid][lv] += 1

    # 没有 risks 的任务（draft/failed），risk_count 清零
    for t in DEMO_TASKS:
        if t["id"] not in count_map:
            count_map[t["id"]] = {"high": 0, "medium": 0, "low": 0, "notice": 0}

    updated = 0
    for tid, rc in count_map.items():
        # 根据 risk_count 推导 risk_level_max
        if rc["high"] > 0:
            level_max = "high"
        elif rc["medium"] > 0:
            level_max = "medium"
        elif rc["low"] > 0:
            level_max = "low"
        elif rc["notice"] > 0:
            level_max = "notice"
        else:
            level_max = None  # draft/failed 任务无风险

        try:
            sb.table("review_tasks").update({
                "risk_count": json.dumps(rc),
                "risk_level_max": level_max,
            }).eq("id", tid).execute()
            updated += 1
        except Exception as e:
            print(f"  ⚠ 同步 task {tid} risk_count 失败：{e}")
    if updated:
        print(f"  ✓ review_tasks.risk_count + risk_level_max: 同步 {updated} 条")


def seed_fields(sb):
    """写入抽取字段（T1/T2/T3 各一套 DEMO_EXTRACTED_FIELDS）"""
    _delete_all(sb, "extracted_fields")
    all_fields = build_all_fields()
    rows = [to_snake_row(f) for f in all_fields]
    if rows:
        sb.table("extracted_fields").insert(rows).execute()
    print(f"  ✓ extracted_fields: 写入 {len(rows)} 条")


def _to_camel_case(row: dict) -> dict:
    """snake_case 行 → camelCase（含嵌套对象）"""
    result = {}
    for k, v in row.items():
        if k.startswith("_"):
            continue
        parts = k.split("_")
        camel = parts[0] + "".join(p.capitalize() for p in parts[1:])
        # 字段/风险中有 text->reviewBasis, text->riskReason 等嵌套字段
        # 保持简单：仅扁平映射
        result[camel] = v
    return result


def _build_report_snapshot(sb, task_id: str, report_row: dict) -> dict:
    """从 DB 查询任务+风险+字段，构建报告 snapshot"""
    raw = sb.table("review_tasks").select("*").eq("id", task_id).single().execute().data
    if not raw:
        return report_row

    risks_raw = sb.table("risks").select("*").eq("review_task_id", task_id).execute().data or []
    fields_raw = sb.table("extracted_fields").select("*").eq("review_task_id", task_id).execute().data or []

    task = _to_camel_case(raw)
    risks = [_to_camel_case(r) for r in risks_raw]
    fields = [_to_camel_case(f) for f in fields_raw]

    risk_count = {"high": 0, "medium": 0, "low": 0, "notice": 0}
    for r in risks:
        lv = r.get("riskLevel", "low")
        if lv in risk_count:
            risk_count[lv] += 1

    weights = {"high": 25, "medium": 12, "low": 4, "notice": 1}
    ws = sum(weights.get(r.get("riskLevel", "low"), 0) for r in risks)
    risk_score = round((ws / (ws + 50)) * 100) if ws > 0 else 0

    overall = "high" if risk_count["high"] > 0 else "medium" if risk_count["medium"] > 0 else "low" if risk_count["low"] > 0 else "notice"
    high_c, medium_c, low_c, notice_c = risk_count["high"], risk_count["medium"], risk_count["low"], risk_count["notice"]
    ai_summary = f"本次共识别 {len(risks)} 项风险，其中高风险 {high_c} 项、中风险 {medium_c} 项、低风险 {low_c} 项、提示项 {notice_c} 项。"

    snapshot = {
        "contractName": task.get("contractName", ""),
        "contractNo": task.get("contractNo", ""),
        "counterparty": task.get("counterparty", ""),
        "amount": task.get("amount", 0),
        "currency": task.get("currency", "CNY"),
        "contractType": task.get("contractType", ""),
        "reviewFocus": task.get("reviewFocus", []),
        "fields": fields,
        "risks": risks,
        "riskCount": risk_count,
        "riskScore": risk_score,
        "overallRiskLevel": overall,
        "aiSummary": ai_summary,
        "legalOpinion": task.get("legalOpinion", ""),
        "legalConclusion": task.get("legalConclusion", "sign_after_modify"),
        "majorRisks": [r for r in risks if r.get("riskLevel") == "high" and r.get("status") in ("confirmed", "accepted", "edited", "manual_review")],
        "disclaimer": "本系统审核结果由AI辅助生成，仅供合同初审参考，不构成正式法律意见，最终结论应由专业人员确认。",
        "generatedAt": task.get("completedAt") or task.get("updatedAt") or report_row.get("createdAt"),
    }

    report_row["snapshot"] = snapshot
    return report_row


def seed_reports(sb):
    """写入 3 个演示报告，每个报告从 DB 动态查询 snapshot"""
    _delete_all(sb, "reports")
    total_risks = 0
    total_fields = 0
    rows = []
    for r in DEMO_REPORTS:
        rpt = _build_report_snapshot(sb, r["reviewTaskId"], dict(r))
        snap = rpt.get("snapshot", {})
        rows.append(to_snake_row(rpt))
        total_risks += len(snap.get("risks", []))
        total_fields += len(snap.get("fields", []))
    if rows:
        sb.table("reports").insert(rows).execute()
    print(f"  ✓ reports: 写入 {len(rows)} 条（snapshot 含 {total_risks} 项风险 / {total_fields} 个字段）")


def _create_demo_docx(task_id: str, file_dir: Path,
                       contract_name: str = "软件系统采购合同",
                       paragraphs: list = None) -> Optional[str]:
    """用 python-docx 创建演示合同 DOCX 文件，返回路径"""
    if not HAS_DOCX:
        return None
    if paragraphs is None:
        paragraphs = DEMO_PARAGRAPHS
    try:
        file_dir.mkdir(parents=True, exist_ok=True)
        # 本地文件名使用对应样例合同名称（可含中文），便于区分
        docx_path = file_dir / f"{contract_name}.docx"
        document = DocxDocument()
        # 设置默认字体
        style = document.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = 112800  # 12pt
        for p in paragraphs:
            para_type = p.get("type", "body")
            docx_para = document.add_paragraph()
            if para_type == "title":
                docx_para.alignment = 1  # CENTER
                run = docx_para.add_run(p["text"])
                run.bold = True
                run.font.size = 169000  # 18pt
            elif para_type == "header":
                run = docx_para.add_run(p["text"])
                run.font.size = 112800  # 12pt
                docx_para.paragraph_format.space_before = 0
            elif para_type == "signature":
                run = docx_para.add_run(p["text"])
                run.font.size = 112800
                docx_para.paragraph_format.space_before = 240  # 12pt
            else:
                # 正文：带条款编号加粗
                text = p["text"]
                clause_no = p.get("clauseNo", "")
                if clause_no and text.startswith(clause_no):
                    first_line_end = text.find("\n")
                    first_line = text[:first_line_end] if first_line_end > 0 else text
                    rest = text[first_line_end:] if first_line_end > 0 else ""
                    run = docx_para.add_run(first_line)
                    run.bold = True
                    run.font.size = 112800
                    if rest:
                        run2 = docx_para.add_run(rest)
                        run2.font.size = 112800
                else:
                    run = docx_para.add_run(text)
                    font.size = 112800
                docx_para.paragraph_format.line_spacing = 1.5
        document.save(str(docx_path))
        return str(docx_path)
    except Exception as e:
        print(f"  ⚠ 创建演示 DOCX 失败：{e}")
        return None


def _create_demo_pdf(task_id: str, file_dir: Path,
                       contract_name: str = "软件系统采购合同",
                       paragraphs: list = None) -> Optional[str]:
    """用 reportlab 创建演示合同 PDF 文件，返回路径"""
    if not HAS_REPORTLAB:
        return None
    if paragraphs is None:
        paragraphs = DEMO_PARAGRAPHS
    try:
        file_dir.mkdir(parents=True, exist_ok=True)
        # 本地文件名使用对应样例合同名称（可含中文），便于区分
        pdf_path = file_dir / f"{contract_name}.pdf"

        # 注册中文字体
        cn_font = "Helvetica"
        for fp in [r"C:\Windows\Fonts\msyh.ttc", r"C:\Windows\Fonts\simhei.ttf",
                   r"C:\Windows\Fonts\msyh.ttf", r"C:\Windows\Fonts\Deng.ttf"]:
            if os.path.exists(fp):
                try:
                    pdfmetrics.registerFont(TTFont("CNDoc", fp))
                    cn_font = "CNDoc"
                    break
                except Exception:
                    continue

        styles = getSampleStyleSheet()
        style_title = ParagraphStyle("CN_Title", parent=styles["Title"],
                                     fontName=cn_font, fontSize=18, leading=24, spaceAfter=16, alignment=1)
        style_h1 = ParagraphStyle("CN_H1", parent=styles["Heading1"],
                                  fontName=cn_font, fontSize=13, leading=18, spaceAfter=8, spaceBefore=12)
        style_normal = ParagraphStyle("CN_Normal", parent=styles["Normal"],
                                      fontName=cn_font, fontSize=10.5, leading=17, spaceAfter=8)

        doc = SimpleDocTemplate(
            str(pdf_path), pagesize=A4, title=contract_name, author="契审智控",
            leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=16 * mm,
        )
        story = []
        for p in paragraphs:
            para_type = p.get("type", "body")
            text = p["text"]
            if para_type == "title":
                story.append(Paragraph(text, style_title))
            elif para_type in ("header", "signature"):
                story.append(Paragraph(text.replace("\n", "<br/>"), style_normal))
                story.append(Spacer(1, 4 * mm))
            else:
                clause_no = p.get("clauseNo", "")
                if clause_no and text.startswith(clause_no):
                    first_line_end = text.find("\n")
                    first_line = text[:first_line_end] if first_line_end > 0 else text
                    rest = text[first_line_end:] if first_line_end > 0 else ""
                    story.append(Paragraph(f"<b>{first_line}</b>{rest.replace(chr(10), '<br/>')}", style_h1))
                else:
                    story.append(Paragraph(text.replace("\n", "<br/>"), style_normal))
        doc.build(story)
        return str(pdf_path)
    except Exception as e:
        print(f"  ⚠ 创建演示 PDF 失败：{e}")
        return None


def seed_documents(sb):
    """为所有演示任务创建 parsed_documents（含 full_text、paragraphs、html_content）

    按任务 sampleId 使用对应样例合同的段落与标题；无 sampleId 的任务使用默认 DEMO_PARAGRAPHS。
    如果有 python-docx + mammoth，会为每个样例创建真实 DOCX 并转换 HTML；
    否则回退到 _text_to_html 基础转换。
    """
    _delete_all(sb, "parsed_documents", "review_task_id")

    # 从段落自动生成章节（与前端 buildSectionsFromParagraphs / 后端 pdf_service 逻辑一致）
    def _build_demo_sections(paragraphs):
        sections = []
        current_paras = []
        current_title = ""
        current_no = ""
        prelude_paras = []

        def flush():
            nonlocal current_paras, current_title, current_no
            if current_paras:
                sections.append({
                    "id": f"sec-{len(sections) + 1}",
                    "title": current_title or "正文",
                    # 标题章节显式设置 current_no=''，需保持空字符串，避免显示「正文」等虚假章节号
                    "clauseNo": "" if current_no == "" else (current_no or "正文"),
                    "paragraphIds": current_paras[:],
                })
                current_paras = []

        for p in paragraphs:
            ptype = p.get("type", "body")
            clause_no = p.get("clauseNo")
            clause_title = p.get("clauseTitle")
            if ptype == "title":
                flush()
                current_title = p["text"][:30]
                # 合同标题不显示「标题」等虚假章节号，避免左栏目录误导
                current_no = ""
                current_paras.extend(prelude_paras)
                prelude_paras = []
            elif ptype in ("header", "signature"):
                # 甲乙方信息、签署落款不单独建章节；若尚无真实章节，先作为前导暂存
                if not current_no:
                    prelude_paras.append(p["id"])
                else:
                    current_paras.append(p["id"])
                continue
            else:
                if clause_no and clause_no != current_no:
                    flush()
                    current_no = clause_no
                    current_title = clause_title or clause_no
                    current_paras.extend(prelude_paras)
                    prelude_paras = []
            current_paras.append(p["id"])

        flush()

        # 若全文都没有标题/条款，只有前导段落，统一归入「正文」一节
        # 避免生成「合同信息」「首部」「签署落款」等误导性目录项
        if prelude_paras and not sections:
            sections.append({
                "id": "sec-1",
                "title": "正文",
                "clauseNo": "正文",
                "paragraphIds": prelude_paras[:],
            })

        return sections

    # 构建 paragraphs JSON（尊重显式 type，不再按 index 硬编码）
    def _build_para_list(paragraphs):
        para_list = []
        for p in paragraphs:
            entry = {"id": p["id"], "index": p["index"], "text": p["text"], "type": p.get("type", "body")}
            if p.get("clauseNo"):
                entry["clauseNo"] = p["clauseNo"]
            if p.get("clauseTitle"):
                entry["clauseTitle"] = p["clauseTitle"]
            para_list.append(entry)
        return para_list

    # 按 sampleId 分组任务；sampleId 为 None 的任务共用默认段落
    task_sample_map = {t["id"]: t.get("sampleId") for t in DEMO_TASKS}
    task_file_name_map = {t["id"]: t["fileName"] for t in DEMO_TASKS}
    sample_tasks = {}
    for tid, sample_id in task_sample_map.items():
        sample_tasks.setdefault(sample_id, []).append(tid)

    # 为每个样例预生成段落、全文、HTML、DOCX、PDF
    sample_cache = {}
    can_generate_real_html = HAS_MAMMOTH and HAS_DOCX
    for sample_id, tids in sample_tasks.items():
        if sample_id and sample_id in SAMPLE_CONTRACTS:
            contract_name = SAMPLE_CONTRACTS[sample_id]["contractName"]
            paragraphs = SAMPLE_CONTRACTS[sample_id]["paragraphs"]
        else:
            contract_name = "软件系统采购合同"
            paragraphs = DEMO_PARAGRAPHS

        para_list = _build_para_list(paragraphs)
        full_text = "\n\n".join(p["text"] for p in paragraphs)
        html = None
        docx_path = None
        pdf_path = None

        if can_generate_real_html:
            sample_dir = Path("/tmp") / "contract_files" / f"sample_{sample_id or 'default'}"
            docx_path = _create_demo_docx(
                f"sample_{sample_id or 'default'}", sample_dir,
                contract_name=contract_name, paragraphs=paragraphs,
            )
            if docx_path and Path(docx_path).exists():
                try:
                    with open(docx_path, "rb") as f:
                        result = mammoth.convert_to_html(f)
                    if result.value.strip():
                        css = """<style>
  body { font-family: 'Microsoft YaHei','SimSun',serif; line-height:1.8; padding:40px; color:#333; max-width:900px; margin:0 auto; }
  table { border-collapse:collapse; width:100%; margin:16px 0; }
  td,th { border:1px solid #999; padding:8px 12px; text-align:left; }
  th { background:#f5f5f5; font-weight:600; }
  p { margin:8px 0; }
  h3 { margin:16px 0 8px; color:#1a1a1a; }
  img { max-width:100%; }
</style>"""
                        html = css + result.value
                    else:
                        html = _text_to_html(full_text)
                except Exception as e:
                    print(f"  ⚠ 样例 {sample_id or 'default'} mammoth 转换失败，回退到 _text_to_html: {e}")
                    html = _text_to_html(full_text)
            else:
                html = _text_to_html(full_text)
        else:
            html = _text_to_html(full_text)

        # 预生成 PDF（供 fileName 为 .pdf 的任务使用）
        sample_pdf_dir = Path("/tmp") / "contract_files" / f"sample_{sample_id or 'default'}_pdf"
        pdf_path = _create_demo_pdf(
            f"sample_{sample_id or 'default'}", sample_pdf_dir,
            contract_name=contract_name, paragraphs=paragraphs,
        )

        sample_cache[sample_id] = {
            "contractName": contract_name,
            "paragraphs": paragraphs,
            "paraList": para_list,
            "fullText": full_text,
            "html": html,
            "docxPath": docx_path,
            "pdfPath": pdf_path,
        }

    if can_generate_real_html:
        print("  ✓ 已用 mammoth 从真实 DOCX 生成 html_content")
    else:
        print("  ⚠ 未使用 mammoth（python-docx 或 mammoth 不可用），使用 _text_to_html 基础转换")

    count = 0
    html_column_ok = True
    for tid in [t["id"] for t in DEMO_TASKS]:
        sample_id = task_sample_map.get(tid)
        sample = sample_cache.get(sample_id)
        if not sample:
            # 理论上 sample_cache 已包含所有 sample_id；兜底使用默认
            sample = sample_cache.get(None)

        row = {
            "review_task_id": tid,
            "title": sample["contractName"],
            "sections": json.dumps(_build_demo_sections(sample["paraList"])),
            "paragraphs": json.dumps(sample["paraList"]),
            "full_text": sample["fullText"],
            "html_content": sample["html"],
        }
        try:
            sb.table("parsed_documents").upsert(row).execute()
            count += 1
        except Exception as e:
            err_msg = str(e)
            if "html_content" in err_msg and ("could not find" in err_msg.lower() or "PGRST204" in err_msg):
                if html_column_ok:
                    html_column_ok = False
                    print("\n" + "=" * 64)
                    print("⚠ 警告：数据库 parsed_documents 表缺少 html_content 列！")
                    print("  已自动降级为不带 html_content 写入（原文格式预览将不可用）。")
                    print("  修复方法：在 Supabase Dashboard → SQL Editor 中执行以下 SQL：")
                    print("-" * 64)
                    print("ALTER TABLE public.parsed_documents ADD COLUMN IF NOT EXISTS html_content TEXT;")
                    print("NOTIFY pgrst, 'reload schema';")
                    print("-" * 64)
                    print("  执行后重新运行 seed 即可启用原文格式预览。")
                    print("=" * 64 + "\n")
                row.pop("html_content", None)
                sb.table("parsed_documents").upsert(row).execute()
                count += 1
            else:
                raise

        # 为每个任务复制源文件到 /tmp/contract_files/{tid}/（本地开发环境）
        # 同时上传到 Supabase Storage，避免 Render /tmp 重启丢失
        # 根据任务 fileName 扩展名决定上传 original.docx 还是 original.pdf
        # 下载端点会根据 fileName 扩展名计算 storage key，二者必须匹配
        task_file_name = task_file_name_map.get(tid, "软件系统采购合同.docx")
        is_pdf_task = task_file_name.lower().endswith(".pdf")
        task_dir = Path("/tmp") / "contract_files" / tid
        task_dir.mkdir(parents=True, exist_ok=True)

        if is_pdf_task:
            pdf_path = sample.get("pdfPath")
            if pdf_path and Path(pdf_path).exists():
                # 本地保留对应样例名称的 PDF（可含中文）
                local_pdf = task_dir / f"{sample['contractName']}.pdf"
                if not local_pdf.exists():
                    import shutil
                    shutil.copy2(pdf_path, str(local_pdf))
                try:
                    _upload_demo_pdf_to_storage(sb, pdf_path, tid)
                except Exception as e:
                    print(f"  ⚠ 演示任务 {tid} 上传 PDF 到 Storage 失败：{e}")
        else:
            docx_path = sample.get("docxPath")
            if docx_path and Path(docx_path).exists():
                # 本地保留对应样例名称的 DOCX（可含中文）
                local_docx = task_dir / f"{sample['contractName']}.docx"
                if not local_docx.exists():
                    import shutil
                    shutil.copy2(docx_path, str(local_docx))
                try:
                    _upload_demo_docx_to_storage(sb, docx_path, tid)
                except Exception as e:
                    print(f"  ⚠ 演示任务 {tid} 上传 DOCX 到 Storage 失败：{e}")

    print(f"  ✓ parsed_documents: 写入 {count} 条" + ("（含 html_content）" if html_column_ok else "（未含 html_content，见上方警告）"))

    # 回读校验：仅在 html_column_ok 时确认数据确实写入
    if html_column_ok:
        try:
            check = sb.table("parsed_documents").select("review_task_id, html_content").limit(1).execute()
            if check.data:
                sample = check.data[0]
                if not sample.get("html_content"):
                    print("  ⚠ 警告：html_content 列存在但值为空，请检查 _text_to_html 函数")
                else:
                    print("  ✓ html_content 列存在且数据正常")
        except Exception as e:
            print(f"  ⚠ html_content 回读校验失败：{e}")


def _text_to_html(text: str) -> str:
    """将纯文本转为基础 HTML 段落（保留换行和简单格式）"""
    paragraphs = []
    for block in text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        lines = block.split("\n")
        if len(lines) >= 3 and all(len(l) <= 50 for l in lines):
            # 可能是表格型数据
            rows_html = "".join(f"<tr><td>{'</td><td>'.join(l.split())}</td></tr>" for l in lines if l.strip())
            if rows_html:
                paragraphs.append(f"<table>{rows_html}</table>")
                continue
        if len(lines) == 1 and len(lines[0]) <= 40:
            paragraphs.append(f"<h3>{lines[0]}</h3>")
        else:
            paragraphs.append("<p>" + "<br/>".join(lines) + "</p>")
    css = """
    <style>
      body { font-family: 'Microsoft YaHei','SimSun',serif; line-height:1.8; padding:40px; color:#333; max-width:900px; margin:0 auto; }
      table { border-collapse:collapse; width:100%; margin:16px 0; }
      td,th { border:1px solid #999; padding:8px 12px; text-align:left; }
      th { background:#f5f5f5; font-weight:600; }
      p { margin:8px 0; }
      h3 { margin:16px 0 8px; color:#1a1a1a; }
    </style>
    """
    return css + "\n".join(paragraphs)


def _upload_demo_docx_to_storage(sb, docx_path: str, task_id: str):
    """把演示 DOCX 上传到 Supabase Storage，保证 Render 重启后仍可下载"""
    bucket = "contract-files"
    filename = "软件系统采购合同.docx"
    # Supabase Storage key 不能含中文或 % 编码，固定用 original.docx
    file_path = f"{task_id}/original.docx"

    # 先尝试创建 bucket（已存在则忽略）
    try:
        sb.storage.create_bucket(bucket, options={"public": False})
    except Exception as e:
        err_msg = str(e)
        if "already exists" not in err_msg.lower() and "duplicate" not in err_msg.lower():
            print(f"  ⚠ 创建 Storage bucket 失败（可能已存在）: {e}")

    content = Path(docx_path).read_bytes()
    try:
        sb.storage.from_(bucket).upload(
            path=file_path,
            file=content,
            file_options={
                "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "upsert": "true",
            },
        )
    except Exception as e:
        err_msg = str(e)
        # 已存在时可能抛错，尝试覆盖
        if "already exists" in err_msg.lower() or "duplicate" in err_msg.lower():
            try:
                sb.storage.from_(bucket).remove([file_path])
                sb.storage.from_(bucket).upload(
                    path=file_path,
                    file=content,
                    file_options={
                        "content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    },
                )
            except Exception as e2:
                raise RuntimeError(f"覆盖 Storage 文件失败：{e2}")
        else:
            raise


def _upload_demo_pdf_to_storage(sb, pdf_path: str, task_id: str):
    """把演示 PDF 上传到 Supabase Storage，key 为 {task_id}/original.pdf

    用于 fileName 为 .pdf 的演示任务，下载端点会根据 fileName 扩展名计算 storage key。
    """
    bucket = "contract-files"
    file_path = f"{task_id}/original.pdf"

    try:
        sb.storage.create_bucket(bucket, options={"public": False})
    except Exception as e:
        err_msg = str(e)
        if "already exists" not in err_msg.lower() and "duplicate" not in err_msg.lower():
            print(f"  ⚠ 创建 Storage bucket 失败（可能已存在）: {e}")

    content = Path(pdf_path).read_bytes()
    try:
        sb.storage.from_(bucket).upload(
            path=file_path,
            file=content,
            file_options={
                "content-type": "application/pdf",
                "upsert": "true",
            },
        )
    except Exception as e:
        err_msg = str(e)
        if "already exists" in err_msg.lower() or "duplicate" in err_msg.lower():
            try:
                sb.storage.from_(bucket).remove([file_path])
                sb.storage.from_(bucket).upload(
                    path=file_path,
                    file=content,
                    file_options={"content-type": "application/pdf"},
                )
            except Exception as e2:
                raise RuntimeError(f"覆盖 Storage PDF 文件失败：{e2}")
        else:
            raise


def seed_audit_logs(sb):
    """写入演示审计日志（操作时间轴数据）"""
    _delete_all(sb, "audit_logs")
    rows = [to_snake_row(log) for log in DEMO_AUDIT_LOGS]
    if rows:
        sb.table("audit_logs").insert(rows).execute()
    print(f"  ✓ audit_logs: 写入 {len(rows)} 条")


def _check_permissions(sb):
    """检查 service_role 是否有表权限。

    若 schema.sql 未执行授权段（GRANT），service_role 会无任何权限，
    此处提前检测并打印修复指引，避免逐表报错。
    """
    GRANT_SQL = """-- 在 Supabase Dashboard → SQL Editor 中执行以下语句后重新运行 seed.py
GRANT ALL ON ALL TABLES IN SCHEMA public TO service_role;
GRANT USAGE, SELECT ON ALL SEQUENCES IN SCHEMA public TO service_role;
GRANT SELECT, INSERT, UPDATE, DELETE ON ALL TABLES IN SCHEMA public TO authenticated;
GRANT SELECT, INSERT, UPDATE, DELETE ON ALL TABLES IN SCHEMA public TO anon;"""
    try:
        sb.table("users").select("id").limit(1).execute()
        return True
    except Exception as e:
        msg = str(e)
        if "42501" in msg or "permission denied" in msg:
            print("\n" + "=" * 64)
            print("ERROR: service_role 缺少表权限（permission denied）")
            print("原因：schema.sql 的授权段未执行，或表创建后未授予 service_role 权限。")
            print("修复方法：在 Supabase Dashboard → SQL Editor 中执行以下 SQL：")
            print("-" * 64)
            print(GRANT_SQL)
            print("-" * 64)
            print("执行完成后，重新运行：python supabase/seed.py")
            print("=" * 64 + "\n")
            return False
        # 其他错误（如表不存在）也在此提示
        print(f"\nERROR: 数据库检查失败 - {msg[:200]}")
        print("请确认 schema.sql 已在 Supabase Dashboard 中执行。")
        return False


def seed_all():
    """调用以上所有函数，按依赖顺序写入种子数据"""
    if not SUPABASE_URL or not SUPABASE_SERVICE_ROLE_KEY:
        print("ERROR: SUPABASE_URL 或 SUPABASE_SERVICE_ROLE_KEY 未配置")
        print("请在 backend/.env 中设置后重试")
        sys.exit(1)

    print(f"连接 Supabase: {SUPABASE_URL}")
    sb = create_client(SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY)

    # 前置权限检查
    if not _check_permissions(sb):
        sys.exit(1)

    print("开始写入种子数据（幂等：先清空后插入）...")
    # 删除顺序：子表 → 父表（避免 FK 冲突，虽然 ON DELETE CASCADE 会处理，这里显式控制）
    # 插入顺序：父表 → 子表（保证 FK 引用存在）
    seed_users(sb)
    seed_rules(sb)
    seed_tasks(sb)
    seed_risks(sb)
    seed_fields(sb)
    seed_reports(sb)
    seed_audit_logs(sb)
    seed_documents(sb)
    print("种子数据写入完成！")

    # 统计校验
    print("\n数据统计：")
    for table in ["users", "rules", "rule_versions", "review_tasks",
                  "risks", "extracted_fields", "reports", "audit_logs"]:
        resp = sb.table(table).select("id", count="exact").execute()
        print(f"  {table}: {resp.count} 条")


if __name__ == "__main__":
    seed_all()
