"""
生成 4 份真实合同样例（DOCX + PDF 各一份）
每份合同包含：图片、表格、图表等非文字内容
用于导入测试合同审核系统

输出目录：backend/test_contracts/
生成文件：
  01_软件采购合同_含图表.docx / .pdf
  02_硬件采购合同_含图表.docx / .pdf
  03_运维服务合同_含图表.docx / .pdf
  04_系统集成合同_含图表.docx / .pdf

用法：
    cd backend
    .\\venv\\Scripts\\python.exe scripts\\generate_sample_contracts_v2.py
"""
import os
import math
from PIL import Image, ImageDraw, ImageFont

# reportlab / python-docx 在函数内按需导入

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.normpath(os.path.join(SCRIPT_DIR, ".."))
OUTPUT_DIR = os.path.join(BACKEND_DIR, "test_contracts")

CN_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\msyhbd.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\msyh.ttf",
    r"C:\Windows\Fonts\Deng.ttf",
    r"C:\Windows\Fonts\SIMKAI.ttf",
]

PARTY_A = "智远科技有限公司"


def _find_cn_font_path():
    """返回第一个可用的中文字体路径，找不到返回 None。"""
    for fp in CN_FONT_CANDIDATES:
        if os.path.exists(fp):
            return fp
    return None


CN_FONT_PATH = _find_cn_font_path()


def _get_font(size):
    """获取 Pillow 中文字体。"""
    if CN_FONT_PATH:
        try:
            return ImageFont.truetype(CN_FONT_PATH, size)
        except Exception:
            pass
    return ImageFont.load_default()


# =====================================================================
# 一、占位图片与图表生成（Pillow）
# =====================================================================

def gen_logo_image(path):
    """公司 logo：蓝色矩形背景 + 白色文字「智远科技」。"""
    w, h = 240, 72
    img = Image.new("RGB", (w, h), color=(22, 119, 255))
    draw = ImageDraw.Draw(img)
    font = _get_font(32)
    text = "智远科技"
    bbox = draw.textbbox((0, 0), text, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((w - tw) // 2, (h - th) // 2 - bbox[1]), text, fill=(255, 255, 255), font=font)
    img.save(path)
    print(f"  ✓ logo 图片：{path}")


def gen_stamp_image(path):
    """签署盖章图：红色圆形印章。"""
    size = 140
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    red = (245, 34, 45, 255)
    draw.ellipse((6, 6, size - 6, size - 6), outline=red, width=4)

    font = _get_font(20)
    cx, cy = size // 2, size // 2 - 24
    r = 10
    star_pts = []
    for i in range(10):
        ang = -math.pi / 2 + i * math.pi / 5
        rr = r if i % 2 == 0 else r * 0.4
        star_pts.append((cx + rr * math.cos(ang), cy + rr * math.sin(ang)))
    draw.polygon(star_pts, fill=red)

    lines = ["合同", "专用章"]
    try:
        lh = draw.textbbox((0, 0), "合同", font=font)[3] - draw.textbbox((0, 0), "合同", font=font)[1]
    except Exception:
        lh = 20
    y0 = cy + 14
    for idx, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        lw = bbox[2] - bbox[0]
        x = (size - lw) // 2
        y = y0 + idx * (lh + 4)
        draw.text((x, y), line, fill=red, font=font)
    img.save(path)
    print(f"  ✓ 印章图片：{path}")


def gen_gantt_chart(path):
    """项目进度甘特图（ERP 项目）。"""
    tasks = ["需求分析", "系统设计", "开发实现", "测试验收"]
    starts = [0, 15, 35, 75]
    widths = [15, 20, 40, 20]
    colors = ["#1677ff", "#13c2c2", "#52c41a", "#fa8c16"]

    w, h = 640, 260
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = _get_font(14)
    title_font = _get_font(18)
    draw.text((20, 14), "图1  ERP项目进度甘特图", fill=(33, 33, 33), font=title_font)

    left, top, right, bottom = 110, 55, 610, 210
    draw.rectangle((left, top, right, bottom), outline="#d9d9d9", width=1)

    # 网格线（每 10 天一条）
    for day in range(0, 101, 10):
        x = left + day * (right - left) / 100
        draw.line((x, top, x, bottom), fill="#f0f0f0", width=1)
        draw.text((x - 8, bottom + 6), str(day), fill="#595959", font=font)

    row_h = (bottom - top) / len(tasks)
    for i, task in enumerate(tasks):
        y0 = top + i * row_h + 8
        y1 = top + (i + 1) * row_h - 8
        draw.text((10, y0 + 6), task, fill="#262626", font=font)
        sx = left + starts[i] * (right - left) / 100
        ex = left + (starts[i] + widths[i]) * (right - left) / 100
        draw.rounded_rectangle((sx, y0, ex, y1), radius=6, fill=colors[i])
        txt = f"{widths[i]}天"
        bbox = draw.textbbox((0, 0), txt, font=font)
        tw = bbox[2] - bbox[0]
        draw.text(((sx + ex - tw) / 2, y0 + 6), txt, fill=(255, 255, 255), font=font)

    draw.text((left, bottom + 28), "单位：天", fill="#8c8c8c", font=font)
    img.save(path)
    print(f"  ✓ 甘特图：{path}")


def gen_pie_chart(path):
    """设备金额分布饼图。"""
    labels = ["服务器", "存储设备", "网络设备", "配件及服务"]
    values = [60, 20, 12, 8]
    colors = ["#1677ff", "#13c2c2", "#52c41a", "#faad14"]

    w, h = 560, 360
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _get_font(18)
    font = _get_font(14)
    draw.text((20, 14), "图1  设备采购金额分布", fill=(33, 33, 33), font=title_font)

    cx, cy, r = 180, 190, 110
    total = sum(values)
    start = -90
    for label, value, color in zip(labels, values, colors):
        angle = value / total * 360
        draw.pieslice((cx - r, cy - r, cx + r, cy + r), start, start + angle, fill=color)
        mid = math.radians(start + angle / 2)
        lx = cx + (r + 30) * math.cos(mid)
        ly = cy + (r + 30) * math.sin(mid)
        draw.text((lx, ly - 7), f"{label}\n{value}%", fill="#262626", font=font)
        start += angle

    # 图例
    for i, (label, color) in enumerate(zip(labels, colors)):
        x = 360
        y = 100 + i * 34
        draw.rounded_rectangle((x, y, x + 18, y + 18), radius=3, fill=color)
        draw.text((x + 26, y), label, fill="#262626", font=font)

    img.save(path)
    print(f"  ✓ 饼图：{path}")


def gen_bar_chart(path):
    """年度费用分配柱状图。"""
    labels = ["服务器运维", "网络运维", "数据库运维", "桌面支持"]
    values = [120, 80, 90, 70]
    colors = ["#1677ff", "#13c2c2", "#52c41a", "#faad14"]

    w, h = 560, 340
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _get_font(18)
    font = _get_font(13)
    draw.text((20, 14), "图1  年度运维费用分配", fill=(33, 33, 33), font=title_font)

    left, top, right, bottom = 60, 70, 520, 270
    draw.line((left, bottom, right, bottom), fill="#d9d9d9", width=1)
    draw.line((left, top, left, bottom), fill="#d9d9d9", width=1)

    max_v = max(values)
    n = len(labels)
    bar_w = (right - left - 60) / n
    gap = 15
    for i, (label, value, color) in enumerate(zip(labels, values, colors)):
        bh = value / max_v * (bottom - top - 20)
        x0 = left + 30 + i * (bar_w + gap)
        y0 = bottom - bh
        x1 = x0 + bar_w
        draw.rounded_rectangle((x0, y0, x1, bottom), radius=4, fill=color)
        draw.text((x0 + 6, y0 - 22), f"¥{value}k", fill="#262626", font=font)
        bbox = draw.textbbox((0, 0), label, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x0 + (bar_w - tw) / 2, bottom + 8), label, fill="#595959", font=font)

    draw.text((left, top - 24), "金额（千元）", fill="#8c8c8c", font=font)
    img.save(path)
    print(f"  ✓ 柱状图：{path}")


def gen_donut_chart(path):
    """项目预算占比环形图。"""
    labels = ["硬件设备", "软件授权", "实施服务", "质保服务", "项目管理"]
    values = [42, 23, 18, 10, 7]
    colors = ["#1677ff", "#13c2c2", "#52c41a", "#faad14", "#eb2f96"]

    w, h = 600, 380
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    title_font = _get_font(18)
    font = _get_font(14)
    draw.text((20, 14), "图1  项目预算占比", fill=(33, 33, 33), font=title_font)

    cx, cy, r = 180, 200, 110
    total = sum(values)
    start = -90
    for value, color in zip(values, colors):
        angle = value / total * 360
        draw.pieslice((cx - r, cy - r, cx + r, cy + r), start, start + angle, fill=color)
        start += angle
    # 中心圆孔
    draw.ellipse((cx - 45, cy - 45, cx + 45, cy + 45), fill=(255, 255, 255))
    draw.text((cx - 22, cy - 10), "预算", fill="#262626", font=font)

    # 图例
    for i, (label, color, value) in enumerate(zip(labels, colors, values)):
        x = 380
        y = 100 + i * 34
        draw.rounded_rectangle((x, y, x + 18, y + 18), radius=3, fill=color)
        draw.text((x + 26, y), f"{label}  {value}%", fill="#262626", font=font)

    img.save(path)
    print(f"  ✓ 环形图：{path}")


def gen_flow_chart(path):
    """故障响应流程图。"""
    w, h = 640, 220
    img = Image.new("RGB", (w, h), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    font = _get_font(14)
    title_font = _get_font(18)
    draw.text((20, 10), "图2  故障响应流程", fill=(33, 33, 33), font=title_font)

    nodes = [
        (60, 80, "故障申报"),
        (200, 80, "工单分级"),
        (340, 80, "工程师处理"),
        (480, 80, "回访确认"),
    ]
    for x, y, text in nodes:
        draw.rounded_rectangle((x, y, x + 110, y + 50), radius=8, fill="#1677ff", outline="#0958d9", width=2)
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
        draw.text((x + (110 - tw) / 2, y + (50 - th) / 2 - bbox[1]), text, fill=(255, 255, 255), font=font)

    # 箭头
    for i in range(len(nodes) - 1):
        x1 = nodes[i][0] + 110
        y1 = nodes[i][1] + 25
        x2 = nodes[i + 1][0]
        y2 = nodes[i + 1][1] + 25
        draw.line((x1, y1, x2, y2), fill="#8c8c8c", width=2)
        draw.polygon([(x2 - 8, y2 - 5), (x2, y2), (x2 - 8, y2 + 5)], fill="#8c8c8c")

    #  SLA 标注
    slas = ["5分钟", "15分钟", "30分钟", "闭环"]
    for i, sla in enumerate(slas):
        x = nodes[i][0] + 55
        y = nodes[i][1] + 60
        bbox = draw.textbbox((0, 0), sla, font=font)
        tw = bbox[2] - bbox[0]
        draw.text((x - tw / 2, y), sla, fill="#595959", font=font)

    img.save(path)
    print(f"  ✓ 流程图：{path}")


# =====================================================================
# 二、4 份合同数据
# =====================================================================

CONTRACTS = [
    {
        "key": "01_软件采购合同_含图表",
        "title": "企业ERP系统采购合同",
        "contract_no": "HT-YZ-2026-ERP-001",
        "party_b": "星河软件有限公司",
        "amount_words": "人民币伍拾捌万元整",
        "amount_number": "¥580,000.00",
        "sections": [
            ("第一条  合同主体", [
                f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
                "乙方（供应方）：星河软件有限公司，法定代表人：刘星河，统一社会信用代码：91440101MA5XXXXXX2，注册地址：上海市浦东新区张江高科技园区博云路2号。",
            ]),
            ("第二条  合同标的", [
                "甲方向乙方采购企业资源计划（ERP）管理系统一套，包含系统软件许可、数据库配置、接口开发及三年运维服务。具体功能模块与交付内容详见本合同附件。",
            ]),
            ("第三条  合同金额与付款方式", [
                "本合同总金额为人民币伍拾捌万元整（¥580,000.00），已含增值税。甲方按附表一《付款计划表》约定节点向乙方支付。",
            ]),
            ("第四条  交付安排", [
                "乙方应于合同生效后 90 个自然日内完成 ERP 系统的需求分析、系统设计、开发实现、测试验收及上线部署，具体进度安排见图1。",
            ]),
            ("第五条  验收标准", [
                "系统交付后，甲方应在 15 个工作日内组织验收。验收以双方确认的《需求规格说明书》为依据，主要功能模块全部通过测试并签署《验收报告》后视为验收合格。",
            ]),
            ("第六条  知识产权", [
                "本项目定制开发的源代码及配套技术文档的知识产权归甲方所有；乙方对标准化产品模块保留所有权并授予甲方永久使用许可。",
            ]),
            ("第七条  保密条款", [
                "双方对合作中知悉的对方商业秘密负保密义务，保密期限自合同生效起 5 年，未经书面同意不得向第三方披露。",
            ]),
            ("第八条  违约责任", [
                "1. 甲方逾期付款的，每日按应付未付金额的 0.3‰ 支付违约金；",
                "2. 乙方延期交付的，每日按合同总额的 0.5‰ 支付违约金，累计不超过合同总额的 10%。",
            ]),
            ("第九条  争议解决", [
                "因本合同引起的争议，由合同签订地（北京市海淀区）有管辖权的人民法院管辖。",
            ]),
        ],
        "tables": [
            ("附表一  付款计划表", [
                ["期次", "付款节点", "比例", "金额（元）", "备注"],
                ["1", "合同签订后 7 个工作日内", "30%", "174,000.00", "预付款"],
                ["2", "系统部署完成", "40%", "232,000.00", "进度款"],
                ["3", "验收合格签署报告", "25%", "145,000.00", "验收款"],
                ["4", "质保期满（12 个月）", "5%", "29,000.00", "质保金"],
                ["合计", "—", "100%", "580,000.00", "—"],
            ], [40, 130, 50, 90, 70]),
            ("附表二  交付清单表", [
                ["序号", "模块名称", "内容描述", "数量", "交付形式"],
                ["1", "财务模块", "总账/应收/应付/固定资产", "1 套", "软件许可+部署"],
                ["2", "供应链模块", "采购/库存/销售/供应商管理", "1 套", "软件许可+部署"],
                ["3", "生产模块", "生产计划/工单/物料清单", "1 套", "软件许可+部署"],
                ["4", "数据接口", "与 OA、CRM 系统对接", "3 个", "定制开发"],
                ["5", "运维服务", "系统巡检/故障处理/版本升级", "3 年", "远程+现场"],
            ], [35, 75, 160, 45, 75]),
        ],
        "charts": ["gantt"],
    },
    {
        "key": "02_硬件采购合同_含图表",
        "title": "服务器硬件采购合同",
        "contract_no": "HT-YZ-2026-HW-002",
        "party_b": "华信硬件设备股份有限公司",
        "amount_words": "人民币壹佰贰拾万元整",
        "amount_number": "¥1,200,000.00",
        "sections": [
            ("第一条  合同主体", [
                f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
                "乙方（供应方）：华信硬件设备股份有限公司，法定代表人：陈华信，统一社会信用代码：91440101MA5XXXXXX3，注册地址：深圳市南山区科技园南区。",
            ]),
            ("第二条  合同标的", [
                "甲方向乙方采购机架式服务器、存储阵列及网络设备一批，具体型号、配置、数量及单价详见附表一《设备配置清单表》。",
            ]),
            ("第三条  合同金额与付款方式", [
                "本合同总金额为人民币壹佰贰拾万元整（¥1,200,000.00），含增值税、运输及安装调试费用。合同签订后 10 个工作日内甲方支付 30% 预付款，货到验收合格后支付 65%，剩余 5% 作为质保金于质保期满后支付。",
            ]),
            ("第四条  交付安排", [
                "乙方应于合同生效后 30 个自然日内将全部设备送达甲方指定地点并完成安装调试。",
            ]),
            ("第五条  验收标准", [
                "设备到货后甲方应在 10 个工作日内完成验收，验收标准详见附表二。设备型号、数量、配置与清单一致，通电运行 72 小时无故障视为初验合格。",
            ]),
            ("第六条  质保服务", [
                "乙方提供三年原厂质保服务，质保期内设备出现非人为故障，乙方应在 48 小时内响应并提供免费维修或更换服务。",
            ]),
            ("第七条  违约责任", [
                "1. 乙方所交设备不符合约定的，应免费更换或退货并承担合同总额 5% 的违约金；",
                "2. 甲方无正当理由拒收的，应承担合同总额 5% 的违约金。",
            ]),
            ("第八条  争议解决", [
                "本合同争议由甲方所在地（北京市海淀区）人民法院管辖。",
            ]),
        ],
        "tables": [
            ("附表一  设备配置清单表", [
                ["序号", "设备名称", "型号/配置", "数量", "单价（元）", "金额（元）"],
                ["1", "机架式服务器", "2U/64核/512G/8×1.92T SSD", "8", "85,000.00", "680,000.00"],
                ["2", "全闪存储阵列", "双控/64TB可用容量", "1", "240,000.00", "240,000.00"],
                ["3", "万兆交换机", "48口/6×100G上联", "2", "72,000.00", "144,000.00"],
                ["4", "光模块及线缆", "10G/25G多模", "40", "350.00", "14,000.00"],
                ["5", "实施服务", "上架/布线/联调", "1", "122,000.00", "122,000.00"],
                ["合计", "—", "—", "—", "—", "1,200,000.00"],
            ], [30, 80, 160, 40, 80, 80]),
            ("附表二  验收标准表", [
                ["验收项", "验收内容", "合格标准", "检验方式"],
                ["外观", "设备包装、外观、标识", "无破损、型号一致", "目视检查"],
                ["配置", "CPU/内存/硬盘/网卡", "与清单一致", "开机自检+软件检测"],
                ["稳定性", "连续运行", "72小时无故障", "通电烤机"],
                ["网络", "交换机端口、链路", "全端口速率达标", "流量测试"],
            ], [70, 140, 120, 120]),
        ],
        "charts": ["pie"],
    },
    {
        "key": "03_运维服务合同_含图表",
        "title": "IT运维服务采购合同",
        "contract_no": "HT-YZ-2026-SVC-003",
        "party_b": "恒通信息技术服务有限公司",
        "amount_words": "人民币叁拾陆万元整",
        "amount_number": "¥360,000.00",
        "sections": [
            ("第一条  合同主体", [
                f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
                "乙方（供应方）：恒通信息技术服务有限公司，法定代表人：王恒通，统一社会信用代码：91440101MA5XXXXXX4，注册地址：广州市天河区珠江新城华夏路30号。",
            ]),
            ("第二条  合同标的", [
                "乙方为甲方提供为期一年的 IT 基础设施运维服务，服务范围包括服务器、网络、数据库及桌面终端的日常运维、故障处理与巡检优化。",
            ]),
            ("第三条  合同金额与付款方式", [
                "本合同总金额为人民币叁拾陆万元整（¥360,000.00），含增值税。按季度平均支付，每季度初 10 个工作日内支付当季服务费 ¥90,000.00。",
            ]),
            ("第四条  服务期限", [
                "服务期自 2026 年 8 月 1 日起至 2027 年 7 月 31 日止，乙方提供 7×24 小时远程支持与工作日现场支持。",
            ]),
            ("第五条  服务标准", [
                "乙方承诺的服务可用率、响应时间、解决时间等 SLA 指标详见附表一。月度系统可用率不低于 99.5%。",
            ]),
            ("第六条  知识产权", [
                "乙方在服务过程中形成的运维文档、巡检报告归甲方所有；乙方自有工具与脚本的知识产权归乙方所有。",
            ]),
            ("第七条  保密条款", [
                "乙方接触到的甲方业务数据、系统配置信息属甲方机密，乙方须严格保密，保密期限为合同终止后 5 年。",
            ]),
            ("第八条  违约责任", [
                "1. 乙方未达服务标准的，按当季服务费 5%—20% 支付违约金；",
                "2. 甲方逾期付款的，每日按 0.3‰ 支付违约金。",
            ]),
            ("第九条  争议解决", [
                "本合同争议由合同履行地（北京市海淀区）人民法院管辖。",
            ]),
        ],
        "tables": [
            ("附表一  SLA服务指标表", [
                ["服务项", "可用率", "响应时间", "解决时间", "服务时段"],
                ["服务器运维", "≥99.5%", "≤15分钟", "≤4小时", "7×24"],
                ["网络运维", "≥99.5%", "≤15分钟", "≤4小时", "7×24"],
                ["数据库运维", "≥99.9%", "≤10分钟", "≤2小时", "7×24"],
                ["桌面支持", "—", "≤30分钟", "≤1工作日", "工作日现场"],
            ], [90, 70, 80, 80, 90]),
            ("附表二  服务内容与报价表", [
                ["序号", "服务项目", "服务内容", "金额（元/年）"],
                ["1", "服务器运维", "20 台物理机巡检、补丁、故障处理", "120,000.00"],
                ["2", "网络运维", "交换机/防火墙/VPN 配置与监控", "80,000.00"],
                ["3", "数据库运维", "Oracle/MySQL 备份、优化、容灾", "90,000.00"],
                ["4", "桌面支持", "300 终端办公电脑软硬件支持", "70,000.00"],
                ["合计", "—", "—", "360,000.00"],
            ], [40, 90, 260, 90]),
        ],
        "charts": ["bar", "flow"],
    },
    {
        "key": "04_系统集成合同_含图表",
        "title": "智慧园区系统集成合同",
        "contract_no": "HT-YZ-2026-INT-004",
        "party_b": "中科系统集成有限公司",
        "amount_words": "人民币贰佰捌拾万元整",
        "amount_number": "¥2,800,000.00",
        "sections": [
            ("第一条  合同主体", [
                f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
                "乙方（供应方）：中科系统集成有限公司，法定代表人：李中科，统一社会信用代码：91440101MA5XXXXXX5，注册地址：武汉市东湖新技术开发区光谷大道1号。",
            ]),
            ("第二条  合同标的", [
                "乙方承担甲方智慧园区综合管理平台系统集成项目，包含视频监控、门禁考勤、能耗管理、停车管理四大子系统的设计、开发、部署与联调。",
            ]),
            ("第三条  合同金额与付款方式", [
                "本合同总金额为人民币贰佰捌拾万元整（¥2,800,000.00），含增值税。甲方按附表一《付款计划表》约定节点向乙方支付。",
            ]),
            ("第四条  交付安排", [
                "项目总工期 120 个自然日，分四个里程碑阶段交付，各阶段成果经甲方验收确认后进入下一阶段，详见附表二。",
            ]),
            ("第五条  验收标准", [
                "各子系统功能按《技术方案书》逐项验收，整体联调完成后进行为期 30 天的试运行，试运行期间无重大故障且各项指标达标后签署最终验收报告。",
            ]),
            ("第六条  知识产权", [
                "本项目定制开发的平台源代码、接口文档及设计资料的知识产权归甲方所有；乙方提供的通用中间件、硬件驱动等保留所有权并授予甲方使用许可。",
            ]),
            ("第七条  保密条款", [
                "双方对项目涉及的技术方案、园区数据、安防信息等承担严格保密义务，保密期限自合同生效起 5 年。",
            ]),
            ("第八条  违约责任", [
                "1. 乙方逾期交付里程碑成果的，每日按合同总额 0.3‰ 支付违约金，累计不超过合同总额 10%；",
                "2. 甲方逾期付款的，每日按应付金额 0.3‰ 支付违约金。",
            ]),
            ("第九条  争议解决", [
                "本合同争议由合同签订地（北京市海淀区）人民法院管辖。",
            ]),
        ],
        "tables": [
            ("附表一  付款计划表", [
                ["期次", "付款节点", "比例", "金额（元）", "备注"],
                ["1", "合同签订后 10 个工作日内", "20%", "560,000.00", "预付款"],
                ["2", "设计方案评审通过", "20%", "560,000.00", "设计款"],
                ["3", "各子系统部署完成", "30%", "840,000.00", "进度款"],
                ["4", "整体试运行结束验收", "25%", "700,000.00", "验收款"],
                ["5", "质保期满（24 个月）", "5%", "140,000.00", "质保金"],
                ["合计", "—", "100%", "2,800,000.00", "—"],
            ], [40, 130, 50, 90, 70]),
            ("附表二  项目里程碑表", [
                ["阶段", "里程碑", "交付物", "工期（天）", "完成节点"],
                ["一", "需求调研与方案设计", "需求规格说明书、技术方案书", "20", "T+20"],
                ["二", "子系统开发与部署", "视频/门禁/能耗/停车四子系统", "50", "T+70"],
                ["三", "系统集成与联调", "综合管理平台、接口联调报告", "30", "T+100"],
                ["四", "试运行与最终验收", "试运行报告、最终验收报告", "20", "T+120"],
            ], [35, 100, 200, 55, 55]),
        ],
        "charts": ["donut"],
    },
]


# =====================================================================
# 三、PDF 生成（reportlab）
# =====================================================================

def _register_pdf_font():
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    fp = CN_FONT_PATH
    if fp:
        try:
            pdfmetrics.registerFont(TTFont("CNDoc", fp))
            return "CNDoc"
        except Exception as e:
            print(f"  ⚠ 注册 PDF 字体失败 {fp}: {e}")
    print("  ⚠ 未找到中文字体，PDF 中文可能显示异常。")
    return "Helvetica"


def _pdf_styles(cn_font):
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("CN_Title", parent=styles["Title"],
                                fontName=cn_font, fontSize=20, leading=28, spaceAfter=16, alignment=1),
        "h1": ParagraphStyle("CN_H1", parent=styles["Heading1"],
                             fontName=cn_font, fontSize=13, leading=18, spaceAfter=8, spaceBefore=12),
        "normal": ParagraphStyle("CN_Normal", parent=styles["Normal"],
                                 fontName=cn_font, fontSize=10.5, leading=17, spaceAfter=8),
        "cell": ParagraphStyle("CN_Cell", parent=styles["Normal"],
                               fontName=cn_font, fontSize=9, leading=13),
        "cell_head": ParagraphStyle("CN_CellHead", parent=styles["Normal"],
                                    fontName=cn_font, fontSize=9, leading=13,
                                    textColor=colors.white, alignment=1),
        "small": ParagraphStyle("CN_Small", parent=styles["Normal"],
                                fontName=cn_font, fontSize=8, leading=12, textColor=colors.grey),
        "info": ParagraphStyle("CN_Info", parent=styles["Normal"],
                               fontName=cn_font, fontSize=10, leading=15, spaceAfter=4),
    }


def _pdf_table(table_data, col_widths, cn_font):
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib import colors

    cell_style = ParagraphStyle("c", fontName=cn_font, fontSize=9, leading=13)
    head_style = ParagraphStyle("h", fontName=cn_font, fontSize=9, leading=13,
                                textColor=colors.white, alignment=1)
    rows = []
    for r, row in enumerate(table_data):
        rows.append([Paragraph(str(c), head_style if r == 0 else cell_style) for c in row])
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1677ff")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bfbfbf")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("PADDING", (0, 0), (-1, -1), 4),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
    ]))
    return t


def _build_pdf_header(story, contract, styles, logo_path=None):
    from reportlab.platypus import Paragraph, Spacer, Image as RLImage
    from reportlab.lib.units import mm
    if logo_path and os.path.exists(logo_path):
        story.append(RLImage(logo_path, width=45 * mm, height=13.5 * mm))
        story.append(Spacer(1, 3 * mm))
    story.append(Paragraph(contract["title"], styles["title"]))
    story.append(Paragraph(f"合同编号：{contract['contract_no']}", styles["info"]))
    story.append(Paragraph(f"甲方（采购方）：{PARTY_A}", styles["info"]))
    story.append(Paragraph(f"乙方（供应方）：{contract['party_b']}", styles["info"]))
    story.append(Paragraph(f"合同金额：{contract['amount_words']}（{contract['amount_number']}）", styles["info"]))
    story.append(Spacer(1, 4 * mm))


def _build_pdf_sections(story, contract, styles):
    from reportlab.platypus import Paragraph
    for title, paras in contract["sections"]:
        story.append(Paragraph(title, styles["h1"]))
        for p in paras:
            story.append(Paragraph(p, styles["normal"]))


def _build_pdf_signature(story, styles, stamp_path=None):
    from reportlab.platypus import Paragraph, Spacer, Image as RLImage
    from reportlab.lib.units import mm
    story.append(Spacer(1, 8 * mm))
    story.append(Paragraph("甲方（盖章）：智远科技有限公司          乙方（盖章）：供应方", styles["normal"]))
    story.append(Paragraph("授权代表：                        授权代表：", styles["normal"]))
    story.append(Paragraph("签订日期：    年    月    日       签订日期：    年    月    日", styles["normal"]))
    if stamp_path and os.path.exists(stamp_path):
        story.append(Spacer(1, 3 * mm))
        story.append(RLImage(stamp_path, width=28 * mm, height=28 * mm))
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("（本合同样例由系统自动生成，用于导入测试）", styles["small"]))


def generate_pdf(contract, output_path, resources):
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    cn_font = _register_pdf_font()
    styles = _pdf_styles(cn_font)

    doc = SimpleDocTemplate(
        output_path, pagesize=A4, title=contract["title"], author="契审智控",
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=18 * mm, bottomMargin=16 * mm,
    )
    story = []
    _build_pdf_header(story, contract, styles, resources.get("logo"))
    _build_pdf_sections(story, contract, styles)

    for title, data, widths in contract["tables"]:
        story.append(Spacer(1, 4 * mm))
        story.append(Paragraph(title, styles["h1"]))
        story.append(_pdf_table(data, widths, cn_font))

    for chart_key in contract["charts"]:
        chart_path = resources.get(chart_key)
        if chart_path and os.path.exists(chart_path):
            story.append(Spacer(1, 5 * mm))
            img = RLImage(chart_path, width=140 * mm, height=_chart_height(chart_key) * mm)
            story.append(img)

    _build_pdf_signature(story, styles, resources.get("stamp"))
    doc.build(story)
    print(f"  ✅ PDF：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


def _chart_height(chart_key):
    return {"gantt": 55, "pie": 55, "bar": 52, "donut": 55, "flow": 40}.get(chart_key, 50)


# =====================================================================
# 四、DOCX 生成（python-docx）
# =====================================================================

def _set_cn_font(run, font_name="微软雅黑", size=None):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = size


def _docx_add_heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        _set_cn_font(run, "微软雅黑", Pt(20))
        p.alignment = 1
    else:
        _set_cn_font(run, "微软雅黑", Pt(13))
        run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
    return p


def _docx_add_para(doc, text):
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, "微软雅黑", Pt(10.5))
    return p


def _docx_add_info(doc, contract):
    from docx.shared import Pt
    for text in [
        f"合同编号：{contract['contract_no']}",
        f"甲方（采购方）：{PARTY_A}",
        f"乙方（供应方）：{contract['party_b']}",
        f"合同金额：{contract['amount_words']}（{contract['amount_number']}）",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_cn_font(run, "微软雅黑", Pt(10))
    doc.add_paragraph()


def _docx_add_table(doc, data, col_widths=None):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = 1
            run = p.add_run(str(data[r][c]))
            _set_cn_font(run, "微软雅黑", Pt(9))
            if r == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "1677FF")
                tcPr.append(shd)

    if col_widths:
        for c, w in enumerate(col_widths):
            for r in range(rows):
                table.cell(r, c).width = Cm(w)
    return table


def _docx_add_signature(doc, stamp_path=None):
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc.add_paragraph()
    _docx_add_para(doc, "甲方（盖章）：智远科技有限公司          乙方（盖章）：供应方")
    _docx_add_para(doc, "授权代表：                        授权代表：")
    _docx_add_para(doc, "签订日期：    年    月    日       签订日期：    年    月    日")
    if stamp_path and os.path.exists(stamp_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(stamp_path, width=Cm(3))
    p = doc.add_paragraph()
    run = p.add_run("（本合同样例由系统自动生成，用于导入测试）")
    _set_cn_font(run, "微软雅黑", Pt(8))


def generate_docx(contract, output_path, resources):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    logo_path = resources.get("logo")
    if logo_path and os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(5))

    _docx_add_heading(doc, contract["title"], level=0)
    _docx_add_info(doc, contract)

    for title, paras in contract["sections"]:
        _docx_add_heading(doc, title, level=1)
        for p in paras:
            _docx_add_para(doc, p)

    for title, data, widths in contract["tables"]:
        _docx_add_heading(doc, title, level=1)
        _docx_add_table(doc, data, widths)

    for chart_key in contract["charts"]:
        chart_path = resources.get(chart_key)
        if chart_path and os.path.exists(chart_path):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(chart_path, width=Cm(14))

    _docx_add_signature(doc, resources.get("stamp"))
    doc.save(output_path)
    print(f"  ✅ DOCX：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


# =====================================================================
# 五、主流程
# =====================================================================

def main():
    print("=" * 60)
    print("  生成 4 份合同样例（DOCX + PDF，含图片/表格/图表）")
    print("=" * 60)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"输出目录：{OUTPUT_DIR}\n")

    print("[1/3] 生成图片与图表资源")
    resources = {
        "logo": os.path.join(OUTPUT_DIR, "_sample_logo.png"),
        "stamp": os.path.join(OUTPUT_DIR, "_sample_stamp.png"),
        "gantt": os.path.join(OUTPUT_DIR, "_sample_gantt.png"),
        "pie": os.path.join(OUTPUT_DIR, "_sample_pie.png"),
        "bar": os.path.join(OUTPUT_DIR, "_sample_bar.png"),
        "donut": os.path.join(OUTPUT_DIR, "_sample_donut.png"),
        "flow": os.path.join(OUTPUT_DIR, "_sample_flow.png"),
    }
    gen_logo_image(resources["logo"])
    gen_stamp_image(resources["stamp"])
    gen_gantt_chart(resources["gantt"])
    gen_pie_chart(resources["pie"])
    gen_bar_chart(resources["bar"])
    gen_donut_chart(resources["donut"])
    gen_flow_chart(resources["flow"])

    print("\n[2/3] 生成 4 份合同文件")
    results = []
    for contract in CONTRACTS:
        docx_path = os.path.join(OUTPUT_DIR, f"{contract['key']}.docx")
        pdf_path = os.path.join(OUTPUT_DIR, f"{contract['key']}.pdf")
        generate_docx(contract, docx_path, resources)
        generate_pdf(contract, pdf_path, resources)
        results.append((contract["key"], docx_path, pdf_path))

    print("\n[3/3] 生成结果汇总")
    print("-" * 70)
    all_ok = True
    for key, docx_path, pdf_path in results:
        docx_ok = os.path.exists(docx_path)
        pdf_ok = os.path.exists(pdf_path)
        docx_size = os.path.getsize(docx_path) / 1024 if docx_ok else 0
        pdf_size = os.path.getsize(pdf_path) / 1024 if pdf_ok else 0
        status = "✅" if docx_ok and pdf_ok else "⚠"
        print(f"  {status} {key}")
        print(f"      DOCX: {docx_size:7.1f} KB  {'✓' if docx_ok else '✗'}")
        print(f"      PDF : {pdf_size:7.1f} KB  {'✓' if pdf_ok else '✗'}")
        if not (docx_ok and pdf_ok):
            all_ok = False
    print("-" * 70)
    if all_ok:
        print("✅ 全部 4 份合同样例（DOCX + PDF）已成功生成。")
    else:
        print("⚠ 部分文件未生成，请检查上方日志。")


if __name__ == "__main__":
    main()
