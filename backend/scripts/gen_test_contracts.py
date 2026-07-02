"""生成带表格/图片的测试合同文件

用于验证合同审核系统的「表格提取」与「图片提取」能力。

生成物（输出到 backend/test_contracts/）：
  1. 测试合同_含表格.pdf        —— 含 2 个表格（付款计划表、交付清单表），无图片
  2. 测试合同_含图片.pdf        —— 含 2 个图片（公司 logo、签署盖章图），无表格
  3. 测试合同_表格图片.docx     —— 含 1 个表格 + 1 个图片
  4. 测试合同_综合版.docx       —— 含 2 个表格 + 2 个图片

用法：
    cd backend
    .\\venv\\Scripts\\python.exe scripts\\gen_test_contracts.py
"""
import os
from PIL import Image, ImageDraw, ImageFont

# reportlab 与 python-docx 在 __main__ 中按需导入，避免顶部导入失败影响整体

# ===== 路径常量 =====
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.normpath(os.path.join(SCRIPT_DIR, ".."))
OUTPUT_DIR = os.path.join(BACKEND_DIR, "test_contracts")

# 临时图片输出到 test_contracts 目录下
LOGO_PATH = os.path.join(OUTPUT_DIR, "_logo_zhiyuan.png")
STAMP_PATH = os.path.join(OUTPUT_DIR, "_stamp_contract.png")

# ===== 中文字体（Pillow / reportlab 共用） =====
CN_FONT_CANDIDATES = [
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\msyhbd.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\msyh.ttf",
    r"C:\Windows\Fonts\Deng.ttf",
    r"C:\Windows\Fonts\SIMKAI.ttf",
]


def _find_cn_font_path():
    """返回第一个可用的中文字体路径，找不到返回 None。"""
    for fp in CN_FONT_CANDIDATES:
        if os.path.exists(fp):
            return fp
    return None


# =====================================================================
# 一、占位图片生成（Pillow）
# =====================================================================

def gen_logo_image(path):
    """生成公司 logo：蓝色矩形背景 + 白色文字「智远科技」，尺寸约 200x60。"""
    w, h = 200, 60
    img = Image.new("RGB", (w, h), color=(22, 119, 255))  # 主题蓝 #1677ff
    draw = ImageDraw.Draw(img)

    font_path = _find_cn_font_path()
    font = None
    if font_path:
        for size in (28, 26, 24, 22):
            try:
                font = ImageFont.truetype(font_path, size)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()

    text = "智远科技"
    try:
        bbox = draw.textbbox((0, 0), text, font=font)
        tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    except AttributeError:
        tw, th = font.getsize(text)
    tx = (w - tw) // 2
    ty = (h - th) // 2 - bbox[1] if font_path else (h - th) // 2
    draw.text((tx, ty), text, fill=(255, 255, 255), font=font)
    img.save(path)
    print(f"  ✓ logo 图片已生成：{path}（{os.path.getsize(path) / 1024:.1f} KB）")


def gen_stamp_image(path):
    """生成签署盖章图：红色圆形 + 「合同专用章」文字，尺寸约 120x120。"""
    size = 120
    img = Image.new("RGBA", (size, size), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)

    # 红色圆形（环 + 浅填充），模拟印章
    red = (245, 34, 45, 255)  # #f5222d
    draw.ellipse((6, 6, size - 6, size - 6), outline=red, width=4)

    font_path = _find_cn_font_path()
    font = None
    if font_path:
        for s in (18, 16, 14):
            try:
                font = ImageFont.truetype(font_path, s)
                break
            except Exception:
                continue
    if font is None:
        font = ImageFont.load_default()

    # 中央五角星
    cx, cy = size // 2, size // 2 - 22
    r = 9
    import math
    star_pts = []
    for i in range(10):
        ang = -math.pi / 2 + i * math.pi / 5
        rr = r if i % 2 == 0 else r * 0.4
        star_pts.append((cx + rr * math.cos(ang), cy + rr * math.sin(ang)))
    draw.polygon(star_pts, fill=red)

    # 文字「合同专用章」竖排两行：合同 / 专用章
    lines = ["合同", "专用章"]
    if font_path:
        try:
            line_h = draw.textbbox((0, 0), "合同", font=font)
            lh = line_h[3] - line_h[1]
        except AttributeError:
            lh = 18
    else:
        lh = 18
    y0 = cy + 12
    for idx, line in enumerate(lines):
        try:
            bbox = draw.textbbox((0, 0), line, font=font)
            lw = bbox[2] - bbox[0]
        except AttributeError:
            lw = font.getsize(line)[0]
        x = (size - lw) // 2
        y = y0 + idx * (lh + 4)
        draw.text((x, y), line, fill=red, font=font)

    img.save(path)
    print(f"  ✓ 印章图片已生成：{path}（{os.path.getsize(path) / 1024:.1f} KB）")


# =====================================================================
# 二、合同正文内容（4 份合同各自一套完整条款）
# =====================================================================

# 通用甲方
PARTY_A = "智远科技有限公司"

# ---- 合同1：软件采购（ERP 系统），含表格 ----
CONTRACT_TABLES = {
    "title": "企业ERP系统采购合同",
    "contract_no": "HT-YZ-2026-ERP-001",
    "party_b": "星河软件有限公司",
    "amount_words": "人民币伍拾捌万元整",
    "amount_number": "¥580,000.00",
    "sections": [
        ("第一条  合同主体", [
            f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
            f"乙方（供应方）：星河软件有限公司，法定代表人：刘星河，统一社会信用代码：91440101MA5XXXXXX2，注册地址：上海市浦东新区张江高科技园区博云路2号。",
        ]),
        ("第二条  合同标的", [
            "甲方向乙方采购企业资源计划（ERP）管理系统一套，包含系统软件许可、数据库配置、接口开发及三年运维服务，具体功能模块详见下表交付清单。",
        ]),
        ("第三条  合同金额与付款方式", [
            "本合同总金额为人民币伍拾捌万元整（¥580,000.00），已含增值税。甲方按以下付款计划向乙方支付：",
        ]),
        ("第四条  交付安排", [
            "乙方应于合同生效后 60 个自然日内完成 ERP 系统的交付、安装部署与初步调试，交付清单如下表所示。",
        ]),
        ("第五条  验收标准", [
            "系统交付后，甲方应在 15 个工作日内组织验收。验收以双方确认的《需求规格说明书》为依据，主要功能模块全部通过测试并签署《验收报告》后视为验收合格。",
        ]),
        ("第六条  知识产权", [
            "本项目定制开发的源代码及配套技术文档的知识产权归甲乙双方共同所有，乙方对标准化产品模块保留所有权并授予甲方永久使用许可。",
        ]),
        ("第七条  保密条款", [
            "双方对合作中知悉的对方商业秘密负保密义务，保密期限自合同生效起 5 年，未经书面同意不得向第三方披露。",
        ]),
        ("第八条  违约责任", [
            "1. 甲方逾期付款的，每日按应付未付金额的 0.3‰ 支付违约金；",
            "2. 乙方延期交付的，每日按合同总额的 0.5‰ 支付违约金，累计不超过合同总额的 10%。",
        ]),
        ("第九条  争议解决", [
            "因本合同引起的争议，由合同签订地（北京市海淀区）有管辖权的人民法院管辖。",
        ]),
    ],
}

# ---- 合同2：硬件采购（服务器），含图片 ----
CONTRACT_IMAGES = {
    "title": "服务器硬件采购合同",
    "contract_no": "HT-YZ-2026-HW-002",
    "party_b": "华信硬件设备股份有限公司",
    "amount_words": "人民币壹佰贰拾万元整",
    "amount_number": "¥1,200,000.00",
    "sections": [
        ("第一条  合同主体", [
            f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
            "乙方（供应方）：华信硬件设备股份有限公司，法定代表人：陈华信，统一社会信用代码：91440101MA5XXXXXX3，注册地址：深圳市南山区科技园南区。",
        ]),
        ("第二条  合同标的", [
            "甲方向乙方采购机架式服务器 20 台及配套存储设备 1 套，型号与配置以附件《硬件配置清单》为准。",
        ]),
        ("第三条  合同金额与付款方式", [
            "本合同总金额为人民币壹佰贰拾万元整（¥1,200,000.00），含增值税与运输安装费用。合同签订后 10 个工作日内甲方支付 30% 预付款，货到验收合格后支付 65%，剩余 5% 作为质保金于质保期满后支付。",
        ]),
        ("第四条  交付安排", [
            "乙方应于合同生效后 30 个自然日内将全部设备送达甲方指定地点（北京市海淀区中关村大街1号）并完成安装调试。",
        ]),
        ("第五条  验收标准", [
            "设备到货后甲方应在 10 个工作日内完成验收，验收标准为：设备型号、数量、配置与清单一致，通电运行 72 小时无故障。",
        ]),
        ("第六条  知识产权", [
            "硬件设备所附带的系统软件及管理工具，乙方授予甲方非独占的、不可转授权的内部使用权，相关知识产权归乙方或其授权方所有。",
        ]),
        ("第七条  保密条款", [
            "双方对在合作中获知的对方技术资料、商业数据承担保密义务，保密期限为合同终止后 3 年。",
        ]),
        ("第八条  违约责任", [
            "1. 乙方所交设备不符合约定的，应免费更换或退货并承担合同总额 5% 的违约金；",
            "2. 甲方无正当理由拒收的，应承担合同总额 5% 的违约金。",
        ]),
        ("第九条  争议解决", [
            "本合同争议由甲方所在地（北京市海淀区）人民法院管辖。",
        ]),
    ],
}

# ---- 合同3：服务采购（IT 运维），含 1 表格 + 1 图片 ----
CONTRACT_TABLE_IMAGE = {
    "title": "IT运维服务采购合同",
    "contract_no": "HT-YZ-2026-SVC-003",
    "party_b": "恒通信息技术服务有限公司",
    "amount_words": "人民币叁拾陆万元整",
    "amount_number": "¥360,000.00",
    "sections": [
        ("第一条  合同主体", [
            f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
            "乙方（供应方）：恒通信息技术服务有限公司，法定代表人：王恒通，统一社会信用代码：91440101MA5XXXXXX4，注册地址：广州市天河区珠江新城华夏路30号。",
        ]),
        ("第二条  合同标的", [
            "乙方为甲方提供为期一年的 IT 基础设施运维服务，服务内容与报价详见下表。",
        ]),
        ("第三条  合同金额与付款方式", [
            "本合同总金额为人民币叁拾陆万元整（¥360,000.00），含增值税。按季度平均支付，每季度初 10 个工作日内支付当季服务费 ¥90,000.00。",
        ]),
        ("第四条  交付安排", [
            "服务期自 2026 年 8 月 1 日起至 2027 年 7 月 31 日止，乙方提供 7×24 小时远程支持与工作日现场支持。",
        ]),
        ("第五条  验收标准", [
            "乙方按月提交《运维服务报告》，甲方对服务质量进行考核，系统可用率不低于 99.5%、故障响应时间不超过 30 分钟视为当月服务合格。",
        ]),
        ("第六条  知识产权", [
            "乙方在服务过程中形成的运维文档、巡检报告归甲方所有；乙方自有工具与脚本的知识产权归乙方所有。",
        ]),
        ("第七条  保密条款", [
            "乙方接触到的甲方业务数据、系统配置信息属甲方机密，乙方须严格保密，保密期限为合同终止后 5 年。",
        ]),
        ("第八条  违约责任", [
            "1. 乙方未达服务标准的，按当季服务费 5%—20% 支付违约金；",
            "2. 甲方逾期付款的，每日按 0.3‰ 支付违约金。",
        ]),
        ("第九条  争议解决", [
            "本合同争议由合同履行地（北京市海淀区）人民法院管辖。",
        ]),
    ],
}

# ---- 合同4：系统集成，含 2 表格 + 2 图片（综合版） ----
CONTRACT_COMPREHENSIVE = {
    "title": "智慧园区系统集成合同",
    "contract_no": "HT-YZ-2026-INT-004",
    "party_b": "中科系统集成有限公司",
    "amount_words": "人民币贰佰捌拾万元整",
    "amount_number": "¥2,800,000.00",
    "sections": [
        ("第一条  合同主体", [
            f"甲方（采购方）：{PARTY_A}，法定代表人：张智远，统一社会信用代码：91440101MA5XXXXXX1，注册地址：北京市海淀区中关村大街1号。",
            "乙方（供应方）：中科系统集成有限公司，法定代表人：李中科，统一社会信用代码：91440101MA5XXXXXX5，注册地址：武汉市东湖新技术开发区光谷大道1号。",
        ]),
        ("第二条  合同标的", [
            "乙方承担甲方智慧园区综合管理平台系统集成项目，包含视频监控、门禁考勤、能耗管理、停车管理四大子系统的设计、开发、部署与联调，项目里程碑详见下表。",
        ]),
        ("第三条  合同金额与付款方式", [
            "本合同总金额为人民币贰佰捌拾万元整（¥2,800,000.00），含增值税。甲方按以下付款计划支付：",
        ]),
        ("第四条  交付安排", [
            "项目总工期 120 个自然日，分四个里程碑阶段交付，各阶段成果经甲方验收确认后进入下一阶段。",
        ]),
        ("第五条  验收标准", [
            "各子系统功能按《技术方案书》逐项验收，整体联调完成后进行为期 30 天的试运行，试运行期间无重大故障且各项指标达标后签署最终验收报告。",
        ]),
        ("第六条  知识产权", [
            "本项目定制开发的平台源代码、接口文档及设计资料的知识产权归甲方所有；乙方提供的通用中间件、硬件驱动等保留所有权并授予甲方使用许可。",
        ]),
        ("第七条  保密条款", [
            "双方对项目涉及的技术方案、园区数据、安防信息等承担严格保密义务，保密期限自合同生效起 5 年。",
        ]),
        ("第八条  违约责任", [
            "1. 乙方逾期交付里程碑成果的，每日按合同总额 0.3‰ 支付违约金，累计不超过合同总额 10%；",
            "2. 甲方逾期付款的，每日按应付金额 0.3‰ 支付违约金。",
        ]),
        ("第九条  争议解决", [
            "本合同争议由合同签订地（北京市海淀区）人民法院管辖。",
        ]),
    ],
}


# =====================================================================
# 三、PDF 生成（reportlab）
# =====================================================================

def _register_pdf_font():
    """注册 reportlab 中文字体，返回字体名。"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    fp = _find_cn_font_path()
    if fp:
        try:
            pdfmetrics.registerFont(TTFont("CNDoc", fp))
            print(f"  ✓ 已注册 PDF 中文字体：{fp}")
            return "CNDoc"
        except Exception as e:
            print(f"  ⚠ 注册 PDF 字体失败 {fp}: {e}")
    print("  ⚠ 未找到中文字体，PDF 中文可能显示异常。")
    return "Helvetica"


def _pdf_styles(cn_font):
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors

    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("CN_Title", parent=styles["Title"],
                                fontName=cn_font, fontSize=20, leading=28, spaceAfter=18, alignment=1),
        "h1": ParagraphStyle("CN_H1", parent=styles["Heading1"],
                             fontName=cn_font, fontSize=13, leading=18, spaceAfter=8, spaceBefore=14),
        "normal": ParagraphStyle("CN_Normal", parent=styles["Normal"],
                                 fontName=cn_font, fontSize=10.5, leading=17, spaceAfter=8),
        "cell": ParagraphStyle("CN_Cell", parent=styles["Normal"],
                               fontName=cn_font, fontSize=9.5, leading=14),
        "cell_head": ParagraphStyle("CN_CellHead", parent=styles["Normal"],
                                    fontName=cn_font, fontSize=9.5, leading=14,
                                    textColor=colors.white, alignment=1),
        "small": ParagraphStyle("CN_Small", parent=styles["Normal"],
                                fontName=cn_font, fontSize=8, leading=12, textColor=colors.grey),
        "info": ParagraphStyle("CN_Info", parent=styles["Normal"],
                               fontName=cn_font, fontSize=10, leading=15, spaceAfter=4),
    }


def _pdf_table(table_data, col_widths, cn_font):
    """构造带表头的 reportlab Table。"""
    from reportlab.platypus import Table, TableStyle, Paragraph
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle

    cell_style = ParagraphStyle("c", fontName=cn_font, fontSize=9.5, leading=14)
    head_style = ParagraphStyle("h", fontName=cn_font, fontSize=9.5, leading=14,
                                textColor=colors.white, alignment=1)
    rows = []
    for r, row in enumerate(table_data):
        rows.append([Paragraph(str(c), head_style if r == 0 else cell_style) for c in row])
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1677ff")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bfbfbf")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("PADDING", (0, 0), (-1, -1), 5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f7fa")]),
    ]))
    return t


def _build_pdf_header(story, contract, styles):
    """构建 PDF 合同抬头（标题 + 编号/甲乙方信息段落）。"""
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import mm

    story.append(Paragraph(contract["title"], styles["title"]))
    story.append(Paragraph(f"合同编号：{contract['contract_no']}", styles["info"]))
    story.append(Paragraph(f"甲方（采购方）：{PARTY_A}", styles["info"]))
    story.append(Paragraph(f"乙方（供应方）：{contract['party_b']}", styles["info"]))
    story.append(Paragraph(
        f"合同金额：{contract['amount_words']}（{contract['amount_number']}）", styles["info"]))
    story.append(Spacer(1, 6 * mm))


def _build_pdf_sections(story, contract, styles):
    """构建 PDF 条款段落。"""
    from reportlab.platypus import Paragraph
    for title, paras in contract["sections"]:
        story.append(Paragraph(title, styles["h1"]))
        for p in paras:
            story.append(Paragraph(p, styles["normal"]))


def _build_pdf_signature(story, styles):
    from reportlab.platypus import Paragraph, Spacer
    from reportlab.lib.units import mm
    story.append(Spacer(1, 10 * mm))
    story.append(Paragraph(
        "甲方（盖章）：智远科技有限公司          乙方（盖章）：供应方", styles["normal"]))
    story.append(Paragraph("授权代表：                        授权代表：", styles["normal"]))
    story.append(Paragraph("签订日期：    年    月    日       签订日期：    年    月    日", styles["normal"]))
    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph(
        "（本测试合同由系统自动生成，用于验证表格/图片提取功能）", styles["small"]))


def gen_pdf_with_tables(output_path):
    """生成含 2 个表格的 PDF（付款计划表 + 交付清单表），无图片。"""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    cn_font = _register_pdf_font()
    styles = _pdf_styles(cn_font)
    contract = CONTRACT_TABLES

    doc = SimpleDocTemplate(
        output_path, pagesize=A4, title=contract["title"], author="契审智控",
        leftMargin=22 * mm, rightMargin=22 * mm, topMargin=20 * mm, bottomMargin=18 * mm,
    )
    story = []
    _build_pdf_header(story, contract, styles)
    _build_pdf_sections(story, contract, styles)

    # 在「第三条 合同金额与付款方式」之后插入付款计划表 —— 直接追加在条款之后
    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("附表一  付款计划表", styles["h1"]))
    pay_data = [
        ["期次", "付款节点", "比例", "金额（元）", "备注"],
        ["1", "合同签订后 7 个工作日内", "30%", "174,000.00", "预付款"],
        ["2", "系统交付部署完成", "40%", "232,000.00", "进度款"],
        ["3", "验收合格签署报告", "25%", "145,000.00", "验收款"],
        ["4", "质保期满（12 个月）", "5%", "29,000.00", "质保金"],
        ["合计", "—", "100%", "580,000.00", "—"],
    ]
    story.append(_pdf_table(pay_data, [40, 130, 50, 90, 70], cn_font))

    story.append(Spacer(1, 6 * mm))
    story.append(Paragraph("附表二  交付清单表", styles["h1"]))
    deliv_data = [
        ["序号", "模块名称", "内容描述", "数量", "交付形式"],
        ["1", "财务模块", "总账/应收/应付/固定资产", "1 套", "软件许可+部署"],
        ["2", "供应链模块", "采购/库存/销售/供应商管理", "1 套", "软件许可+部署"],
        ["3", "生产模块", "生产计划/工单/物料清单", "1 套", "软件许可+部署"],
        ["4", "数据接口", "与 OA、CRM 系统对接", "3 个", "定制开发"],
        ["5", "运维服务", "系统巡检/故障处理/版本升级", "3 年", "远程+现场"],
    ]
    story.append(_pdf_table(deliv_data, [35, 75, 160, 45, 75], cn_font))

    _build_pdf_signature(story, styles)
    doc.build(story)
    print(f"  ✅ PDF（含表格）已生成：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


def gen_pdf_with_images(output_path, logo_path, stamp_path):
    """生成含 2 个图片的 PDF（公司 logo + 签署盖章图），无表格。"""
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm

    cn_font = _register_pdf_font()
    styles = _pdf_styles(cn_font)
    contract = CONTRACT_IMAGES

    doc = SimpleDocTemplate(
        output_path, pagesize=A4, title=contract["title"], author="契审智控",
        leftMargin=22 * mm, rightMargin=22 * mm, topMargin=20 * mm, bottomMargin=18 * mm,
    )
    story = []

    # 图片1：公司 logo 置于标题上方
    logo_img = RLImage(logo_path, width=50 * mm, height=15 * mm)
    story.append(logo_img)
    story.append(Spacer(1, 4 * mm))

    _build_pdf_header(story, contract, styles)
    _build_pdf_sections(story, contract, styles)
    _build_pdf_signature(story, styles)

    # 图片2：签署盖章图置于落款右下方
    story.append(Spacer(1, 4 * mm))
    stamp_img = RLImage(stamp_path, width=30 * mm, height=30 * mm)
    story.append(Paragraph("（盖章处）", styles["small"]))
    story.append(stamp_img)

    doc.build(story)
    print(f"  ✅ PDF（含图片）已生成：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


# =====================================================================
# 四、DOCX 生成（python-docx）
# =====================================================================

def _set_cn_font(run, font_name="微软雅黑", size=None):
    """为 docx run 设置中文字体（需同时设置 w:eastAsia）。"""
    from docx.oxml.ns import qn
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = size


def _docx_add_heading(doc, text, level=1):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        _set_cn_font(run, "微软雅黑", Pt(20))
        p.alignment = 1  # 居中
    else:
        _set_cn_font(run, "微软雅黑", Pt(13))
        run.font.color.rgb = RGBColor(0x16, 0x77, 0xFF)
    return p


def _docx_add_para(doc, text):
    from docx.shared import Pt
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_cn_font(run, "微软雅黑", Pt(10.5))
    return p


def _docx_add_info(doc, contract):
    """添加合同编号/甲乙方/金额信息行。"""
    from docx.shared import Pt
    for text in [
        f"合同编号：{contract['contract_no']}",
        f"甲方（采购方）：{PARTY_A}",
        f"乙方（供应方）：{contract['party_b']}",
        f"合同金额：{contract['amount_words']}（{contract['amount_number']}）",
    ]:
        p = doc.add_paragraph()
        run = p.add_run(text)
        _set_cn_font(run, "微软雅黑", Pt(10))
    doc.add_paragraph()


def _docx_add_table(doc, data, col_widths=None):
    """添加带表头的表格，data[0] 为表头。"""
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r in range(rows):
        for c in range(cols):
            cell = table.cell(r, c)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = 1  # 居中
            run = p.add_run(str(data[r][c]))
            _set_cn_font(run, "微软雅黑", Pt(9.5))
            if r == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                # 表头底色
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:fill"), "1677FF")
                tcPr.append(shd)

    if col_widths:
        for c, w in enumerate(col_widths):
            for r in range(rows):
                table.cell(r, c).width = Cm(w)
    return table


def _docx_add_signature(doc):
    from docx.shared import Pt
    doc.add_paragraph()
    _docx_add_para(doc, "甲方（盖章）：智远科技有限公司          乙方（盖章）：供应方")
    _docx_add_para(doc, "授权代表：                        授权代表：")
    _docx_add_para(doc, "签订日期：    年    月    日       签订日期：    年    月    日")
    p = doc.add_paragraph()
    run = p.add_run("（本测试合同由系统自动生成，用于验证表格/图片提取功能）")
    _set_cn_font(run, "微软雅黑", Pt(8))


def _docx_build_sections(doc, contract):
    for title, paras in contract["sections"]:
        _docx_add_heading(doc, title, level=1)
        for p in paras:
            _docx_add_para(doc, p)


def gen_docx_table_image(output_path, logo_path):
    """生成含 1 个表格 + 1 个图片的 DOCX（IT 运维服务采购）。"""
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    contract = CONTRACT_TABLE_IMAGE
    doc = Document()

    # 图片1：公司 logo（居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(5))

    _docx_add_heading(doc, contract["title"], level=0)
    _docx_add_info(doc, contract)
    _docx_build_sections(doc, contract)

    # 表格：服务内容与报价表（1 个表格）
    _docx_add_heading(doc, "附表  服务内容与报价表", level=1)
    svc_data = [
        ["序号", "服务项目", "服务内容", "频次", "金额（元/年）"],
        ["1", "服务器运维", "20 台物理机巡检、补丁、故障处理", "7×24", "120,000.00"],
        ["2", "网络运维", "交换机/防火墙/VPN 配置与监控", "7×24", "80,000.00"],
        ["3", "数据库运维", "Oracle/MySQL 备份、优化、容灾", "工作日", "90,000.00"],
        ["4", "桌面支持", "300 终端办公电脑软硬件支持", "工作日现场", "70,000.00"],
        ["合计", "—", "—", "—", "360,000.00"],
    ]
    _docx_add_table(doc, svc_data, col_widths=[1.2, 2.5, 6.5, 2.5, 3.0])

    _docx_add_signature(doc)
    doc.save(output_path)
    print(f"  ✅ DOCX（表格+图片）已生成：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


def gen_docx_comprehensive(output_path, logo_path, stamp_path):
    """生成含 2 个表格 + 2 个图片的 DOCX 综合版（智慧园区系统集成）。"""
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    contract = CONTRACT_COMPREHENSIVE
    doc = Document()

    # 图片1：公司 logo（居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(logo_path, width=Cm(5))

    _docx_add_heading(doc, contract["title"], level=0)
    _docx_add_info(doc, contract)
    _docx_build_sections(doc, contract)

    # 表格1：付款计划表
    _docx_add_heading(doc, "附表一  付款计划表", level=1)
    pay_data = [
        ["期次", "付款节点", "比例", "金额（元）", "备注"],
        ["1", "合同签订后 10 个工作日内", "20%", "560,000.00", "预付款"],
        ["2", "设计方案评审通过", "20%", "560,000.00", "设计款"],
        ["3", "各子系统部署完成", "30%", "840,000.00", "进度款"],
        ["4", "整体试运行结束验收", "25%", "700,000.00", "验收款"],
        ["5", "质保期满（24 个月）", "5%", "140,000.00", "质保金"],
        ["合计", "—", "100%", "2,800,000.00", "—"],
    ]
    _docx_add_table(doc, pay_data, col_widths=[1.2, 4.5, 1.5, 3.0, 2.5])

    # 表格2：项目里程碑表
    _docx_add_heading(doc, "附表二  项目里程碑表", level=1)
    ms_data = [
        ["阶段", "里程碑", "交付物", "工期（天）", "完成节点"],
        ["一", "需求调研与方案设计", "需求规格说明书、技术方案书", "20", "T+20"],
        ["二", "子系统开发与部署", "视频/门禁/能耗/停车四子系统", "50", "T+70"],
        ["三", "系统集成与联调", "综合管理平台、接口联调报告", "30", "T+100"],
        ["四", "试运行与最终验收", "试运行报告、最终验收报告", "20", "T+120"],
    ]
    _docx_add_table(doc, ms_data, col_widths=[1.2, 3.5, 5.5, 2.0, 2.3])

    _docx_add_signature(doc)

    # 图片2：签署盖章图（居中）
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(stamp_path, width=Cm(3))

    doc.save(output_path)
    print(f"  ✅ DOCX（综合版）已生成：{output_path}（{os.path.getsize(output_path) / 1024:.1f} KB）")


# =====================================================================
# 五、主流程
# =====================================================================

def main():
    print("=" * 60)
    print("  生成带表格/图片的测试合同文件")
    print("=" * 60)

    # 1. 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"输出目录：{OUTPUT_DIR}")

    # 2. 生成临时占位图片
    print("\n[1/3] 生成占位图片")
    gen_logo_image(LOGO_PATH)
    gen_stamp_image(STAMP_PATH)

    # 3. 生成 4 份合同
    print("\n[2/3] 生成测试合同")
    gen_pdf_with_tables(os.path.join(OUTPUT_DIR, "测试合同_含表格.pdf"))
    gen_pdf_with_images(os.path.join(OUTPUT_DIR, "测试合同_含图片.pdf"), LOGO_PATH, STAMP_PATH)
    gen_docx_table_image(os.path.join(OUTPUT_DIR, "测试合同_表格图片.docx"), LOGO_PATH)
    gen_docx_comprehensive(os.path.join(OUTPUT_DIR, "测试合同_综合版.docx"), LOGO_PATH, STAMP_PATH)

    # 4. 汇总结果
    print("\n[3/3] 生成结果汇总")
    print("-" * 60)
    targets = [
        "测试合同_含表格.pdf",
        "测试合同_含图片.pdf",
        "测试合同_表格图片.docx",
        "测试合同_综合版.docx",
    ]
    all_ok = True
    for name in targets:
        fp = os.path.join(OUTPUT_DIR, name)
        if os.path.exists(fp):
            size_kb = os.path.getsize(fp) / 1024
            print(f"  ✓ {name:28s}  {size_kb:7.1f} KB")
        else:
            print(f"  ✗ {name:28s}  未生成！")
            all_ok = False
    print("-" * 60)
    if all_ok:
        print("✅ 全部 4 份测试合同已成功生成。")
    else:
        print("⚠ 部分文件未生成，请检查上方日志。")


if __name__ == "__main__":
    main()
